from __future__ import annotations

import importlib.util
import re
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

from resume_tailor.domain.models import (
    ClaimSupport,
    EducationRecord,
    EntityKind,
    GraduationStatus,
    ResumeItem,
    ResumeStrategy,
    StructuredBullet,
    StructuredResume,
    TechnicalSkillCategory,
)
from resume_tailor.infrastructure.adaptive_docx import render_structured_resume
from resume_tailor.infrastructure.rendering import (
    ManagedResumeRenderer,
    MicrosoftWordDocxPageCountProvider,
    PageCountMeasurement,
    PageCountVerificationError,
)
from resume_tailor.infrastructure.static_template_docx import render_template_v1_resume
from resume_tailor.infrastructure.template_v1 import (
    TEMPLATE_V1_DOCX_SHA256,
    TEMPLATE_V1_REFERENCE_SHA256,
    load_template_v1_layout_profile,
    template_v1_docx_path,
)

REFERENCE = Path("manual-test/reference-resume.docx")
PLACEHOLDER_PATTERN = re.compile(r"\{\{[A-Z0-9_]+\}\}")
SECTION_LABELS = {"Education", "Technical Skills", "Technical Experience", "Projects"}


def _strategy() -> ResumeStrategy:
    return ResumeStrategy(
        role_family="embedded_firmware",
        primary_focus="verified firmware evidence",
        rationale="Static template structural fixture.",
    )


def _bullet(identifier: str, text: str) -> StructuredBullet:
    return StructuredBullet(
        id=identifier,
        text=text,
        evidence_ids=[identifier],
        support=ClaimSupport.DIRECT,
    )


def _complete_resume() -> StructuredResume:
    first_experience = ResumeItem(
        id="experience-1",
        title="Firmware Engineer",
        subtitle="C, STM32",
        kind=EntityKind.EXPERIENCE,
        organization="Example Robotics",
        start_date="May 2024",
        end_date="Aug. 2025",
        location="Toronto, Ontario",
    )
    second_experience = ResumeItem(
        id="experience-2",
        title="Embedded Systems Engineer",
        technology_label="C++, RTOS",
        kind=EntityKind.EXPERIENCE,
        organization="Example Controls",
        start_date="Sep. 2023",
        end_date="Apr. 2024",
        location="Hamilton, Ontario",
    )
    first_project = ResumeItem(
        id="project-1",
        title="Sensor Platform",
        technology_label="Python, CAN",
        kind=EntityKind.PROJECT,
        organization="Example Design Team",
        start_date="Jan. 2024",
        end_date="Apr. 2024",
        location="Toronto, Ontario",
    )
    second_project = ResumeItem(
        id="project-2",
        title="Motor Controller",
        technologies=["C", "SPI"],
        award_or_placement="Finalist",
        kind=EntityKind.PROJECT,
        start_date="Sep. 2023",
        end_date="Dec. 2023",
    )
    return StructuredResume(
        profile_id="profile-static-template",
        profile_version=1,
        posting_id="posting-static-template",
        template_id="application-viego-resume-v1",
        display_name="Alex Example",
        contact_line=("alex@example.test | 555-0100 | Toronto, Ontario | example.test/portfolio"),
        strategy=_strategy(),
        entity_titles={
            first_experience.id: first_experience.title,
            second_experience.id: second_experience.title,
            first_project.id: first_project.title,
            second_project.id: second_project.title,
        },
        education=[
            EducationRecord(
                school="Example Polytechnic Institute",
                program="Bachelor of Applied Engineering",
                minor_or_specialization="Embedded Systems",
                start_date="Sep. 2021",
                expected_graduation_date="Apr. 2026",
                graduation_status=GraduationStatus.EXPECTED,
                location="Example City, ZZ",
                gpa="3.8/4.0",
                awards=["Example Merit Award"],
                relevant_coursework=["Control Systems", "Digital Logic"],
            )
        ],
        technical_skills=[
            TechnicalSkillCategory(
                category="Programming & Scripting",
                values=["C", "Python"],
            ),
            TechnicalSkillCategory(
                category="Embedded Systems",
                values=["STM32", "SPI"],
            ),
        ],
        experiences=[first_experience, second_experience],
        projects=[first_project, second_project],
        experience_bullets={
            first_experience.id: [
                _bullet(
                    "evidence-exp-1",
                    "Integrated verified embedded controls with deterministic timing.",
                ),
                _bullet(
                    "evidence-exp-2",
                    "Validated a verified sensor interface across documented operating modes.",
                ),
            ],
            second_experience.id: [
                _bullet(
                    "evidence-exp-3",
                    "Documented a verified firmware release and its test evidence.",
                ),
                _bullet(
                    "evidence-exp-4",
                    "Reduced verified diagnostic latency through measured interface changes.",
                ),
            ],
        },
        project_bullets={
            first_project.id: [
                _bullet(
                    "evidence-project-1",
                    "Built a verified sensor prototype from reviewed project evidence.",
                ),
                _bullet(
                    "evidence-project-2",
                    "Tested the verified prototype against documented requirements.",
                ),
            ],
            second_project.id: [
                _bullet(
                    "evidence-project-3",
                    "Implemented a verified motor-control interface using reviewed evidence.",
                ),
                _bullet(
                    "evidence-project-4",
                    "Measured the verified control response under documented test conditions.",
                ),
            ],
        },
        selected_coursework=["Control Systems", "Digital Logic"],
    )


def _paragraph(document: Document, prefix: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix))


def _canonical_xml(element: etree._Element | None) -> bytes:
    if element is None:
        return b""
    clone = deepcopy(element)
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for node in clone.iter():
        for attribute in list(node.attrib):
            namespace, _, local_name = attribute.rpartition("}")
            if namespace == f"{{{word_namespace}" and local_name.startswith("rsid"):
                del node.attrib[attribute]
    return etree.tostring(clone, method="c14n")


def _canonical_ppr_without_tabs(paragraph) -> bytes:
    properties = deepcopy(paragraph._p.pPr)
    if properties is None:
        return b""
    tabs = properties.find(qn("w:tabs"))
    if tabs is not None:
        properties.remove(tabs)
    return _canonical_xml(properties)


def _manual_sparse_resume() -> StructuredResume:
    path = Path("manual-test/render_deterministic_docx.py")
    spec = importlib.util.spec_from_file_location("manual_static_template_render", path)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module.build_manual_resume()[2]


def test_runtime_opens_packaged_docx_instead_of_blank_document(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import resume_tailor.infrastructure.static_template_docx as static_module

    actual_document = static_module.Document
    calls: list[str | None] = []

    def tracked_document(path: str | None = None):
        calls.append(path)
        return actual_document(path)

    monkeypatch.setattr(static_module, "Document", tracked_document)

    class ExactOnePageProvider:
        def measure(self, docx_path: Path) -> PageCountMeasurement:
            assert docx_path.is_file()
            return PageCountMeasurement(
                page_count=1,
                provider="runtime-template-open-test",
                confidence="exact",
                exact=True,
            )

    resume = _complete_resume().model_copy(
        update={
            "education": [],
            "technical_skills": [],
            "experiences": [],
            "projects": [],
            "experience_bullets": {},
            "project_bullets": {},
        }
    )

    ManagedResumeRenderer(page_count_provider=ExactOnePageProvider()).render_docx(
        resume,
        tmp_path / "opened-template.docx",
    )

    assert calls == [str(template_v1_docx_path())]
    assert all(call is not None for call in calls)


def test_packaged_template_is_content_neutral_and_has_explicit_prototypes() -> None:
    template_path = template_v1_docx_path()
    assert TEMPLATE_V1_REFERENCE_SHA256 == (
        "2b9dd1474b9e4a303a87b8a147f3511460988104efde7cfa053cad64294369cd"
    )
    assert TEMPLATE_V1_DOCX_SHA256 == (
        "2b4eeae9bed52ff27b86cb1e9f75516d0a9935359658849589b37ffef0a5974e"
    )
    document = Document(template_path)
    assert document.paragraphs
    for paragraph in document.paragraphs:
        residual = PLACEHOLDER_PATTERN.sub("", paragraph.text).strip()
        assert not residual or residual in SECTION_LABELS
        assert paragraph.text.strip()

    with ZipFile(template_path) as package:
        package_text = "\n".join(
            package.read(name).decode("utf-8", errors="ignore")
            for name in package.namelist()
            if name.endswith((".xml", ".rels"))
        )
        denied = {
            "Shiv Arora",
            "shiv2006arora",
            "647-997-0896",
            "linkedin.com/in/shiv-arora",
            "University of Toronto",
            "EXL Services",
            "Stush Foods",
            "Telebotics",
            "Crest",
            "MPC Hacks",
        }
        assert not {value for value in denied if value.casefold() in package_text.casefold()}
        assert "docProps/custom.xml" not in package.namelist()
        assert 'relationships/hyperlink"' not in package_text
        bookmark_names = set(re.findall(r'w:name="(TPL_[A-Z0-9_]+)"', package_text))
        assert bookmark_names == {
            "TPL_EDUCATION_ENTRY",
            "TPL_SKILL_CATEGORY_ROW",
            "TPL_EXPERIENCE_ENTRY_FIRST",
            "TPL_EXPERIENCE_ENTRY_REPEAT",
            "TPL_PROJECT_ENTRY",
        }


def test_section_geometry_and_source_package_styles_are_preserved(tmp_path: Path) -> None:
    output = tmp_path / "geometry.docx"
    render_template_v1_resume(_complete_resume(), output)
    reference = Document(REFERENCE)
    template = Document(template_v1_docx_path())
    generated = Document(output)

    expected_geometry = (
        reference.sections[0].page_width.twips,
        reference.sections[0].page_height.twips,
        reference.sections[0].top_margin.twips,
        reference.sections[0].right_margin.twips,
        reference.sections[0].bottom_margin.twips,
        reference.sections[0].left_margin.twips,
        reference.sections[0].header_distance.twips,
        reference.sections[0].footer_distance.twips,
    )
    for document in (template, generated):
        section = document.sections[0]
        assert (
            section.page_width.twips,
            section.page_height.twips,
            section.top_margin.twips,
            section.right_margin.twips,
            section.bottom_margin.twips,
            section.left_margin.twips,
            section.header_distance.twips,
            section.footer_distance.twips,
        ) == expected_geometry

    with (
        ZipFile(REFERENCE) as reference_package,
        ZipFile(template_v1_docx_path()) as template_package,
    ):
        for part in (
            "word/styles.xml",
            "word/numbering.xml",
            "word/theme/theme1.xml",
            "word/fontTable.xml",
        ):
            reference_root = etree.fromstring(reference_package.read(part))
            template_root = etree.fromstring(template_package.read(part))
            assert _canonical_xml(template_root) == _canonical_xml(reference_root)


def test_styles_direct_formatting_and_education_rows_are_source_derived() -> None:
    reference = Document(REFERENCE)
    template = Document(template_v1_docx_path())

    assert _canonical_xml(template.paragraphs[0].runs[0]._r.rPr) == _canonical_xml(
        reference.paragraphs[0].runs[0]._r.rPr
    )
    assert _canonical_ppr_without_tabs(template.paragraphs[3]) == _canonical_ppr_without_tabs(
        reference.paragraphs[4]
    )
    assert _canonical_ppr_without_tabs(template.paragraphs[4]) == _canonical_ppr_without_tabs(
        reference.paragraphs[5]
    )
    assert _canonical_xml(template.paragraphs[5]._p.pPr) == _canonical_xml(
        reference.paragraphs[6]._p.pPr
    )
    assert _canonical_xml(template.paragraphs[6]._p.pPr) == _canonical_xml(
        reference.paragraphs[7]._p.pPr
    )
    assert _canonical_xml(template.paragraphs[12].runs[0]._r.rPr) == _canonical_xml(
        reference.paragraphs[17].runs[0]._r.rPr
    )
    assert _canonical_xml(template.paragraphs[19].runs[0]._r.rPr) == _canonical_xml(
        reference.paragraphs[35].runs[0]._r.rPr
    )


def test_metadata_rows_keep_static_right_tabs_and_education_spacing(tmp_path: Path) -> None:
    output = tmp_path / "metadata-tabs.docx"
    render_template_v1_resume(_complete_resume(), output)
    document = Document(output)
    reference = Document(REFERENCE)
    metadata_prefixes = (
        "Example Polytechnic Institute",
        "Bachelor of Applied Engineering",
        "Firmware Engineer",
        "Example Robotics",
        "Embedded Systems Engineer",
        "Example Controls",
        "Sensor Platform",
        "Example Design Team",
        "Motor Controller",
    )
    metadata_rows = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(metadata_prefixes)
    ]
    assert metadata_rows
    assert all(
        len(paragraph.paragraph_format.tab_stops) == 1
        and paragraph.paragraph_format.tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
        and paragraph.paragraph_format.tab_stops[0].position.twips == 11_160
        for paragraph in metadata_rows
    )

    institution = _paragraph(document, "Example Polytechnic Institute")
    program = _paragraph(document, "Bachelor of Applied Engineering")
    awards = _paragraph(document, "Awards:")
    coursework = _paragraph(document, "Relevant Courses:")
    assert institution.paragraph_format.space_before.twips == (
        reference.paragraphs[4].paragraph_format.space_before.twips
    )
    assert program.paragraph_format.space_before.twips == (
        reference.paragraphs[5].paragraph_format.space_before.twips
    )
    assert awards.paragraph_format.space_before.twips == (
        reference.paragraphs[6].paragraph_format.space_before.twips
    )
    assert coursework.paragraph_format.space_before == (
        reference.paragraphs[7].paragraph_format.space_before
    )
    assert institution.runs[0].bold is True
    assert program.runs[0].italic is True


def test_skill_rows_preserve_category_and_value_formatting(tmp_path: Path) -> None:
    output = tmp_path / "skills.docx"
    render_template_v1_resume(_complete_resume(), output)
    document = Document(output)
    skill_rows = [
        _paragraph(document, "Programming & Scripting:"),
        _paragraph(document, "Embedded Systems:"),
    ]

    assert [row.text for row in skill_rows] == [
        "Programming & Scripting: C, Python",
        "Embedded Systems: STM32, SPI",
    ]
    assert all(row.runs[0].bold is True for row in skill_rows)
    assert all(row.runs[-1].bold is not True for row in skill_rows)
    assert all(row.paragraph_format.left_indent.twips == 108 for row in skill_rows)


def test_experience_project_and_bullet_prototypes_clone_with_canonical_geometry(
    tmp_path: Path,
) -> None:
    output = tmp_path / "cloned-blocks.docx"
    render_template_v1_resume(_complete_resume(), output)
    document = Document(output)
    reference = Document(REFERENCE)
    texts = [paragraph.text for paragraph in document.paragraphs]

    experience_titles = ("Firmware Engineer", "Embedded Systems Engineer")
    assert sum(text.startswith(experience_titles) for text in texts) == 2
    assert sum(text.startswith(("Sensor Platform", "Motor Controller")) for text in texts) == 2
    for expected in (
        "Integrated verified embedded controls",
        "Validated a verified sensor interface",
        "Documented a verified firmware release",
        "Reduced verified diagnostic latency",
        "Built a verified sensor prototype",
        "Tested the verified prototype",
        "Implemented a verified motor-control interface",
        "Measured the verified control response",
    ):
        assert any(text.startswith(expected) for text in texts)

    bullets = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    ]
    education_bullets = [
        paragraph
        for paragraph in bullets
        if paragraph.text.startswith(("Awards:", "Relevant Courses:"))
    ]
    body_bullets = [paragraph for paragraph in bullets if paragraph not in education_bullets]
    assert len(body_bullets) == 8
    assert all(
        paragraph.paragraph_format.left_indent.twips
        == reference.paragraphs[17].paragraph_format.left_indent.twips
        and paragraph.paragraph_format.first_line_indent.twips
        == reference.paragraphs[17].paragraph_format.first_line_indent.twips
        and paragraph._p.pPr.numPr.numId.val == 1
        for paragraph in body_bullets
    )


def test_unused_prototypes_leave_no_placeholders_duplicates_or_blank_paragraphs(
    tmp_path: Path,
) -> None:
    resume = _complete_resume().model_copy(
        update={
            "education": [],
            "technical_skills": [],
            "selected_skills": [],
            "projects": [],
            "project_bullets": {},
            "experiences": [_complete_resume().experiences[0]],
            "experience_bullets": {},
        }
    )
    output = tmp_path / "unused-prototypes.docx"
    render_template_v1_resume(resume, output)
    document = Document(output)
    texts = [paragraph.text for paragraph in document.paragraphs]

    assert all(text.strip() for text in texts)
    assert not PLACEHOLDER_PATTERN.search("\n".join(texts))
    assert texts.count("Technical Experience") == 1
    assert "Education" not in texts
    assert "Technical Skills" not in texts
    assert "Projects" not in texts
    with ZipFile(output) as package:
        package_text = "\n".join(
            package.read(name).decode("utf-8", errors="ignore")
            for name in package.namelist()
            if name.endswith((".xml", ".rels"))
        )
    assert "TPL_" not in package_text


def test_static_population_preserves_input_model_and_visible_selected_content(
    tmp_path: Path,
) -> None:
    resume = _complete_resume()
    before = resume.model_dump(mode="json")
    static_output = tmp_path / "static.docx"
    reconstructed_output = tmp_path / "reconstructed.docx"

    render_template_v1_resume(resume, static_output)
    render_structured_resume(
        resume,
        load_template_v1_layout_profile(),
        reconstructed_output,
    )

    assert resume.model_dump(mode="json") == before
    assert [paragraph.text for paragraph in Document(static_output).paragraphs] == [
        paragraph.text for paragraph in Document(reconstructed_output).paragraphs
    ]


def test_static_population_adds_no_text_outside_the_existing_content_contract(
    tmp_path: Path,
) -> None:
    resume = _complete_resume()
    output = tmp_path / "truthful.docx"
    comparison = tmp_path / "content-contract.docx"
    render_template_v1_resume(resume, output)
    render_structured_resume(resume, load_template_v1_layout_profile(), comparison)

    generated_text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    contract_text = "\n".join(paragraph.text for paragraph in Document(comparison).paragraphs)
    assert generated_text == contract_text
    assert "unsupported" not in generated_text.casefold()
    assert not PLACEHOLDER_PATTERN.search(generated_text)


def test_sparse_manual_resume_is_one_word_page_when_word_is_available(
    tmp_path: Path,
) -> None:
    output = tmp_path / "sparse-manual-static.docx"
    render_template_v1_resume(_manual_sparse_resume(), output)
    try:
        measurement = MicrosoftWordDocxPageCountProvider().measure(output)
    except PageCountVerificationError as error:
        pytest.skip(f"Microsoft Word verification unavailable: {error}")

    assert measurement.exact is True
    assert measurement.page_count == 1


def test_managed_sparse_manual_resume_preserves_severe_underfill_status(
    tmp_path: Path,
) -> None:
    class ExactOnePageProvider:
        def measure(self, docx_path: Path) -> PageCountMeasurement:
            assert docx_path.is_file()
            return PageCountMeasurement(
                page_count=1,
                provider="controlled-one-page-verification",
                confidence="exact",
                exact=True,
            )

    renderer = ManagedResumeRenderer(page_count_provider=ExactOnePageProvider())
    renderer.render_docx(
        _manual_sparse_resume(),
        tmp_path / "sparse-manual-underfill.docx",
    )

    assert renderer.last_page_utilization is not None
    assert renderer.last_page_utilization.status.value == "severe_underfill"
    assert renderer.last_page_utilization.uncontrolled_blank_paragraph_count == 0
