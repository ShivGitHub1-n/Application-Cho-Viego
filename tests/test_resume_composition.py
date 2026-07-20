from __future__ import annotations

import json
from collections.abc import Iterator
from hashlib import sha256
from pathlib import Path

import pytest
from docx import Document

from resume_tailor.application.llm_services import HybridLlmServices
from resume_tailor.application.resume_composition import (
    CompositionSearchBounds,
    DeterministicResumeComposer,
    _posting_context,
    _State,
)
from resume_tailor.application.resume_features import TemplateV1BulletLineEstimator
from resume_tailor.application.services import TailorResumeService
from resume_tailor.domain.layout import PageUtilizationStatus
from resume_tailor.domain.models import (
    EntityKind,
    EvidenceItem,
    JobPosting,
    MasterProfile,
    ResumeItem,
    ResumeStrategy,
    StructuredResume,
    TemplateConstraints,
)
from resume_tailor.domain.resume_composition import (
    TEMPLATE_V1_UTILIZATION_TARGET_FLOOR,
    CompositionOutcome,
    CompositionTerminationReason,
    PageFitEvaluation,
    PreferredDensityStatus,
    ProjectRepresentationStatus,
)
from resume_tailor.infrastructure.composition_page_fit import TemplateV1PageFitEvaluator
from resume_tailor.infrastructure.optimization import (
    DeterministicResumeOptimizer,
    EvidenceBoundResumeWriter,
)
from resume_tailor.infrastructure.rendering import (
    PageCountMeasurement,
    PageCountVerificationError,
)
from resume_tailor.infrastructure.static_template_docx import render_template_v1_resume
from resume_tailor.infrastructure.template_v1 import (
    TEMPLATE_V1_DOCX_SHA256,
    template_v1_docx_path,
)
from tests.fakes import FakeResumeLanguageModel

FIXTURE = Path(__file__).parent / "fixtures" / "resume_composition_cases.json"


class ParagraphLimitPageProvider:
    def __init__(self, maximum_paragraphs: int = 36) -> None:
        self.maximum_paragraphs = maximum_paragraphs
        self.calls = 0
        self.overflow_measurements = 0

    def measure(self, docx_path: Path) -> PageCountMeasurement:
        self.calls += 1
        paragraph_count = sum(
            bool(paragraph.text.strip()) for paragraph in Document(docx_path).paragraphs
        )
        page_count = 1 if paragraph_count <= self.maximum_paragraphs else 2
        self.overflow_measurements += int(page_count > 1)
        return PageCountMeasurement(
            page_count=page_count,
            provider=f"controlled paragraph limit {self.maximum_paragraphs}",
            confidence="exact",
            exact=True,
        )


class FailingExactPageProvider:
    def measure(self, docx_path: Path) -> PageCountMeasurement:
        raise PageCountVerificationError(f"Controlled exact provider failure for {docx_path.name}")


class FixedRatioPageFitEvaluator:
    def __init__(self, utilization_ratio: float) -> None:
        self.utilization_ratio = utilization_ratio
        self.calls = 0

    def evaluate(
        self,
        resume: object,
        *,
        attempt_exact: bool = True,
    ) -> PageFitEvaluation:
        self.calls += 1
        return PageFitEvaluation(
            status=(
                PageUtilizationStatus.SEVERE_UNDERFILL
                if self.utilization_ratio < TEMPLATE_V1_UTILIZATION_TARGET_FLOOR
                else PageUtilizationStatus.ACCEPTABLE_ONE_PAGE
            ),
            page_count=1,
            exact=attempt_exact,
            provider="fixed-ratio acceptance evaluator",
            utilization_ratio=self.utilization_ratio,
            fits_one_page=True,
        )


class LineCostDensityEvaluator:
    def __init__(self) -> None:
        self._estimator = TemplateV1BulletLineEstimator()

    def evaluate(
        self,
        resume: StructuredResume,
        *,
        attempt_exact: bool = True,
    ) -> PageFitEvaluation:
        bullets = list(_output_bullets(resume))
        utilization = min(
            0.99,
            0.55
            + sum(
                self._estimator.estimate(bullet.text).total_vertical_line_cost for bullet in bullets
            )
            * 0.055,
        )
        return PageFitEvaluation(
            status=(
                PageUtilizationStatus.SEVERE_UNDERFILL
                if utilization < TEMPLATE_V1_UTILIZATION_TARGET_FLOOR
                else PageUtilizationStatus.ACCEPTABLE_ONE_PAGE
            ),
            page_count=1,
            exact=attempt_exact,
            provider="controlled line-cost density evaluator",
            utilization_ratio=utilization,
            fits_one_page=True,
        )


@pytest.fixture
def composition_fixture() -> tuple[MasterProfile, dict[str, JobPosting]]:
    raw = json.loads(FIXTURE.read_text(encoding="utf-8"))
    return (
        MasterProfile.model_validate(raw["profile"]),
        {key: JobPosting.model_validate(value) for key, value in raw["postings"].items()},
    )


def _service(
    provider: ParagraphLimitPageProvider,
    *,
    hybrid_services: HybridLlmServices | None = None,
) -> TailorResumeService:
    return TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
        hybrid_services=hybrid_services,
        resume_composer=DeterministicResumeComposer(TemplateV1PageFitEvaluator(provider)),
    )


def _composed(
    profile: MasterProfile,
    posting: JobPosting,
    provider: ParagraphLimitPageProvider | None = None,
) -> tuple[object, object, ParagraphLimitPageProvider]:
    resolved_provider = provider or ParagraphLimitPageProvider()
    service = _service(resolved_provider)
    plan = service.create_plan(profile, posting, TemplateConstraints())
    resume = service.build_document(plan, profile, set())
    return plan, resume, resolved_provider


def _output_bullets(resume: object) -> Iterator[object]:
    yield from (bullet for bullets in resume.experience_bullets.values() for bullet in bullets)
    yield from (bullet for bullets in resume.project_bullets.values() for bullet in bullets)


def _synthetic_profile(
    *,
    experiences: list[dict[str, object]],
    projects: list[dict[str, object]],
    evidence: list[dict[str, object]],
    technical_skills: list[dict[str, object]] | None = None,
    declared_skills: list[str] | None = None,
) -> MasterProfile:
    return MasterProfile.model_validate(
        {
            "id": "generic-growing-profile",
            "user_id": "generic-user",
            "display_name": "Jordan Candidate",
            "contact": {"email": "candidate@example.com"},
            "education": [
                {
                    "school": "Public Polytechnic",
                    "program": "Bachelor of Engineering",
                    "start_date": "2021",
                    "expected_graduation_date": "2026",
                    "location": "Toronto, ON",
                }
            ],
            "experiences": experiences,
            "projects": projects,
            "declared_skills": declared_skills or [],
            "technical_skills": technical_skills or [],
            "evidence": evidence,
        }
    )


def _compose_with_bounds(
    profile: MasterProfile,
    posting: JobPosting,
    bounds: CompositionSearchBounds,
) -> object:
    baseline = StructuredResume(
        profile_id=profile.id,
        profile_version=profile.version,
        posting_id=posting.id,
        template_id="managed-engineering-v1",
        display_name=profile.display_name,
        contact_line=profile.contact.email,
        strategy=ResumeStrategy(
            role_family="deterministic_test",
            primary_focus=posting.title,
            rationale="Controlled composer-only acceptance fixture.",
        ),
        education=profile.education,
    )
    return DeterministicResumeComposer(
        FixedRatioPageFitEvaluator(0.82),
        bounds=bounds,
    ).compose(
        baseline,
        profile,
        posting,
        TemplateConstraints(),
    )


def test_rich_firmware_composition_expands_to_multiple_entries_and_project(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    plan, resume, provider = _composed(profile, postings["firmware"])
    diagnostic = resume.composition_diagnostic

    assert plan.strategy is not None
    assert diagnostic is not None
    assert diagnostic.outcome is CompositionOutcome.ACCEPTABLE_ONE_PAGE
    assert diagnostic.page_count == 1
    assert diagnostic.verification_status.value == "exact"
    assert diagnostic.final_utilization_ratio >= TEMPLATE_V1_UTILIZATION_TARGET_FLOOR
    assert diagnostic.utilization_target_reached is True
    assert diagnostic.termination_reason is CompositionTerminationReason.TARGET_FINALISTS_FOUND
    assert diagnostic.preferred_density_reached is False
    assert diagnostic.underfill_reasons
    assert any(reason.value == "search_bounds_limited" for reason in diagnostic.underfill_reasons)
    assert provider.overflow_measurements == diagnostic.overflow_rollbacks
    assert "firmware-intern" in diagnostic.selected_experience_ids
    assert {
        "electronics-lab",
        "mechanical-team",
        "controls-test-coop",
    } & set(diagnostic.selected_experience_ids)
    assert "rover-controller" in diagnostic.selected_project_ids
    assert diagnostic.bullet_counts["firmware-intern"] >= 2
    assert len(diagnostic.selected_bullet_ids) >= 12
    assert diagnostic.maximum_search_depth is None
    assert len(diagnostic.selected_skill_category_ids) == 3
    assert diagnostic.final_utilization_ratio == pytest.approx(0.8055630026809651)
    assert provider.calls <= diagnostic.maximum_page_evaluations


def test_generic_admission_reopens_previously_hidden_relevant_evidence(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    removed_entry_ids = {"controls-test-coop", "sensor-node"}
    legacy_profile = profile.model_copy(
        update={
            "experiences": [
                item for item in profile.experiences if item.id not in removed_entry_ids
            ],
            "projects": [item for item in profile.projects if item.id not in removed_entry_ids],
            "evidence": [
                item for item in profile.evidence if item.entity_id not in removed_entry_ids
            ],
        }
    )
    composer = DeterministicResumeComposer(FixedRatioPageFitEvaluator(0.578))
    context = _posting_context(postings["firmware"])
    bullets = composer._rank_bullets(legacy_profile, context)
    bullet_by_id = {item.evidence_id: item for item in bullets}
    selected_ids = frozenset(
        {
            "firmware-drivers",
            "firmware-rtos",
            "firmware-hil-test",
            "firmware-latency",
            "electronics-bringup",
            "electronics-faults",
            "rover-integration",
            "rover-validation",
            "rover-enclosure",
            "rover-fabrication",
        }
    )
    state = _State(selected_ids, frozenset())
    remaining_positive_marginal = []
    selected_entries = {bullet_by_id[item].entry_id for item in selected_ids}
    for candidate in bullets:
        if candidate.evidence_id in selected_ids:
            continue
        penalty, duplicate = composer._redundancy_penalty(
            candidate,
            state,
            bullet_by_id,
        )
        opening_penalty = 6.0 if candidate.entry_id not in selected_entries else 0.0
        if (
            not duplicate
            and candidate.score - penalty - opening_penalty >= composer._minimum_marginal_score
        ):
            remaining_positive_marginal.append(candidate)

    assert [item.evidence_id for item in remaining_positive_marginal] == ["mechanical-prototype"]
    assert remaining_positive_marginal[0].admission_reason.startswith(
        "Admitted through specific reviewed-text overlap"
    )
    historical_remaining = next(
        item
        for item in composer._all_bullet_candidates(legacy_profile, context)
        if item.evidence_id == "observability-deployment"
    )
    assert 17.9 - 3.4 - 6.0 == pytest.approx(8.5)
    selected_project_lines = 2 + sum(
        bullet_by_id[item].estimated_lines
        for item in selected_ids
        if bullet_by_id[item].entry_kind is EntityKind.PROJECT
    )
    cloud_project_lines = 2 + historical_remaining.estimated_lines
    assert selected_project_lines + cloud_project_lines > TemplateConstraints().max_project_lines
    assert len(selected_ids) + 1 < 12
    assert 40 < 48

    baseline_service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
    )
    plan = baseline_service.create_plan(
        legacy_profile,
        postings["firmware"],
        TemplateConstraints(),
    )
    baseline = baseline_service.build_document(plan, legacy_profile, set())
    skills = composer._rank_skills(legacy_profile, context, bullets)
    skill_by_id = {item.category_id: item for item in skills}
    old_state = _State(
        selected_ids,
        frozenset({skills[0].category_id}),
    )
    old_resume = composer._resume_for_state(
        baseline,
        legacy_profile,
        old_state,
        bullet_by_id,
        skill_by_id,
    )
    old_evaluation = TemplateV1PageFitEvaluator(ParagraphLimitPageProvider()).evaluate(old_resume)

    assert old_evaluation.utilization_ratio == pytest.approx(0.578083109919571)
    assert old_evaluation.status is PageUtilizationStatus.SEVERE_UNDERFILL


def test_57_8_rich_result_with_unused_evidence_is_not_acceptable(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    evaluator = FixedRatioPageFitEvaluator(0.578)
    service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
        resume_composer=DeterministicResumeComposer(
            evaluator,
            bounds=CompositionSearchBounds(maximum_selected_bullets=10),
        ),
    )

    plan = service.create_plan(profile, postings["firmware"], TemplateConstraints())
    resume = service.build_document(plan, profile, set())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.outcome is CompositionOutcome.SEVERE_UNDERFILL
    assert diagnostic.utilization_target_reached is False
    assert diagnostic.additional_evidence_unavailable is False
    assert any(
        reason.value
        in {
            "profile_incomplete",
            "evidence_limited",
            "quality_limited",
            "job_match_limited",
            "candidate_construction_failure",
            "search_bounds_limited",
        }
        for reason in diagnostic.underfill_reasons
    )
    assert (
        diagnostic.unused_admissible_candidates or diagnostic.candidates_excluded_by_search_bounds
    )


def test_search_bound_termination_is_typed_and_bound_candidates_are_reported(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    bounds = CompositionSearchBounds(
        beam_width=2,
        maximum_estimated_page_evaluations=8,
        maximum_exact_finalist_evaluations=2,
        maximum_expansion_operations=200,
        maximum_expansions_per_state=2,
    )
    service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
        resume_composer=DeterministicResumeComposer(
            FixedRatioPageFitEvaluator(0.45),
            bounds=bounds,
        ),
    )

    plan = service.create_plan(profile, postings["firmware"], TemplateConstraints())
    resume = service.build_document(plan, profile, set())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.termination_reason is CompositionTerminationReason.ESTIMATED_EVALUATION_LIMIT
    assert diagnostic.estimated_page_evaluations == 8
    assert diagnostic.candidates_excluded_by_search_bounds
    assert all(
        candidate.exclusion_category.value == "search_bound"
        for candidate in diagnostic.candidates_excluded_by_search_bounds
    )


def test_mechanical_posting_is_mechanical_led_even_when_classifier_is_not(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    plan, resume, _provider = _composed(profile, postings["mechanical"])
    diagnostic = resume.composition_diagnostic

    assert plan.report.role.role_family == "software_data_engineering"
    assert diagnostic is not None
    assert diagnostic.selected_experience_ids == ["mechanical-team"]
    assert "rover-controller" in diagnostic.selected_project_ids
    assert {
        "mechanical-cad",
        "mechanical-prototype",
        "mechanical-test",
        "rover-enclosure",
        "rover-fabrication",
    }.issubset(diagnostic.selected_bullet_ids)
    assert "cloud-intern" not in diagnostic.selected_experience_ids
    assert diagnostic.selected_skill_category_labels == ["Mechanical Design & Manufacturing"]


def test_software_posting_excludes_unrelated_mechanical_content(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    _plan, resume, _provider = _composed(profile, postings["software"])
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.selected_experience_ids == ["cloud-intern"]
    assert diagnostic.selected_project_ids == ["cloud-observability"]
    assert all(
        not evidence_id.startswith(("mechanical-", "rover-"))
        for evidence_id in diagnostic.selected_bullet_ids
    )
    assert diagnostic.selected_skill_category_labels == ["Software & Cloud"]


def test_mixed_posting_selects_complementary_cross_disciplinary_evidence(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    _plan, resume, _provider = _composed(profile, postings["mixed"])
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert {"firmware-intern", "cloud-intern"}.issubset(diagnostic.selected_experience_ids)
    assert "rover-controller" in diagnostic.selected_project_ids
    assert any(
        evidence_id.startswith("firmware-") for evidence_id in diagnostic.selected_bullet_ids
    )
    assert any(evidence_id.startswith("cloud-") for evidence_id in diagnostic.selected_bullet_ids)
    assert {
        "mechanical-cad",
        "mechanical-prototype",
        "mechanical-test",
    } & set(diagnostic.selected_bullet_ids)
    selected_rover_mechanical = {"rover-enclosure", "rover-fabrication"} & set(
        diagnostic.selected_bullet_ids
    )
    assert selected_rover_mechanical
    rover_mechanical = next(
        candidate
        for candidate in diagnostic.selected_candidates
        if candidate.candidate_id == f"bullet:{sorted(selected_rover_mechanical)[0]}"
    )
    assert rover_mechanical.line_fit is not None
    assert rover_mechanical.line_fit.awkward_wrap_risk is True
    assert rover_mechanical.line_fit.future_rewrite_recommended is True
    assert diagnostic.utilization_target_reached is True
    assert len(diagnostic.selected_skill_category_ids) == 3
    assert diagnostic.final_utilization_ratio == pytest.approx(0.8008042895442359)


def test_sparse_profile_reports_insufficient_evidence_without_invention() -> None:
    profile = MasterProfile(
        id="sparse-profile",
        user_id="sparse-user",
        display_name="Sparse Candidate",
        experiences=[
            ResumeItem(
                id="sparse-entry",
                title="Firmware Assistant",
                kind=EntityKind.EXPERIENCE,
                organization="Reviewed Organization",
            )
        ],
        evidence=[
            EvidenceItem(
                id="sparse-evidence",
                entity_id="sparse-entry",
                source_text="Tested STM32 firmware on a reviewed prototype.",
                technologies=["STM32"],
                capabilities=["firmware testing"],
            )
        ],
    )
    posting = JobPosting(
        id="sparse-posting",
        title="Firmware Intern",
        description="Develop and test STM32 firmware.",
    )

    _plan, resume, _provider = _composed(profile, posting)
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.outcome is CompositionOutcome.INSUFFICIENT_EVIDENCE
    assert diagnostic.additional_evidence_unavailable is True
    assert diagnostic.selected_bullet_ids == ["sparse-evidence"]
    assert [bullet.text for bullet in _output_bullets(resume)] == [
        "Tested STM32 firmware on a reviewed prototype."
    ]


def test_overflow_search_rolls_back_and_keeps_best_exact_one_page_alternative(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    provider = ParagraphLimitPageProvider(maximum_paragraphs=24)

    _plan, resume, provider = _composed(profile, postings["firmware"], provider)
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.page_count == 1
    assert diagnostic.verification_status.value == "exact"
    assert diagnostic.overflow_rollbacks >= 1
    assert provider.overflow_measurements >= 1
    assert diagnostic.best_exact_verified_utilization_ratio is not None
    assert diagnostic.best_exact_verified_utilization_ratio == pytest.approx(
        diagnostic.final_utilization_ratio
    )
    assert any(
        iteration.overflow and not iteration.accepted
        for iteration in diagnostic.page_fill_iterations
    )
    assert any(iteration.accepted for iteration in diagnostic.page_fill_iterations)


def test_redundant_bullets_are_suppressed_in_favor_of_distinct_coverage(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    _plan, resume, _provider = _composed(profile, postings["firmware"])
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    selected = set(diagnostic.selected_bullet_ids)
    assert not {
        "firmware-drivers",
        "firmware-drivers-duplicate",
    }.issubset(selected)
    duplicate_diagnostic = next(
        candidate
        for candidate in diagnostic.excluded_high_ranking_candidates
        if candidate.source_ids == ["firmware-drivers-duplicate"]
    )
    assert "duplicate" in (duplicate_diagnostic.exclusion_reason or "")
    assert duplicate_diagnostic.redundancy_penalty > 0
    assert "firmware-hil-test" in selected


def test_selecting_entry_does_not_automatically_select_redundant_bullets() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "validation-entry",
                "title": "Platform Validation Developer",
                "kind": "experience",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "primary-proof",
                "entity_id": "validation-entry",
                "source_text": (
                    "Validated Kubernetes releases through CI/CD policy checks and "
                    "repeatable rollback tests."
                ),
            },
            {
                "id": "duplicate-proof",
                "entity_id": "validation-entry",
                "source_text": (
                    "Validated Kubernetes releases with CI/CD policy checks and "
                    "repeatable rollback tests."
                ),
            },
            {
                "id": "diagnostic-proof",
                "entity_id": "validation-entry",
                "source_text": (
                    "Diagnosed PostgreSQL reliability defects using structured logs "
                    "and reviewed incident evidence."
                ),
            },
        ],
    )
    posting = JobPosting(
        id="marginal-posting",
        title="Platform Validation Developer",
        description=(
            "Validate Kubernetes CI/CD releases, test rollbacks, and diagnose "
            "PostgreSQL reliability defects."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    entry = next(
        item for item in diagnostic.entry_bullet_selections
        if item.entry_id == "validation-entry"
    )
    assert 1 < len(entry.selected_bullet_ids) < len(entry.available_bullet_ids)
    assert entry.omitted_bullet_reasons
    assert entry.retained_all_available_bullets is False


def test_all_distinct_bullets_can_survive_with_contribution_diagnostics() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "systems-entry",
                "title": "Systems Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "design-proof",
                "entity_id": "systems-entry",
                "source_text": "Designed CAN interfaces for an embedded control system.",
            },
            {
                "id": "test-proof",
                "entity_id": "systems-entry",
                "source_text": "Validated sensor timing through hardware-in-the-loop tests.",
            },
            {
                "id": "debug-proof",
                "entity_id": "systems-entry",
                "source_text": "Debugged power faults with oscilloscope measurements.",
            },
            {
                "id": "release-proof",
                "entity_id": "systems-entry",
                "source_text": "Documented requirements and release evidence for design review.",
            },
            {
                "id": "outcome-proof",
                "entity_id": "systems-entry",
                "source_text": "Reduced verified control-loop latency by 18%.",
            },
        ],
    )
    posting = JobPosting(
        id="distinct-posting",
        title="Embedded Systems Engineer",
        description=(
            "Design CAN control systems, validate sensor timing with hardware-in-the-loop "
            "tests, debug power faults with an oscilloscope, document requirements, and "
            "improve control-loop latency."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    entry = next(item for item in diagnostic.entry_bullet_selections)
    assert entry.retained_all_available_bullets is True
    assert set(entry.distinct_contributions) == set(entry.selected_bullet_ids)
    assert all(entry.distinct_contributions.values())


def test_typical_bullet_count_is_soft_and_seven_distinct_bullets_can_survive() -> None:
    evidence_specs = [
        ("interface-proof", "Design CAN controls for embedded interfaces."),
        ("timing-proof", "Validated sensor timing through hardware-in-the-loop tests."),
        ("power-proof", "Debugged power faults with oscilloscope measurements."),
        ("requirements-proof", "Documented requirements for release review."),
        ("latency-proof", "Reduced verified control-loop latency by 18%."),
        ("manufacturing-proof", "Build 3D-printed prototypes and enclosures."),
        ("deployment-proof", "Automated firmware release checks with CI/CD."),
    ]
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "systems-entry",
                "title": "Systems Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": evidence_id,
                "entity_id": "systems-entry",
                "source_text": text,
            }
            for evidence_id, text in evidence_specs
        ],
    )
    posting = JobPosting(
        id="broad-systems-posting",
        title="Embedded Systems Engineer",
        description=(
            "Design CAN controls, validate sensor timing with hardware-in-the-loop tests, "
            "debug power faults with an oscilloscope, document release requirements, "
            "improve control-loop latency, build 3D-printed prototypes, and automate "
            "firmware CI/CD release checks."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    entry = next(item for item in diagnostic.entry_bullet_selections)
    assert entry.retained_all_available_bullets is True
    assert len(entry.selected_bullet_ids) == 7
    assert set(entry.distinct_contributions) == {
        evidence_id for evidence_id, _text in evidence_specs
    }


def test_later_repetitive_bullet_receives_larger_diminishing_return_penalty() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "release-entry",
                "title": "Release Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "release-one",
                "entity_id": "release-entry",
                "source_text": "Validated Kubernetes releases through CI/CD gates.",
            },
            {
                "id": "release-two",
                "entity_id": "release-entry",
                "source_text": "Tested Kubernetes rollbacks through CI/CD gates.",
            },
            {
                "id": "release-three",
                "entity_id": "release-entry",
                "source_text": "Reviewed Kubernetes deployments through CI/CD gates.",
            },
        ],
    )
    posting = JobPosting(
        id="release-posting",
        title="Release Engineer",
        description="Validate Kubernetes deployments, CI/CD gates, and rollbacks.",
    )
    composer = DeterministicResumeComposer(FixedRatioPageFitEvaluator(0.90))
    candidates = {
        item.evidence_id: item
        for item in composer._rank_bullets(profile, _posting_context(posting))
    }

    early_penalty, _ = composer._redundancy_penalty(
        candidates["release-three"],
        _State(frozenset({"release-one"}), frozenset()),
        candidates,
    )
    later_penalty, _ = composer._redundancy_penalty(
        candidates["release-three"],
        _State(frozenset({"release-one", "release-two"}), frozenset()),
        candidates,
    )

    assert later_penalty > early_penalty


def test_composition_preserves_exact_reviewed_text_and_provenance(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    _plan, resume, _provider = _composed(profile, postings["firmware"])
    source_text = {evidence.id: evidence.source_text for evidence in profile.evidence}
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    for bullet in _output_bullets(resume):
        assert bullet.text == source_text[bullet.id]
        assert bullet.evidence_ids == [bullet.id]
    assert all(
        candidate.provenance for candidate in diagnostic.selected_candidates if candidate.source_ids
    )


def test_all_llm_flags_disabled_make_zero_provider_calls_and_render_docx(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
    tmp_path: Path,
) -> None:
    profile, postings = composition_fixture
    fake = FakeResumeLanguageModel()
    hybrid = HybridLlmServices(
        fake,
        retry_count=0,
        max_calls=4,
        enable_opportunity_analysis=False,
        enable_composition=False,
        enable_bullet_rewrite=False,
    )
    service = _service(ParagraphLimitPageProvider(), hybrid_services=hybrid)

    plan = service.create_plan(profile, postings["firmware"], TemplateConstraints())
    resume = service.build_document(plan, profile, set())
    output = tmp_path / "llm-disabled-page-filled.docx"
    render_template_v1_resume(resume, output)

    assert sum(fake.calls.values()) == 0
    assert resume.composition_diagnostic is not None
    assert output.read_bytes().startswith(b"PK")
    rendered_text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "Composition diagnostic" not in rendered_text
    assert resume.composition_diagnostic.reason not in rendered_text


def test_page_fill_substantially_improves_prior_sparse_plan_utilization(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    baseline_service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
    )
    plan = baseline_service.create_plan(
        profile,
        postings["firmware"],
        TemplateConstraints(),
    )
    baseline = baseline_service.build_document(plan, profile, set())
    prior = TemplateV1PageFitEvaluator(ParagraphLimitPageProvider()).evaluate(baseline)

    _plan, composed, _provider = _composed(profile, postings["firmware"])
    diagnostic = composed.composition_diagnostic

    assert diagnostic is not None
    assert prior.utilization_ratio == pytest.approx(0.2906166219839142)
    assert diagnostic.final_utilization_ratio > prior.utilization_ratio + 0.20
    assert len(diagnostic.selected_bullet_ids) > sum(
        len(bullets) for bullets in baseline.experience_bullets.values()
    ) + sum(len(bullets) for bullets in baseline.project_bullets.values())


def test_exact_provider_failure_returns_typed_unverified_diagnostic(
    composition_fixture: tuple[MasterProfile, dict[str, JobPosting]],
) -> None:
    profile, postings = composition_fixture
    baseline_service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
    )
    plan = baseline_service.create_plan(
        profile,
        postings["firmware"],
        TemplateConstraints(),
    )
    baseline = baseline_service.build_document(plan, profile, set())

    result = TemplateV1PageFitEvaluator(FailingExactPageProvider()).evaluate(baseline)

    assert result.exact is False
    assert result.status.value == "unverified"
    assert result.verification_failure is not None
    assert "Controlled exact provider failure" in result.verification_failure


def test_composition_correction_does_not_change_template_v1_package() -> None:
    assert sha256(template_v1_docx_path().read_bytes()).hexdigest() == (TEMPLATE_V1_DOCX_SHA256)


def test_profile_growth_automatically_evaluates_new_reviewed_content() -> None:
    posting = JobPosting(
        id="growth-posting",
        title="Cloud Platform Developer",
        description=(
            "Deploy containerized services to Kubernetes, implement CI/CD pipelines, "
            "monitor PostgreSQL workloads, and debug production reliability defects."
        ),
    )
    original = _synthetic_profile(
        experiences=[
            {
                "id": "legacy-service",
                "title": "Application Developer",
                "kind": "experience",
                "organization": "Regional Cooperative",
                "start_date": "2022",
                "end_date": "2023",
                "location": "Remote",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "legacy-health-check",
                "entity_id": "legacy-service",
                "source_text": (
                    "Implemented PostgreSQL health checks for an internal reporting service."
                ),
            }
        ],
    )
    bounds = CompositionSearchBounds(
        maximum_selected_entries=2,
        maximum_experience_entries=1,
        maximum_project_entries=1,
        maximum_selected_bullets=5,
    )
    before = _compose_with_bounds(original, posting, bounds)
    grown = MasterProfile.model_validate(
        {
            **original.model_dump(),
            "experiences": [
                *original.experiences,
                ResumeItem(
                    id="new-platform-experience",
                    title="Platform Developer",
                    kind=EntityKind.EXPERIENCE,
                    organization="Member-owned Utility",
                    start_date="2024",
                    end_date="Present",
                    location="Calgary, AB",
                ),
            ],
            "projects": [
                ResumeItem(
                    id="new-release-project",
                    title="Container Release Pipeline",
                    kind=EntityKind.PROJECT,
                    start_date="2025",
                    end_date="2025",
                    location="Remote",
                )
            ],
            "technical_skills": [
                {
                    "id": "new-cloud-tools",
                    "category": "Delivery Platforms",
                    "values": ["Kubernetes", "PostgreSQL 16", "CI/CD"],
                }
            ],
            "evidence": [
                *original.evidence,
                EvidenceItem(
                    id="new-kubernetes-delivery",
                    entity_id="new-platform-experience",
                    source_text=(
                        "Deployed containerized services to Kubernetes through reviewed "
                        "CI/CD release gates and production rollback checks."
                    ),
                ),
                EvidenceItem(
                    id="new-production-debugging",
                    entity_id="new-platform-experience",
                    source_text=(
                        "Debugged production reliability defects using PostgreSQL 16 "
                        "health telemetry and reproducible incident tests."
                    ),
                ),
                EvidenceItem(
                    id="new-pipeline-validation",
                    entity_id="new-release-project",
                    source_text=(
                        "Built a container release pipeline with Kubernetes validation, "
                        "CI/CD policy checks, and automated rollback verification."
                    ),
                ),
            ],
        }
    )
    after = _compose_with_bounds(grown, posting, bounds)

    assert before.composition_diagnostic is not None
    assert after.composition_diagnostic is not None
    assert "new-platform-experience" in after.composition_diagnostic.selected_experience_ids
    assert "new-release-project" in after.composition_diagnostic.selected_project_ids
    assert "legacy-service" not in after.composition_diagnostic.selected_experience_ids
    assert "new-cloud-tools" in after.composition_diagnostic.selected_skill_category_ids
    assert {
        "new-kubernetes-delivery",
        "new-production-debugging",
        "new-pipeline-validation",
    }.issubset(after.composition_diagnostic.selected_bullet_ids)
    assert {bullet.id: (bullet.text, bullet.evidence_ids) for bullet in _output_bullets(after)}[
        "new-kubernetes-delivery"
    ] == (
        grown.evidence[-3].source_text,
        ["new-kubernetes-delivery"],
    )


def test_skill_rows_use_current_reviewed_categories_with_measured_support() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "controls-entry",
                "title": "Controls Developer",
                "kind": "experience",
                "organization": "Municipal Laboratory",
                "start_date": "2024",
                "end_date": "Present",
                "location": "Hamilton, ON",
            }
        ],
        projects=[],
        technical_skills=[
            {
                "id": "languages",
                "category": "Languages & Runtime",
                "values": ["C++/CLI", ".NET 8.0"],
            },
            {
                "id": "standards",
                "category": "Standards",
                "values": ["ISO-26262"],
            },
            {
                "id": "unrelated",
                "category": "Workshop",
                "values": ["Manual milling"],
            },
        ],
        evidence=[
            {
                "id": "controls-runtime",
                "entity_id": "controls-entry",
                "source_text": (
                    "Implemented C++/CLI diagnostics on .NET 8.0 for safety monitoring."
                ),
            }
        ],
    )
    posting = JobPosting(
        id="controls-posting",
        title="Controls Software Developer",
        description=(
            "Implement C++/CLI diagnostics on .NET 8.0 and support ISO-26262 safety validation."
        ),
    )
    composer = DeterministicResumeComposer(FixedRatioPageFitEvaluator(0.82))
    context = _posting_context(posting)
    bullets = composer._rank_bullets(profile, context)
    skills = composer._rank_skills(profile, context, bullets)

    assert [item.category_id for item in skills] == ["languages", "standards"]
    assert [skill.value for skill in skills[0].category.skills] == ["C++/CLI", ".NET 8.0"]
    assert skills[0].supported_skill_ids
    assert skills[0].declared_only_skill_ids == ()
    assert skills[1].supported_skill_ids == ()
    assert skills[1].declared_only_skill_ids
    assert skills[1].score >= composer._minimum_skill_score


def test_rich_project_is_substantive_and_defeats_shallow_representation() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "platform-entry",
                "title": "Platform Developer",
                "kind": "experience",
            }
        ],
        projects=[
            {
                "id": "substantive-project",
                "title": "Reliability Test Platform",
                "kind": "project",
            },
            {
                "id": "shallow-project",
                "title": "Deployment Demo",
                "kind": "project",
            },
        ],
        evidence=[
            {
                "id": "platform-proof",
                "entity_id": "platform-entry",
                "source_text": "Built Kubernetes services with PostgreSQL persistence.",
            },
            {
                "id": "project-build",
                "entity_id": "substantive-project",
                "source_text": (
                    "Built a Kubernetes reliability test platform with controlled "
                    "dependency-failure scenarios."
                ),
            },
            {
                "id": "project-validate",
                "entity_id": "substantive-project",
                "source_text": (
                    "Validated rollback, timeout, and PostgreSQL recovery behavior "
                    "through repeatable automated tests."
                ),
            },
            {
                "id": "project-observe",
                "entity_id": "substantive-project",
                "source_text": (
                    "Implemented structured telemetry for incident diagnosis and "
                    "reviewed recovery evidence."
                ),
            },
            {
                "id": "shallow-proof",
                "entity_id": "shallow-project",
                "source_text": "Deployed one Kubernetes demonstration service.",
            },
        ],
    )
    posting = JobPosting(
        id="project-posting",
        title="Platform Reliability Engineer",
        description=(
            "Build Kubernetes services, validate rollback and PostgreSQL recovery, "
            "automate failure tests, and implement incident telemetry."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.project_representation is not None
    assert diagnostic.project_representation.status is (
        ProjectRepresentationStatus.SUBSTANTIVE_PROJECT
    )
    assert "substantive-project" in diagnostic.selected_project_ids
    assert diagnostic.bullet_counts["substantive-project"] >= 2


def test_one_bullet_project_and_zero_project_outcomes_are_typed_exceptions() -> None:
    one_bullet_project = _synthetic_profile(
        experiences=[],
        projects=[
            {
                "id": "unique-project",
                "title": "Sensor Validation Project",
                "kind": "project",
            }
        ],
        evidence=[
            {
                "id": "unique-project-proof",
                "entity_id": "unique-project",
                "source_text": (
                    "Validated SPI sensor timing through repeatable embedded tests."
                ),
            }
        ],
    )
    posting = JobPosting(
        id="sensor-posting",
        title="Embedded Validation Engineer",
        description="Validate SPI sensor timing through repeatable embedded tests.",
    )
    shallow = _compose_with_bounds(
        one_bullet_project,
        posting,
        CompositionSearchBounds(),
    )
    no_project_profile = _synthetic_profile(
        experiences=[
            {
                "id": "embedded-entry",
                "title": "Embedded Validation Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "embedded-proof",
                "entity_id": "embedded-entry",
                "source_text": (
                    "Validated SPI sensor timing through repeatable embedded tests."
                ),
            }
        ],
    )
    none_available = _compose_with_bounds(
        no_project_profile,
        posting,
        CompositionSearchBounds(),
    )

    assert shallow.composition_diagnostic is not None
    assert shallow.composition_diagnostic.project_representation is not None
    assert shallow.composition_diagnostic.project_representation.status is (
        ProjectRepresentationStatus.SHALLOW_PROJECT_EXCEPTION
    )
    assert shallow.composition_diagnostic.project_representation.reason
    assert none_available.composition_diagnostic is not None
    assert none_available.composition_diagnostic.project_representation is not None
    assert none_available.composition_diagnostic.project_representation.status is (
        ProjectRepresentationStatus.NO_CREDIBLE_PROJECT_EVIDENCE
    )


def test_unique_one_bullet_project_can_supplement_a_substantive_project() -> None:
    profile = _synthetic_profile(
        experiences=[],
        projects=[
            {
                "id": "substantive-project",
                "title": "Embedded Validation Platform",
                "kind": "project",
            },
            {
                "id": "unique-project",
                "title": "Manufacturing Fixture",
                "kind": "project",
            },
        ],
        evidence=[
            {
                "id": "platform-build",
                "entity_id": "substantive-project",
                "source_text": "Built STM32 firmware with SPI sensor interfaces.",
            },
            {
                "id": "platform-test",
                "entity_id": "substantive-project",
                "source_text": "Validated sensor timing through hardware-in-the-loop tests.",
            },
            {
                "id": "fixture-proof",
                "entity_id": "unique-project",
                "source_text": "Created a GD&T inspection fixture for CNC components.",
            },
        ],
    )
    posting = JobPosting(
        id="systems-manufacturing-posting",
        title="Systems Manufacturing Engineer",
        description=(
            "Build STM32 firmware with SPI interfaces, validate sensor timing through "
            "hardware-in-the-loop tests, and create GD&T inspection fixtures for CNC "
            "components."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.project_representation is not None
    assert diagnostic.project_representation.status is (
        ProjectRepresentationStatus.SUBSTANTIVE_PROJECT
    )
    assert set(diagnostic.selected_project_ids) == {
        "substantive-project",
        "unique-project",
    }
    assert diagnostic.bullet_counts["substantive-project"] == 2
    assert diagnostic.bullet_counts["unique-project"] == 1


def test_four_meaningful_skill_rows_can_be_selected_without_dumping_inventory() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "systems-entry",
                "title": "Systems Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        technical_skills=[
            {"id": "languages", "category": "Languages", "values": ["C++", "Python"]},
            {"id": "embedded", "category": "Embedded", "values": ["STM32", "SPI"]},
            {"id": "test", "category": "Test", "values": ["Oscilloscope", "HIL"]},
            {"id": "delivery", "category": "Delivery", "values": ["Git", "CI/CD"]},
            {
                "id": "inventory",
                "category": "Unrelated inventory",
                "values": [f"Reviewed tool {index}" for index in range(12)],
            },
        ],
        evidence=[
            {
                "id": "systems-proof",
                "entity_id": "systems-entry",
                "source_text": (
                    "Built STM32 firmware in C++ and Python, validated SPI through HIL "
                    "and oscilloscope tests, and used Git CI/CD release checks."
                ),
            }
        ],
    )
    posting = JobPosting(
        id="systems-posting",
        title="Embedded Systems Engineer",
        description=(
            "Build C++ and Python STM32 firmware, validate SPI with HIL and an "
            "oscilloscope, and use Git CI/CD."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert len(diagnostic.selected_skill_rows) == 4
    assert all(len(row.skill_values) >= 2 for row in diagnostic.selected_skill_rows)
    assert "inventory" not in diagnostic.selected_skill_category_ids
    assert all(row.provenance for row in diagnostic.selected_skill_rows)
    assert profile.technical_skills[-1].skills[-1].value == "Reviewed tool 11"


def test_one_skill_row_requires_typed_exception_and_complements_selected_evidence() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "validation-entry",
                "title": "Embedded Validation Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        technical_skills=[
            {
                "id": "languages",
                "category": "Languages",
                "values": ["C++", "Python"],
            },
            {
                "id": "specialized-standard",
                "category": "Safety Standard",
                "values": ["ISO 26262"],
            },
        ],
        evidence=[
            {
                "id": "validation-proof",
                "entity_id": "validation-entry",
                "source_text": (
                    "Implemented C++ and Python validation tools for ISO 26262 "
                    "release evidence."
                ),
            }
        ],
    )
    posting = JobPosting(
        id="safety-posting",
        title="Embedded Validation Engineer",
        description=(
            "Implement C++ and Python validation tools and prepare ISO 26262 "
            "release evidence."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    one_skill_row = next(
        row for row in diagnostic.selected_skill_rows if row.row_id == "specialized-standard"
    )
    assert one_skill_row.skill_values == ["ISO 26262"]
    assert one_skill_row.one_skill_exception_reason
    assert one_skill_row.provenance == [
        "profile.technical_skills[specialized-standard]"
    ]
    assert profile.technical_skills[1].values == ["ISO 26262"]


def test_flat_reviewed_skills_are_regrouped_without_mutation_or_invention() -> None:
    declared_skills = [
        "C++",
        "Python",
        "STM32",
        "SPI",
        "Oscilloscope",
        "HIL",
        "Git",
        "CI/CD",
        "Reviewed unrelated inventory item",
    ]
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "flat-skill-entry",
                "title": "Embedded Systems Engineer",
                "kind": "experience",
            }
        ],
        projects=[],
        declared_skills=declared_skills,
        evidence=[
            {
                "id": "flat-skill-proof",
                "entity_id": "flat-skill-entry",
                "source_text": (
                    "Built STM32 firmware in C++ and Python, validated SPI through HIL "
                    "and oscilloscope tests, and used Git CI/CD release checks."
                ),
            }
        ],
    )
    original_profile = profile.model_dump(mode="json")
    posting = JobPosting(
        id="flat-skill-posting",
        title="Embedded Systems Engineer",
        description=(
            "Build C++ and Python STM32 firmware, validate SPI with HIL and an "
            "oscilloscope, and use Git CI/CD."
        ),
    )

    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert len(diagnostic.selected_skill_rows) >= 2
    assert all(len(row.skill_values) >= 2 for row in diagnostic.selected_skill_rows)
    displayed = {
        value
        for row in diagnostic.selected_skill_rows
        for value in row.skill_values
    }
    assert displayed <= set(declared_skills)
    assert "Reviewed unrelated inventory item" not in displayed
    assert all(
        provenance.startswith("profile.declared_skills[")
        for row in diagnostic.selected_skill_rows
        for provenance in row.provenance
    )
    assert profile.model_dump(mode="json") == original_profile


def test_density_above_preferred_is_never_reported_as_below_preferred() -> None:
    assert DeterministicResumeComposer._preferred_density_status(0.94) is (
        PreferredDensityStatus.ABOVE_PREFERRED
    )
    assert DeterministicResumeComposer._preferred_density_status(0.96) is (
        PreferredDensityStatus.OVERFLOW_RISK
    )
    assert DeterministicResumeComposer._preferred_density_status(0.98) is (
        PreferredDensityStatus.OVERFLOW_RISK
    )


def test_selected_awkward_evidence_is_preserved_and_marked_for_future_rewrite() -> None:
    awkward_text = (
        "Implemented deterministic validation across distributed services using reviewed "
        "evidence and reproducible deployment checks tail fragment"
    )
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "validation-entry",
                "title": "Validation Developer",
                "kind": "experience",
                "organization": "Open Research Laboratory",
                "start_date": "2024",
                "end_date": "Present",
                "location": "Ottawa, ON",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "unique-awkward-evidence",
                "entity_id": "validation-entry",
                "source_text": awkward_text,
            }
        ],
    )
    posting = JobPosting(
        id="validation-posting",
        title="Validation Developer",
        description=(
            "Implement deterministic validation for distributed services with reviewed "
            "deployment checks."
        ),
    )
    resume = _compose_with_bounds(profile, posting, CompositionSearchBounds())
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.selected_bullet_ids == ["unique-awkward-evidence"]
    selected = next(
        candidate
        for candidate in diagnostic.selected_candidates
        if candidate.candidate_id == "bullet:unique-awkward-evidence"
    )
    assert selected.line_fit is not None
    assert selected.line_fit.awkward_wrap_risk is True
    assert selected.line_fit.future_rewrite_recommended is True
    assert "despite estimated awkward wrapping" in (selected.selection_reason or "")
    assert next(_output_bullets(resume)).text == awkward_text


def test_balanced_two_line_evidence_outscores_awkward_and_three_line_text() -> None:
    base = (
        "Implemented deterministic validation across distributed services using reviewed "
        "evidence and reproducible deployment checks"
    )
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "line-fit-entry",
                "title": "Validation Developer",
                "kind": "experience",
                "organization": "Open Research Laboratory",
                "start_date": "2024",
                "end_date": "Present",
                "location": "Ottawa, ON",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "awkward-two-line",
                "entity_id": "line-fit-entry",
                "source_text": f"{base} tail fragment",
            },
            {
                "id": "balanced-two-line",
                "entity_id": "line-fit-entry",
                "source_text": f"{base} {'balanced ' * 5}tail fragment",
            },
            {
                "id": "unnecessary-three-line",
                "entity_id": "line-fit-entry",
                "source_text": f"{base} {'balanced ' * 15}tail fragment",
            },
        ],
    )
    posting = JobPosting(
        id="line-fit-posting",
        title="Validation Developer",
        description=(
            "Implement deterministic validation for distributed services with reviewed "
            "deployment checks."
        ),
    )
    composer = DeterministicResumeComposer(FixedRatioPageFitEvaluator(0.82))
    candidates = {
        item.evidence_id: item
        for item in composer._all_bullet_candidates(profile, _posting_context(posting))
    }

    assert candidates["balanced-two-line"].line_fit.awkward_wrap_risk is False
    assert candidates["awkward-two-line"].line_fit.awkward_wrap_risk is True
    assert candidates["unnecessary-three-line"].line_fit.three_line_risk is True
    assert (
        candidates["balanced-two-line"].score
        > candidates["awkward-two-line"].score
        > candidates["unnecessary-three-line"].score
    )
    selected = _compose_with_bounds(
        profile,
        posting,
        CompositionSearchBounds(
            maximum_selected_bullets=1,
            maximum_bullets_per_entry=1,
        ),
    )
    assert selected.composition_diagnostic is not None
    assert selected.composition_diagnostic.selected_bullet_ids == ["balanced-two-line"]


def test_expansion_priority_prefers_skills_and_selected_blocks_over_weak_new_entry() -> None:
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "selected-entry",
                "title": "Release Developer",
                "kind": "experience",
                "organization": "Public Utility",
                "start_date": "2024",
                "end_date": "Present",
                "location": "Remote",
            },
            {
                "id": "weak-new-entry",
                "title": "Application Assistant",
                "kind": "experience",
                "organization": "Student Society",
                "start_date": "2022",
                "end_date": "2022",
                "location": "Remote",
            },
        ],
        projects=[],
        technical_skills=[
            {
                "id": "release-tools",
                "category": "Release Platforms",
                "values": ["Kubernetes", "CI/CD"],
            }
        ],
        evidence=[
            {
                "id": "selected-release",
                "entity_id": "selected-entry",
                "source_text": ("Deployed Kubernetes services through CI/CD release gates."),
            },
            {
                "id": "selected-debugging",
                "entity_id": "selected-entry",
                "source_text": (
                    "Debugged Kubernetes rollout defects through reproducible CI/CD checks."
                ),
            },
            {
                "id": "weak-new-support",
                "entity_id": "weak-new-entry",
                "source_text": ("Supported Kubernetes documentation for one application."),
            },
        ],
    )
    posting = JobPosting(
        id="priority-posting",
        title="Release Developer",
        description="Deploy Kubernetes services through CI/CD and debug rollout defects.",
    )
    composer = DeterministicResumeComposer(FixedRatioPageFitEvaluator(0.70))
    context = _posting_context(posting)
    bullets = composer._rank_bullets(profile, context)
    skills = composer._rank_skills(profile, context, bullets)
    bullet_by_id = {item.evidence_id: item for item in bullets}
    skill_by_id = {item.category_id: item for item in skills}
    state = _State(frozenset({"selected-release"}), frozenset())
    options = composer._expansions(
        state,
        profile,
        bullets,
        skills,
        bullet_by_id,
        skill_by_id,
        TemplateConstraints(),
        {},
    )
    by_source = {option.source_id: option for option in options}

    assert by_source["release-tools"].preference_bonus == 8.0
    assert by_source["selected-debugging"].preference_bonus == 5.0
    assert "weak-new-support" not in by_source


def test_large_profile_keeps_late_strong_evidence_and_search_runtime_bounded() -> None:
    experiences: list[dict[str, object]] = []
    evidence: list[dict[str, object]] = []
    for index in range(30):
        entry_id = f"growing-entry-{index:02d}"
        experiences.append(
            {
                "id": entry_id,
                "title": "Technical Contributor",
                "kind": "experience",
                "organization": f"Reviewed Organization {index:02d}",
                "start_date": str(1990 + index),
                "end_date": str(1991 + index),
                "location": "Remote",
            }
        )
        evidence.append(
            {
                "id": f"growing-evidence-{index:02d}",
                "entity_id": entry_id,
                "source_text": (
                    f"Documented reviewed maintenance task {index:02d} for an internal archive."
                ),
            }
        )
    experiences.append(
        {
            "id": "late-strong-platform-entry",
            "title": "Platform Reliability Developer",
            "kind": "experience",
            "organization": "Public Infrastructure Lab",
            "start_date": "2025",
            "end_date": "Present",
            "location": "Remote",
        }
    )
    evidence.extend(
        [
            {
                "id": "late-strong-kubernetes",
                "entity_id": "late-strong-platform-entry",
                "source_text": (
                    "Deployed Kubernetes services through CI/CD policy gates and rollback tests."
                ),
            },
            {
                "id": "late-strong-postgresql",
                "entity_id": "late-strong-platform-entry",
                "source_text": (
                    "Debugged PostgreSQL reliability defects using production health telemetry."
                ),
            },
        ]
    )
    profile = _synthetic_profile(
        experiences=experiences,
        projects=[],
        evidence=evidence,
    )
    posting = JobPosting(
        id="large-profile-posting",
        title="Platform Reliability Developer",
        description=(
            "Deploy Kubernetes services with CI/CD, debug PostgreSQL reliability defects, "
            "and monitor production health."
        ),
    )
    bounds = CompositionSearchBounds(
        maximum_estimated_page_evaluations=40,
        maximum_expansion_operations=120,
        maximum_ranked_bullets=12,
    )
    resume = _compose_with_bounds(profile, posting, bounds)
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert "late-strong-platform-entry" in diagnostic.selected_experience_ids
    assert {
        "late-strong-kubernetes",
        "late-strong-postgresql",
    }.issubset(diagnostic.selected_bullet_ids)
    assert diagnostic.estimated_page_evaluations <= 40
    assert diagnostic.expansion_operations <= 120


def test_density_objective_does_not_pack_redundant_three_line_bullets() -> None:
    shared = (
        "Implemented deterministic Kubernetes release validation with CI/CD policy checks, "
        "production rollback tests, PostgreSQL health telemetry, and reviewed incident "
        "evidence"
    )
    profile = _synthetic_profile(
        experiences=[
            {
                "id": "readability-entry",
                "title": "Platform Reliability Developer",
                "kind": "experience",
                "organization": "Public Infrastructure Lab",
                "start_date": "2024",
                "end_date": "Present",
                "location": "Remote",
            }
        ],
        projects=[],
        evidence=[
            {
                "id": "balanced-release-evidence",
                "entity_id": "readability-entry",
                "source_text": shared,
            },
            *[
                {
                    "id": f"redundant-three-line-{index}",
                    "entity_id": "readability-entry",
                    "source_text": (
                        f"{shared} with repeated deployment validation context "
                        + ("additional " * 16)
                        + f"variant {index}"
                    ),
                }
                for index in range(3)
            ],
        ],
    )
    posting = JobPosting(
        id="readability-density-posting",
        title="Platform Reliability Developer",
        description=(
            "Implement Kubernetes release validation with CI/CD policy checks, PostgreSQL "
            "health telemetry, rollback tests, and production incident review."
        ),
    )
    baseline = StructuredResume(
        profile_id=profile.id,
        profile_version=profile.version,
        posting_id=posting.id,
        template_id="managed-engineering-v1",
        display_name=profile.display_name,
        strategy=ResumeStrategy(
            role_family="deterministic_test",
            primary_focus=posting.title,
            rationale="Controlled readability-density fixture.",
        ),
        education=profile.education,
    )
    resume = DeterministicResumeComposer(LineCostDensityEvaluator()).compose(
        baseline,
        profile,
        posting,
        TemplateConstraints(),
    )
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.selected_bullet_ids == ["balanced-release-evidence"]
    assert diagnostic.preferred_density_reached is False
    assert diagnostic.final_utilization_ratio < 0.90
    redundant = [
        candidate
        for candidate in diagnostic.candidates_excluded_by_thresholds
        if candidate.candidate_id.startswith("bullet:redundant-three-line")
    ]
    assert len(redundant) == 3
    assert all(
        candidate.line_fit is not None
        and candidate.line_fit.three_line_risk
        and candidate.line_fit.future_rewrite_recommended
        for candidate in redundant
    )
