from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from resume_tailor.application.profile_editor import (
    editor_state_to_profile,
    profile_to_editor_state,
)
from resume_tailor.domain.models import MasterProfile, ResumeStrategy, StructuredResume
from resume_tailor.infrastructure.optimization import EvidenceBoundResumeWriter
from resume_tailor.infrastructure.profile_repository import SQLiteMasterProfileRepository
from resume_tailor.infrastructure.resume_extraction import (
    EmptyResumeFileError,
    ImageOnlyResumeError,
    ResumeExtractionError,
    UnsupportedResumeFileError,
    extract_resume_text,
)
from resume_tailor.infrastructure.static_template_docx import render_template_v1_resume


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Jane Candidate")
    document.add_paragraph("Engineer | Toronto")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _add_hyperlink(paragraph: object, display: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = display
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _contact_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Alex Candidate | alex@example.test | +1 555 010 0200")
    table = document.add_table(rows=1, cols=1)
    _add_hyperlink(
        table.cell(0, 0).paragraphs[0],
        "Portfolio",
        "https://portfolio.example.test/alex",
    )
    header = document.sections[0].header
    _add_hyperlink(
        header.paragraphs[0],
        "LinkedIn",
        "https://www.linkedin.com/in/synthetic-candidate",
    )
    footer = document.sections[0].footer
    _add_hyperlink(
        footer.paragraphs[0],
        "GitHub",
        "https://github.com/synthetic-candidate",
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_bytes(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(output.tell())
        output.write(f"{index} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode())
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode())
    output.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
    )
    return output.getvalue()


def test_docx_text_extraction() -> None:
    result = extract_resume_text("resume.docx", _docx_bytes())
    assert result.source_format == "docx"
    assert "Jane Candidate" in result.text
    assert "Engineer | Toronto" in result.text


def test_docx_contact_links_from_tables_headers_and_footers_survive_save_and_render(
    tmp_path,
) -> None:
    result = extract_resume_text("synthetic-contact.docx", _contact_docx_bytes())
    expected_links = [
        "https://portfolio.example.test/alex",
        "https://www.linkedin.com/in/synthetic-candidate",
        "https://github.com/synthetic-candidate",
    ]
    assert "alex@example.test" in result.text
    assert "+1 555 010 0200" in result.text
    assert all(result.text.count(link) == 1 for link in expected_links)
    assert [link.model_dump() for link in result.contact_links] == [
        {
            "display_text": "Portfolio",
            "destination": "https://portfolio.example.test/alex",
        },
        {
            "display_text": "LinkedIn",
            "destination": "https://www.linkedin.com/in/synthetic-candidate",
        },
        {
            "display_text": "GitHub",
            "destination": "https://github.com/synthetic-candidate",
        },
    ]

    imported = MasterProfile.model_validate(
        {
            "id": "synthetic-contact-profile",
            "user_id": "synthetic-user",
            "display_name": "Alex Candidate",
            "contact": {
                "email": "alex@example.test",
                "phone": "+1 555 010 0200",
                "location": "Toronto, ON",
                "links": expected_links,
                "hyperlinks": [link.model_dump() for link in result.contact_links],
            },
        }
    )
    reviewed = editor_state_to_profile(profile_to_editor_state(imported))
    repository = SQLiteMasterProfileRepository(tmp_path / "profiles.sqlite3")
    repository.save(reviewed)
    loaded = repository.get(reviewed.id)
    assert loaded is not None
    contact_line = EvidenceBoundResumeWriter._contact_line(loaded)
    contact_items = EvidenceBoundResumeWriter._contact_items(loaded)
    resume = StructuredResume(
        profile_id=loaded.id,
        profile_version=loaded.version,
        posting_id="synthetic-posting",
        template_id="managed-engineering-v1",
        display_name=loaded.display_name,
        contact_line=contact_line,
        contact_items=contact_items,
        strategy=ResumeStrategy(
            role_family="embedded_systems",
            primary_focus="Embedded systems",
            rationale="Synthetic rendering regression.",
        ),
    )
    output = tmp_path / "contact.docx"
    render_template_v1_resume(resume, output)
    with ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode()
        relationships = package.read("word/_rels/document.xml.rels").decode()

    assert loaded.contact.links == expected_links
    assert loaded.contact.hyperlinks == list(result.contact_links)
    assert contact_line is not None
    assert [item.display_text for item in contact_items] == [
        "Email",
        "+1 555 010 0200",
        "Toronto, ON",
        "Portfolio",
        "LinkedIn",
        "GitHub",
    ]
    assert all(value in relationships for value in ["mailto:alex@example.test", *expected_links])
    assert all(label in document_xml for label in ["Portfolio", "LinkedIn", "GitHub"])
    assert not any(value in document_xml for value in expected_links)
    assert "+1 555 010 0200" in document_xml
    assert "Toronto, ON" in document_xml


def test_text_pdf_extraction() -> None:
    result = extract_resume_text("resume.pdf", _pdf_bytes("Jane Candidate"))
    assert result.source_format == "pdf"
    assert "Jane Candidate" in result.text


def test_unsupported_corrupt_empty_and_image_only_files_are_clear() -> None:
    with pytest.raises(UnsupportedResumeFileError):
        extract_resume_text("resume.txt", b"text")
    with pytest.raises(EmptyResumeFileError):
        extract_resume_text("resume.docx", b"")
    with pytest.raises(ResumeExtractionError):
        extract_resume_text("resume.docx", b"not-a-docx")
    with pytest.raises(ImageOnlyResumeError):
        extract_resume_text("resume.pdf", _pdf_bytes(""))
