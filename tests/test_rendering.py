import subprocess
from dataclasses import dataclass
from pathlib import Path

import pytest
from docx import Document

from resume_tailor.domain.models import (
    ClaimSupport,
    ResumeStrategy,
    StructuredBullet,
    StructuredResume,
)
from resume_tailor.infrastructure.rendering import (
    ManagedResumeRenderer,
    MicrosoftWordDocxPageCountProvider,
    PageCountMeasurement,
    PageCountVerificationError,
    PageOverflowError,
)


@dataclass
class _SequencePageCountProvider:
    counts: list[int]

    def measure(self, docx_path: Path) -> PageCountMeasurement:
        count = self.counts.pop(0) if len(self.counts) > 1 else self.counts[0]
        return PageCountMeasurement(
            page_count=count,
            provider="test-sequence-provider",
            confidence="exact",
            exact=True,
        )


class _EstimatedPageCountProvider:
    def measure(self, docx_path: Path) -> PageCountMeasurement:
        return PageCountMeasurement(
            page_count=1,
            provider="estimate-only",
            confidence="estimated",
            exact=False,
        )


@dataclass
class _PathCapturingPageCountProvider:
    path: Path | None = None

    def measure(self, docx_path: Path) -> PageCountMeasurement:
        self.path = docx_path
        assert docx_path.is_absolute()
        assert docx_path.is_file()
        assert docx_path.stat().st_size > 0
        return PageCountMeasurement(
            page_count=1,
            provider="path-capturing-provider",
            confidence="exact",
            exact=True,
        )


def test_managed_renderer_exports_docx_and_one_page_pdf(tmp_path: Path) -> None:
    resume = StructuredResume(
        profile_id="profile-1",
        profile_version=1,
        posting_id="posting-1",
        template_id="managed-engineering-v1",
        display_name="Avery Engineer",
        strategy=ResumeStrategy(
            role_family="embedded_firmware",
            primary_focus="embedded firmware development",
            rationale="Use verified evidence.",
        ),
        entity_titles={"experience-1": "Firmware Intern"},
        experience_bullets={
            "experience-1": [
                StructuredBullet(
                    id="evidence-1",
                    text="Developed STM32 firmware and validated SPI sensor communication.",
                    evidence_ids=["evidence-1"],
                    support=ClaimSupport.DIRECT,
                )
            ]
        },
        selected_skills=["STM32", "C", "SPI"],
    )

    result = ManagedResumeRenderer(page_count_provider=_SequencePageCountProvider([1])).render(
        resume, tmp_path
    )

    assert result.page_count == 1
    assert result.exact_page_count_verified is True
    assert result.docx_path.exists()
    assert result.pdf_path.read_bytes().startswith(b"%PDF")


def test_page_count_provider_receives_absolute_existing_docx_with_spaces(tmp_path: Path) -> None:
    provider = _PathCapturingPageCountProvider()
    output = tmp_path / "directory with spaces" / "nested output"
    resume = StructuredResume(
        profile_id="profile-path",
        profile_version=1,
        posting_id="posting-path",
        template_id="managed-engineering-v1",
        display_name="Candidate",
        strategy=ResumeStrategy(
            role_family="embedded_firmware",
            primary_focus="verified engineering work",
            rationale="Use verified evidence.",
        ),
    )

    result = ManagedResumeRenderer(page_count_provider=provider).render(resume, output)

    assert provider.path is not None
    assert provider.path.parent == output.resolve()
    assert result.docx_path.is_absolute()


def test_managed_renderer_rejects_overflow(tmp_path: Path) -> None:
    bullet = StructuredBullet(
        id="evidence-overflow",
        text="Verified firmware validation evidence " * 30,
        evidence_ids=["evidence-overflow"],
        support=ClaimSupport.DIRECT,
    )
    resume = StructuredResume(
        profile_id="profile-1",
        profile_version=1,
        posting_id="posting-1",
        template_id="managed-engineering-v1",
        display_name="Avery Engineer",
        strategy=ResumeStrategy(
            role_family="embedded_firmware",
            primary_focus="embedded firmware development",
            rationale="Use verified evidence.",
        ),
        entity_titles={"experience-1": "Firmware Intern"},
        experience_bullets={"experience-1": [bullet] * 30},
    )

    with pytest.raises(PageOverflowError):
        ManagedResumeRenderer(page_count_provider=_SequencePageCountProvider([1])).render(
            resume, tmp_path
        )


def test_strict_docx_page_gate_reduces_optional_content_until_one_page(tmp_path: Path) -> None:
    bullet = StructuredBullet(
        id="evidence-reduce",
        text="Verified firmware validation evidence.",
        evidence_ids=["evidence-reduce"],
        support=ClaimSupport.DIRECT,
    )
    resume = StructuredResume(
        profile_id="profile-1",
        profile_version=1,
        posting_id="posting-1",
        template_id="managed-engineering-v1",
        display_name="Avery Engineer",
        strategy=ResumeStrategy(
            role_family="embedded_firmware",
            primary_focus="embedded firmware development",
            rationale="Use verified evidence.",
        ),
        entity_titles={"experience-1": "Firmware Intern"},
        experience_bullets={"experience-1": [bullet]},
    )
    provider = _SequencePageCountProvider([2, 1])
    renderer = ManagedResumeRenderer(page_count_provider=provider)
    output = renderer.render_docx(resume, tmp_path / "reduced.docx")

    assert output.exists()
    assert renderer.last_measurement is not None
    assert renderer.last_measurement.page_count == 1
    assert renderer.last_measurement.exact is True
    assert renderer.last_overflow_reduction_count == 1
    assert "Verified firmware" not in "\n".join(
        paragraph.text for paragraph in Document(output).paragraphs
    )


def test_docx_page_gate_returns_typed_unverified_estimate(tmp_path: Path) -> None:
    renderer = ManagedResumeRenderer(page_count_provider=_EstimatedPageCountProvider())
    resume = StructuredResume(
        profile_id="profile-1",
        profile_version=1,
        posting_id="posting-1",
        template_id="managed-engineering-v1",
        display_name="Avery Engineer",
        strategy=ResumeStrategy(
            role_family="embedded_firmware",
            primary_focus="embedded firmware development",
            rationale="Use verified evidence.",
        ),
    )

    output = renderer.render_docx(resume, tmp_path / "estimated.docx")

    assert output.exists()
    assert renderer.last_measurement is not None
    assert renderer.last_measurement.exact is False
    assert renderer.last_page_utilization is not None
    assert renderer.last_page_utilization.status.value == "unverified"
    assert renderer.last_verification_failure is not None
    assert "not exact" in renderer.last_verification_failure


def test_underfill_expansion_is_disabled_and_geometry_is_not_a_page_fit_control(
    tmp_path: Path,
) -> None:
    provider = _SequencePageCountProvider([2, 1])
    renderer = ManagedResumeRenderer(page_count_provider=provider)
    assert renderer.underfill_expansion_enabled is False

    bullet = StructuredBullet(
        id="evidence-geometry",
        text="Verified firmware validation evidence.",
        evidence_ids=["evidence-geometry"],
        support=ClaimSupport.DIRECT,
    )
    resume = StructuredResume(
        profile_id="profile-1",
        profile_version=1,
        posting_id="posting-1",
        template_id="managed-engineering-v1",
        display_name="Avery Engineer",
        strategy=ResumeStrategy(
            role_family="embedded_firmware",
            primary_focus="embedded firmware development",
            rationale="Use verified evidence.",
        ),
        entity_titles={"experience-1": "Firmware Intern"},
        experience_bullets={"experience-1": [bullet]},
    )
    output = renderer.render_docx(resume, tmp_path / "geometry-stable.docx")
    document = Document(output)
    page = renderer.layout_profile.page
    section = document.sections[0]
    assert section.left_margin.twips == page.left_margin_twips
    assert section.right_margin.twips == page.right_margin_twips
    assert section.top_margin.twips == page.top_margin_twips
    assert section.bottom_margin.twips == page.bottom_margin_twips


def test_word_timeout_returns_typed_failure_and_cleans_only_owned_process(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "candidate.docx"
    docx_path.write_bytes(b"controlled-docx")
    commands: list[str] = []

    def fake_run(command: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        commands.append(command[-1])
        if len(commands) == 1:
            environment = kwargs["env"]
            assert isinstance(environment, dict)
            Path(environment["RESUME_WORD_OWNED_PROCESS_PATH"]).write_text(
                "4321|638885280000000000",
                encoding="ascii",
            )
            raise subprocess.TimeoutExpired(command, timeout=0.01)
        return subprocess.CompletedProcess(command, 0, "", "")

    monkeypatch.setattr(subprocess, "run", fake_run)
    monkeypatch.setattr(
        "resume_tailor.infrastructure.rendering.shutil.which",
        lambda _name: "powershell.exe",
    )

    with pytest.raises(PageCountVerificationError, match="timed out after"):
        MicrosoftWordDocxPageCountProvider(timeout_seconds=0.01).measure(docx_path)

    assert len(commands) == 2
    assert "Stop-Process -Id $ownedProcessId" in commands[1]
    assert "Get-Process WINWORD" not in "".join(commands)


def test_word_batch_uses_one_com_session_for_all_finalists(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    paths = [tmp_path / "one.docx", tmp_path / "two.docx"]
    for path in paths:
        path.write_bytes(b"controlled-docx")
    calls = 0

    def fake_run(command: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        nonlocal calls
        calls += 1
        environment = kwargs["env"]
        assert isinstance(environment, dict)
        path_lines = Path(environment["RESUME_DOCX_PATHS_FILE"]).read_text(
            encoding="utf-8"
        )
        assert len(path_lines.splitlines()) == 2
        return subprocess.CompletedProcess(command, 0, "1,2\n", "")

    monkeypatch.setattr(subprocess, "run", fake_run)
    monkeypatch.setattr(
        "resume_tailor.infrastructure.rendering.shutil.which",
        lambda _name: "powershell.exe",
    )

    measurements = MicrosoftWordDocxPageCountProvider().measure_many(paths)

    assert calls == 1
    assert [measurement.page_count for measurement in measurements] == [1, 2]
