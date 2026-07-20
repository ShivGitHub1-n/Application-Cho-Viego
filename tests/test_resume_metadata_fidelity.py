from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from resume_tailor.application.resume_composition import (
    CompositionSearchBounds,
    DeterministicResumeComposer,
)
from resume_tailor.application.services import TailorResumeService
from resume_tailor.domain.models import (
    EducationRecord,
    JobPosting,
    MasterProfile,
    StructuredResume,
    TemplateConstraints,
)
from resume_tailor.domain.resume_composition import (
    CompositionTerminationReason,
    CompositionUnderfillReason,
    PageFitEvaluation,
)
from resume_tailor.domain.resume_metadata import (
    DatePrecision,
    ResumeMetadataIntegrityError,
    compose_date_range,
    validate_structured_resume_metadata,
)
from resume_tailor.infrastructure.composition_page_fit import TemplateV1PageFitEvaluator
from resume_tailor.infrastructure.optimization import (
    DeterministicResumeOptimizer,
    EvidenceBoundResumeWriter,
)
from resume_tailor.infrastructure.rendering import (
    PageCountMeasurement,
    PageCountVerificationError,
    diagnose_docx_page_utilization,
)
from resume_tailor.infrastructure.static_template_docx import render_template_v1_resume
from resume_tailor.infrastructure.template_v1 import load_template_v1_layout_profile
from tests.convergence_cases import (
    mechanical_manufacturing_case,
    rich_mixed_case,
    software_cloud_case,
)


class FailingExactPageProvider:
    def measure(self, docx_path: Path) -> PageCountMeasurement:
        raise PageCountVerificationError(f"Controlled exact provider failure for {docx_path.name}")


class FixedPageFit:
    def __init__(self, utilization: float = 0.91) -> None:
        self.utilization = utilization
        self.calls = 0

    def evaluate(
        self,
        resume: StructuredResume,
        *,
        attempt_exact: bool = True,
    ) -> PageFitEvaluation:
        self.calls += 1
        return PageFitEvaluation(
            status="acceptable_one_page",
            page_count=1,
            exact=attempt_exact,
            provider="controlled exact metadata evaluator",
            utilization_ratio=self.utilization,
            fits_one_page=True,
        )


def _service(
    *,
    bounds: CompositionSearchBounds | None = None,
) -> TailorResumeService:
    return TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
        resume_composer=DeterministicResumeComposer(
            TemplateV1PageFitEvaluator(FailingExactPageProvider()),
            bounds=bounds,
        ),
    )


def _compose(
    profile: MasterProfile,
    posting: JobPosting,
    *,
    bounds: CompositionSearchBounds | None = None,
) -> StructuredResume:
    service = _service(bounds=bounds)
    plan = service.create_plan(profile, posting, TemplateConstraints())
    return service.build_document(plan, profile, set())


def test_multiple_bullets_and_evidence_records_render_entry_metadata_once(
    tmp_path: Path,
) -> None:
    profile, posting = mechanical_manufacturing_case()
    resume = _compose(profile, posting)
    selected = next(
        item for item in resume.experiences if len(resume.experience_bullets[item.id]) >= 3
    )
    output = tmp_path / "metadata-once.docx"

    render_template_v1_resume(resume, output)
    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
    rendered_date = compose_date_range(selected.start_date, selected.end_date)

    assert rendered_date is not None
    assert sum(rendered_date in paragraph for paragraph in paragraphs) == 1
    assert sum(selected.title in paragraph for paragraph in paragraphs) == 1
    assert (
        sum(
            (selected.organization or "") in paragraph
            for paragraph in paragraphs
            if selected.organization
        )
        == 1
    )
    assert len(resume.experience_bullets[selected.id]) >= 3


def test_month_and_year_date_precision_survives_complete_pipeline_and_rendering(
    tmp_path: Path,
) -> None:
    profile, posting = software_cloud_case()
    resume = _compose(profile, posting)
    report = validate_structured_resume_metadata(resume)
    output = tmp_path / "mixed-date-precision.docx"
    render_template_v1_resume(resume, output)
    rendered_text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

    by_id = {item.entry_id: item for item in report.entries}
    assert by_id["backend-platform"].source_start_date == "May 2023"
    assert by_id["backend-platform"].start_date_precision is DatePrecision.MONTH_YEAR
    assert by_id["backend-platform"].rendered_date_text == "May 2023 \u2013 Present"
    assert "May 2023 \u2013 Present" in rendered_text
    assert by_id["application-security"].start_date_precision is DatePrecision.YEAR
    assert by_id["application-security"].end_date_precision is DatePrecision.YEAR
    assert by_id["application-security"].rendered_date_text == "2020 \u2013 2021"
    assert "2020 \u2013 2021" in rendered_text
    assert "Jan. 2021 \u2013 Apr. 2023" in rendered_text


def test_composed_metadata_rejects_accumulated_ranges_and_repeated_fields(
    tmp_path: Path,
) -> None:
    profile, posting = software_cloud_case()
    resume = _compose(profile, posting)
    entry = resume.experiences[0]
    malformed_entry = entry.model_copy(
        update={
            "start_date": "May 2023 \u2013 Present \u2013 May 2023 \u2013 Present",
            "end_date": "May 2023 \u2013 Present",
        }
    )
    malformed = resume.model_copy(
        update={
            "experiences": [malformed_entry, *resume.experiences[1:]],
        }
    )

    with pytest.raises(
        ResumeMetadataIntegrityError,
        match="accumulated date range|repeated composed metadata",
    ):
        render_template_v1_resume(malformed, tmp_path / "malformed.docx")

    repeated_organization = entry.model_copy(
        update={"organization": "Digital Services Lab | Digital Services Lab"}
    )
    malformed = resume.model_copy(
        update={
            "experiences": [repeated_organization, *resume.experiences[1:]],
        }
    )
    with pytest.raises(ResumeMetadataIntegrityError, match="repeated composed metadata"):
        validate_structured_resume_metadata(malformed)


def test_rich_education_survives_and_missing_optional_fields_are_not_invented(
    tmp_path: Path,
) -> None:
    rich_profile, posting = mechanical_manufacturing_case()
    rich_resume = _compose(rich_profile, posting)
    rich_output = tmp_path / "rich-education.docx"
    render_template_v1_resume(rich_resume, rich_output)
    rich_text = "\n".join(paragraph.text for paragraph in Document(rich_output).paragraphs)

    assert "Manufacturing systems option" in rich_text
    assert "Co-operative Education" in rich_text
    assert "GPA: 3.7/4.0" in rich_text
    assert "Awards: Manufacturing Design Award" in rich_text
    assert "Relevant Courses: Machine Design" in rich_text

    sparse_education = EducationRecord(
        school="Public Polytechnic University",
        program="Bachelor of Engineering in Mechanical Engineering",
    )
    sparse_profile = rich_profile.model_copy(
        update={"education": [sparse_education], "coursework": []}
    )
    sparse_resume = _compose(sparse_profile, posting)
    sparse_output = tmp_path / "sparse-education.docx"
    render_template_v1_resume(sparse_resume, sparse_output)
    sparse_text = "\n".join(paragraph.text for paragraph in Document(sparse_output).paragraphs)

    assert "GPA:" not in sparse_text
    assert "Awards:" not in sparse_text
    assert "Relevant Courses:" not in sparse_text
    assert "Expected" not in sparse_text


def test_education_details_participate_in_template_page_budget(tmp_path: Path) -> None:
    profile, posting = mechanical_manufacturing_case()
    rich_resume = _compose(profile, posting)
    sparse_resume = rich_resume.model_copy(
        update={
            "education": [
                EducationRecord(
                    school=profile.education[0].school,
                    program=profile.education[0].program,
                )
            ],
            "selected_coursework": [],
        }
    )
    profile_layout = load_template_v1_layout_profile()
    measurement = PageCountMeasurement(
        page_count=1,
        provider="controlled estimate",
        confidence="estimated",
        exact=False,
    )
    rich_path = tmp_path / "rich.docx"
    sparse_path = tmp_path / "sparse.docx"
    render_template_v1_resume(rich_resume, rich_path)
    render_template_v1_resume(sparse_resume, sparse_path)

    rich = diagnose_docx_page_utilization(rich_path, profile_layout, measurement)
    sparse = diagnose_docx_page_utilization(sparse_path, profile_layout, measurement)

    assert rich.estimated_utilization_ratio > sparse.estimated_utilization_ratio
    assert len(rich_resume.education[0].relevant_coursework) == 3


def test_mechanical_completion_deepens_selected_entries_to_target() -> None:
    profile, posting = mechanical_manufacturing_case()

    resume = _compose(profile, posting)
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert diagnostic.final_utilization_ratio >= 0.90
    assert diagnostic.best_estimated_utilization_ratio >= 0.90
    assert diagnostic.termination_reason is CompositionTerminationReason.TARGET_FINALISTS_FOUND
    assert CompositionUnderfillReason.SEARCH_BOUNDS_LIMITED not in diagnostic.underfill_reasons
    assert sum(diagnostic.bullet_counts.values()) >= 17
    assert diagnostic.bullet_counts["modular-fixture"] >= 3
    assert diagnostic.bullet_counts["robot-chassis"] >= 3


def test_fuller_mechanical_plan_defeats_shallow_bounds_limited_plan() -> None:
    profile, posting = mechanical_manufacturing_case()
    shallow = _compose(
        profile,
        posting,
        bounds=CompositionSearchBounds(
            beam_width=2,
            maximum_estimated_page_evaluations=24,
            maximum_exact_finalist_evaluations=2,
            maximum_expansion_operations=120,
            maximum_expansions_per_state=2,
            maximum_selected_bullets=10,
        ),
    )
    converged = _compose(profile, posting)
    shallow_diagnostic = shallow.composition_diagnostic
    converged_diagnostic = converged.composition_diagnostic

    assert shallow_diagnostic is not None
    assert converged_diagnostic is not None
    assert converged_diagnostic.final_utilization_ratio > (
        shallow_diagnostic.final_utilization_ratio + 0.08
    )
    assert sum(converged_diagnostic.bullet_counts.values()) > sum(
        shallow_diagnostic.bullet_counts.values()
    )
    assert converged_diagnostic.termination_reason is (
        CompositionTerminationReason.TARGET_FINALISTS_FOUND
    )


@pytest.mark.parametrize(
    "case_factory",
    [software_cloud_case, rich_mixed_case],
)
def test_software_and_mixed_controlled_plans_remain_in_preferred_band(
    case_factory: object,
) -> None:
    profile, posting = case_factory()  # type: ignore[operator]
    resume = _compose(profile, posting)
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert 0.90 <= diagnostic.final_utilization_ratio <= 0.95
    assert len(diagnostic.selected_experience_ids) >= 3
    assert len(diagnostic.selected_project_ids) >= 2
    assert len(diagnostic.selected_skill_category_ids) >= 3


def test_irrelevant_added_content_is_not_used_as_density_filler() -> None:
    profile, posting = mechanical_manufacturing_case()
    unrelated_profile, _ = software_cloud_case()
    unrelated_entry = unrelated_profile.experiences[0].model_copy(
        update={"id": "unrelated-marketing-entry"}
    )
    unrelated_evidence = unrelated_profile.evidence[0].model_copy(
        update={
            "id": "unrelated-marketing-evidence",
            "entity_id": unrelated_entry.id,
            "source_text": (
                "Supported social media calendars and drafted general promotional copy."
            ),
        }
    )
    grown = profile.model_copy(
        update={
            "experiences": [*profile.experiences, unrelated_entry],
            "evidence": [*profile.evidence, unrelated_evidence],
        }
    )

    resume = _compose(grown, posting)
    diagnostic = resume.composition_diagnostic

    assert diagnostic is not None
    assert unrelated_entry.id not in diagnostic.selected_experience_ids
    assert unrelated_evidence.id not in diagnostic.selected_bullet_ids


def test_page_fit_iterations_do_not_call_any_llm_provider() -> None:
    profile, posting = mechanical_manufacturing_case()
    evaluator = FixedPageFit()
    service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
        resume_composer=DeterministicResumeComposer(evaluator),
    )
    resume = service.build_document(
        service.create_plan(profile, posting, TemplateConstraints()),
        profile,
        set(),
    )

    assert resume.hybrid_diagnostic is not None
    assert resume.hybrid_diagnostic.provider_call_count == 0
    assert evaluator.calls > 1
