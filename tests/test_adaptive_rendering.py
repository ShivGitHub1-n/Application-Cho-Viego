import importlib.util
from pathlib import Path
from types import ModuleType
from zipfile import ZipFile

from docx import Document

from resume_tailor.domain.models import (
    ClaimSupport,
    EducationRecord,
    EntityKind,
    ResumeItem,
    ResumeStrategy,
    StructuredBullet,
    StructuredResume,
    TechnicalSkillCategory,
)
from resume_tailor.infrastructure.adaptive_docx import render_structured_resume
from resume_tailor.infrastructure.reference_docx import analyze_reference_docx

REFERENCE = Path("manual-test/reference-resume.docx")


def _load_manual_render_module() -> ModuleType:
    path = Path("manual-test/render_deterministic_docx.py")
    spec = importlib.util.spec_from_file_location("manual_render_deterministic_docx", path)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _resume() -> StructuredResume:
    experience = ResumeItem(
        id="experience-alpha",
        title="Systems Integration Engineer",
        kind=EntityKind.EXPERIENCE,
        organization="Example Research Cooperative",
        start_date="May 2024",
        end_date="Aug 2025",
        location="Example City, ZZ",
        subtitle="Embedded Controls",
    )
    project = ResumeItem(
        id="project-beta",
        title="Adaptive Sensor Platform",
        kind=EntityKind.PROJECT,
        start_date="Jan 2024",
        end_date="Apr 2024",
        technology_label="Python, ROS 2",
    )
    return StructuredResume(
        profile_id="adaptive-profile",
        profile_version=1,
        posting_id="adaptive-posting",
        template_id="reference-derived",
        display_name="Alex Example",
        contact_line=(
            "alex@example.test | +1 555 0100 | example.test/portfolio | Example City, ZZ"
        ),
        strategy=ResumeStrategy(
            role_family="software_data_engineering",
            primary_focus="systems integration",
            rationale="Synthetic grounded fixture.",
        ),
        education=[
            EducationRecord(
                school="Example Polytechnic Institute",
                program="Bachelor of Applied Engineering, Systems Option",
                start_date="Sep 2021",
                expected_graduation_date="Apr 2026",
                location="Example City, ZZ",
                gpa="3.8/4.0",
                awards=["Example Merit Award", "Design Showcase Finalist"],
                relevant_coursework=["Control Systems", "Embedded Computing"],
            )
        ],
        technical_skills=[
            TechnicalSkillCategory(
                category="Future Integration Toolchain",
                values=["Python", "ROS 2", "Docker", "A Very Long Verified Tool Name"],
            ),
            TechnicalSkillCategory(
                category="Prototype Hardware",
                values=["STM32", "SPI"],
            ),
        ],
        experiences=[experience],
        projects=[project],
        entity_titles={experience.id: experience.title, project.id: project.title},
        experience_bullets={
            experience.id: [
                StructuredBullet(
                    id="experience-evidence",
                    text=(
                        "Integrated verified embedded controls and sensor interfaces for a "
                        "reusable prototype platform."
                    ),
                    evidence_ids=["experience-evidence"],
                    support=ClaimSupport.DIRECT,
                )
            ]
        },
        project_bullets={
            project.id: [
                StructuredBullet(
                    id="project-evidence",
                    text="Built a verified telemetry workflow using Python and ROS 2.",
                    evidence_ids=["project-evidence"],
                    support=ClaimSupport.DIRECT,
                )
            ]
        },
    )


def _paragraph(document: Document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def _transition(
    profile,
    source_role: str,
    destination_role: str,
    destination_section_first_role: str | None = None,
):
    matches = [
        item
        for item in profile.transition_spacings
        if item.source_role == source_role
        and item.destination_role == destination_role
        and (
            destination_role != "section_heading"
            or item.destination_section_first_role == destination_section_first_role
        )
    ]
    assert matches, (source_role, destination_role, destination_section_first_role)
    return max(matches, key=lambda item: item.occurrence_count)


def _resolved_before(transition) -> int | None:
    observed = (
        transition.resolved_destination_space_before_twips
        or transition.destination_space_before_twips
    )
    return observed.value if isinstance(observed.value, int) else None


def _resolved_after(transition) -> int | None:
    observed = transition.resolved_source_space_after_twips or transition.source_space_after_twips
    return observed.value if isinstance(observed.value, int) else None


def _with_second_experience(resume: StructuredResume) -> StructuredResume:
    first = resume.experiences[0]
    second = ResumeItem(
        id="experience-second",
        title="Verification Systems Engineer",
        kind=EntityKind.EXPERIENCE,
        organization="Example Verification Lab",
        start_date="Jan 2023",
        end_date="Apr 2024",
        location="Example City, ZZ",
    )
    bullets = {
        **resume.experience_bullets,
        second.id: [
            StructuredBullet(
                id="second-bullet-one",
                text="Validated a verified integration workflow.",
                evidence_ids=["second-bullet-one"],
                support=ClaimSupport.DIRECT,
            ),
            StructuredBullet(
                id="second-bullet-two",
                text="Documented a verified test result.",
                evidence_ids=["second-bullet-two"],
                support=ClaimSupport.DIRECT,
            ),
        ],
    }
    return resume.model_copy(
        update={
            "experiences": [first, second],
            "experience_bullets": bullets,
            "entity_titles": {
                **resume.entity_titles,
                second.id: second.title,
            },
        }
    )


def test_adaptive_renderer_maps_complete_structured_resume_without_mutation(tmp_path: Path) -> None:
    before_reference = REFERENCE.read_bytes()
    profile = analyze_reference_docx(REFERENCE)
    resume = _resume()
    before_resume = resume.model_dump(mode="json")
    output = tmp_path / "adaptive.docx"

    render_structured_resume(resume, profile, output)

    assert output.exists()
    assert REFERENCE.read_bytes() == before_reference
    assert resume.model_dump(mode="json") == before_resume
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for expected in (
        "Alex Example",
        "Example Polytechnic Institute",
        "Sep 2021",
        "Apr 2026",
        "Bachelor of Applied Engineering, Systems Option",
        "Example City, ZZ",
        "GPA: 3.8/4.0",
        "Example Merit Award",
        "Control Systems",
        "Future Integration Toolchain:",
        "Python, ROS 2, Docker, A Very Long Verified Tool Name",
        "Prototype Hardware:",
        "Systems Integration Engineer | Embedded Controls",
        "Example Research Cooperative",
        "May 2024",
        "Aug 2025",
        "Adaptive Sensor Platform | Python, ROS 2",
        "Jan 2024",
        "Apr 2024",
    ):
        assert expected in text
    assert text.count("Integrated verified embedded controls") == 1
    assert text.count("Built a verified telemetry workflow") == 1


def test_page_header_contact_and_section_formatting_come_from_profile(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    changed_page = profile.page.model_copy(
        update={"left_margin_twips": profile.page.left_margin_twips + 137}
    )
    changed_profile = profile.model_copy(update={"page": changed_page})
    output = tmp_path / "geometry.docx"
    render_structured_resume(_resume(), changed_profile, output)
    document = Document(output)
    section = document.sections[0]

    assert section.page_width.twips == changed_page.width_twips
    assert section.page_height.twips == changed_page.height_twips
    assert section.left_margin.twips == changed_page.left_margin_twips
    assert section.right_margin.twips == changed_page.right_margin_twips
    name = document.paragraphs[0]
    contact = document.paragraphs[1]
    if profile.semantic_roles["name"].paragraph.alignment.value == "center":
        assert name.alignment == 1
    assert contact.alignment == 1
    expected_family = profile.semantic_roles["name"].primary_typography.font_family.value
    if expected_family:
        assert name.runs[0].font.name == expected_family
    expected_size = profile.semantic_roles["name"].primary_typography.font_size_half_points.value
    if isinstance(expected_size, int):
        assert name.runs[0].font.size.pt == expected_size / 2
    for label in ("Education", "Technical Skills", "Technical Experience", "Projects"):
        heading = _paragraph(document, label)
        assert heading.style.name == "Normal"
        assert (
            heading._p.pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr"
            )
            is not None
        )
        assert (
            heading.runs[0].font.color.rgb is None
            or str(heading.runs[0].font.color.rgb) != "2F5496"
        )


def test_generated_geometry_and_tabs_stay_inside_profile_usable_area(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "geometry-boundary.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    section = document.sections[0]

    assert section.page_width.twips == profile.page.width_twips
    assert section.page_height.twips == profile.page.height_twips
    assert section.left_margin.twips == profile.page.left_margin_twips
    assert section.right_margin.twips == profile.page.right_margin_twips
    assert section.top_margin.twips == profile.page.top_margin_twips
    assert section.bottom_margin.twips == profile.page.bottom_margin_twips
    assert all(
        tab.position.twips <= profile.page.usable_width_twips
        for paragraph in document.paragraphs
        for tab in paragraph.paragraph_format.tab_stops
    )
    assert all("    " not in paragraph.text for paragraph in document.paragraphs)


def test_reference_margin_changes_propagate_without_renderer_constants(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    changed_page = profile.page.model_copy(
        update={
            "left_margin_twips": profile.page.left_margin_twips + 113,
            "right_margin_twips": profile.page.right_margin_twips + 97,
            "top_margin_twips": profile.page.top_margin_twips + 41,
            "bottom_margin_twips": profile.page.bottom_margin_twips + 37,
        }
    )
    changed_profile = profile.model_copy(update={"page": changed_page})
    output = tmp_path / "changed-geometry.docx"
    render_structured_resume(_resume(), changed_profile, output)
    section = Document(output).sections[0]
    assert section.left_margin.twips == changed_page.left_margin_twips
    assert section.right_margin.twips == changed_page.right_margin_twips
    assert section.top_margin.twips == changed_page.top_margin_twips
    assert section.bottom_margin.twips == changed_page.bottom_margin_twips


def test_long_metadata_uses_line_fallback_without_truncation_or_dangling_spaces(
    tmp_path: Path,
) -> None:
    profile = analyze_reference_docx(REFERENCE)
    long_title = (
        "A Very Long Verified Systems Integration Engineering Role Title That Must Remain Intact"
    )
    long_organization = (
        "A Very Long Verified Research Organization Name With Multiple Operational Divisions"
    )
    resume = _resume().model_copy(
        update={
            "experiences": [
                _resume()
                .experiences[0]
                .model_copy(
                    update={
                        "title": long_title,
                        "organization": long_organization,
                        "location": "A Very Long Verified City, Region, Country Location",
                        "start_date": "January 2024",
                        "end_date": "December 2026",
                    }
                )
            ]
        }
    )
    output = tmp_path / "long-metadata.docx"
    render_structured_resume(resume, profile, output)
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert long_title in text
    assert long_organization in text
    assert "A Very Long Verified City, Region, Country Location" in text
    assert "January 2024" in text and "December 2026" in text
    assert all("    " not in paragraph.text for paragraph in document.paragraphs)
    assert any("\n" in paragraph.text for paragraph in document.paragraphs)


def test_contact_hyperlinks_are_compact_and_missing_fields_leave_no_separators(
    tmp_path: Path,
) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "contact.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    contact = document.paragraphs[1]
    assert contact.text.count("|") == 3
    with ZipFile(output) as package:
        relationships = package.read("word/_rels/document.xml.rels").decode()
        assert "mailto:alex@example.test" in relationships
        assert "https://example.test/portfolio" in relationships

    missing = _resume().model_copy(update={"contact_line": "alex@example.test"})
    missing_output = tmp_path / "missing-contact.docx"
    render_structured_resume(missing, profile, missing_output)
    missing_contact = Document(missing_output).paragraphs[1]
    assert missing_contact.text == "alex@example.test"
    assert "|" not in missing_contact.text


def test_metadata_uses_real_tabs_and_skills_remain_separate_categories(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "tabs-skills.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    skill_rows = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(("Future Integration Toolchain:", "Prototype Hardware:"))
    ]
    assert len(skill_rows) == 2
    assert skill_rows[0].runs[0].bold is True
    assert skill_rows[0].runs[1].bold is False
    assert skill_rows[0].paragraph_format.left_indent is not None
    with ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode()
        assert "<w:tabs>" in document_xml
        assert "<w:tab/>" in document_xml or "<w:tab " in document_xml
        assert "May 2024    Aug 2025" not in document_xml


def test_categorized_skills_override_conflicting_legacy_flat_field(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    resume = _resume().model_copy(update={"selected_skills": ["Legacy Flat Value"]})
    output = tmp_path / "categorized-authority.docx"

    render_structured_resume(resume, profile, output)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "Future Integration Toolchain: Python, ROS 2, Docker" in text
    assert "Prototype Hardware: STM32, SPI" in text
    assert "Legacy Flat Value" not in text
    assert "Skills:" not in text


def test_complete_metadata_and_project_section_are_consumed_from_structured_resume(
    tmp_path: Path,
) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "complete-consumption.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for value in (
        "Sep 2021",
        "Apr 2026",
        "Example City, ZZ",
        "Example Merit Award",
        "Control Systems",
        "May 2024",
        "Aug 2025",
        "Example Research Cooperative",
        "Adaptive Sensor Platform",
        "Projects",
    ):
        assert value in text
    with ZipFile(output) as package:
        xml = package.read("word/document.xml").decode()
        assert xml.count("<w:tab/>") >= 4


def test_empty_projects_do_not_create_an_empty_section(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    resume = _resume().model_copy(update={"projects": [], "project_bullets": {}})
    output = tmp_path / "no-projects.docx"
    render_structured_resume(resume, profile, output)
    assert "Projects" not in [paragraph.text for paragraph in Document(output).paragraphs]


def test_manual_generation_uses_current_plan_and_structured_resume_flow() -> None:
    module = _load_manual_render_module()
    master_profile, plan, resume = module.build_manual_resume()
    diagnostic = module.field_presence_diagnostic(master_profile, plan, resume)

    assert resume.education == plan.education
    assert resume.technical_skills == plan.technical_skills
    assert resume.experiences == plan.selected_experiences
    assert resume.projects == plan.selected_projects
    assert diagnostic["structured_resume"]["technical_skill_category_count"] == len(
        resume.technical_skills
    )
    assert "contact" not in str(diagnostic).casefold()
    assert "bullet_text" not in str(diagnostic).casefold()


def test_bullets_use_profile_marker_indentation_and_role_spacing(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "bullets.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    experience_bullet = next(
        paragraph
        for paragraph in document.paragraphs
        if "Integrated verified embedded controls" in paragraph.text
    )
    bullet_role = profile.semantic_roles["experience_bullet"]
    assert (
        experience_bullet.paragraph_format.left_indent.twips == bullet_role.bullet.left_indent_twips
    )
    assert (
        experience_bullet.paragraph_format.first_line_indent.twips
        == -bullet_role.bullet.hanging_indent_twips
    )
    assert experience_bullet._p.pPr.numPr is not None
    assert len(experience_bullet.runs) == 1
    assert experience_bullet.runs[0].bold is False
    title = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Systems Integration Engineer")
    )
    employer = _paragraph(document, "Example Research Cooperative\tExample City, ZZ")
    assert title.paragraph_format.space_before != experience_bullet.paragraph_format.space_before
    assert employer.paragraph_format.space_before != title.paragraph_format.space_before


def test_reference_numbering_marker_font_is_separate_from_plain_bullet_text(
    tmp_path: Path,
) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "numbered-bullets.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    bullets = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    ]

    assert bullets
    assert all("\uf0a7" not in paragraph.text for paragraph in bullets)
    assert all("\u25a1" not in paragraph.text for paragraph in bullets)
    assert all(len(paragraph.runs) == 1 for paragraph in bullets)
    assert all(paragraph.runs[0].font.name == "Times New Roman" for paragraph in bullets)
    assert all(paragraph.runs[0].bold is False for paragraph in bullets)
    with ZipFile(output) as package:
        numbering_xml = package.read("word/numbering.xml").decode("utf-8")
        document_xml = package.read("word/document.xml").decode("utf-8")
        relationships = package.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "Wingdings" in numbering_xml
        assert '<w:numFmt w:val="bullet"' in numbering_xml
        assert "<w:numPr>" in document_xml
        assert "relationships/numbering" in relationships
        assert "\uf0a7" not in document_xml


def test_awards_and_gpa_share_a_compact_row_and_coursework_stays_dedicated(
    tmp_path: Path,
) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "education-details.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    details = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(("Awards:", "Relevant Courses:"))
    ]

    assert [paragraph.text.split(":", 1)[0] for paragraph in details] == [
        "Awards",
        "Relevant Courses",
    ]
    assert "GPA: 3.8/4.0" in details[0].text
    assert all(paragraph._p.pPr.numPr is not None for paragraph in details)
    assert all(len(paragraph.runs) == 1 for paragraph in details)
    assert all(paragraph.runs[0].bold is False for paragraph in details)
    expected = profile.semantic_roles["education_detail_bullet"].bullet
    assert expected is not None
    assert all(
        paragraph.paragraph_format.left_indent.twips == expected.left_indent_twips
        for paragraph in details
    )
    assert all(
        paragraph.paragraph_format.first_line_indent.twips == -expected.hanging_indent_twips
        for paragraph in details
    )


def test_metadata_rows_share_profile_derived_canonical_tab_column(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "metadata-columns.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    metadata_rows = [
        paragraph
        for paragraph in document.paragraphs
        if "\t" in paragraph.text
        and paragraph.text.startswith(
            (
                "Example Polytechnic Institute",
                "Bachelor of Applied Engineering",
                "Systems Integration Engineer",
                "Example Research Cooperative",
                "Adaptive Sensor Platform",
            )
        )
    ]
    positions = []
    for paragraph in metadata_rows:
        stops = list(paragraph.paragraph_format.tab_stops)
        assert len(stops) == 1
        positions.append(stops[0].position.twips)
        assert "    " not in paragraph.text
    secondary = max(
        group.representative_position_twips
        for group in profile.metadata_anchor_groups
        if {
            "education_institution_date_row",
            "education_program_location_row",
            "employer_location_row",
        }
        & set(group.role_groups)
    )
    primary = max(
        group.representative_position_twips
        for group in profile.metadata_anchor_groups
        if {
            "experience_title_date_row",
            "interior_entry_transition",
        }
        & set(group.role_groups)
    )
    actual = {
        next(
            prefix
            for prefix in (
                "Example Polytechnic Institute",
                "Bachelor of Applied Engineering",
                "Systems Integration Engineer",
                "Example Research Cooperative",
                "Adaptive Sensor Platform",
            )
            if paragraph.text.startswith(prefix)
        ): tab.position.twips
        for paragraph in metadata_rows
        for tab in paragraph.paragraph_format.tab_stops
    }
    assert actual == {
        "Example Polytechnic Institute": secondary,
        "Bachelor of Applied Engineering": secondary,
        "Systems Integration Engineer": primary,
        "Example Research Cooperative": secondary,
        "Adaptive Sensor Platform": primary,
    }


def test_semantic_transition_spacing_and_no_blank_paragraphs(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "semantic-spacing.docx"
    render_structured_resume(_resume(), profile, output)
    document = Document(output)
    title = next(
        p for p in document.paragraphs if p.text.startswith("Systems Integration Engineer")
    )
    employer = next(
        p for p in document.paragraphs if p.text.startswith("Example Research Cooperative")
    )
    bullet = next(p for p in document.paragraphs if p.text.startswith("Integrated verified"))
    title_to_employer = max(
        (
            item
            for item in profile.transition_spacings
            if item.source_role == "experience_title_date_row"
            and item.destination_role == "employer_location_row"
        ),
        key=lambda item: item.occurrence_count,
    )
    employer_to_bullet = max(
        (
            item
            for item in profile.transition_spacings
            if item.source_role == "employer_location_row"
            and item.destination_role == "experience_bullet"
        ),
        key=lambda item: item.occurrence_count,
    )
    assert (
        employer.paragraph_format.space_before.twips
        == title_to_employer.destination_space_before_twips.value
    )
    assert (
        bullet.paragraph_format.space_before.twips
        == employer_to_bullet.destination_space_before_twips.value
    )
    assert title.paragraph_format.space_before != employer.paragraph_format.space_before
    assert all(paragraph.text for paragraph in document.paragraphs)


def test_reference_derived_transition_resolution_covers_section_and_detail_rhythm() -> None:
    profile = analyze_reference_docx(REFERENCE)
    section_to_education = _transition(profile, "section_heading", "education_institution_date_row")
    section_to_skills = _transition(profile, "section_heading", "skill_category_row")
    final_education_to_skills = _transition(
        profile,
        "final_paragraph_in_section",
        "section_heading",
        "skill_category_row",
    )
    skill_to_experience = _transition(
        profile,
        "skill_category_row",
        "section_heading",
        "experience_title_date_row",
    )
    assert (
        _resolved_before(section_to_education)
        == section_to_education.destination_space_before_twips.value
    )
    assert (
        _resolved_before(section_to_skills)
        == section_to_skills.destination_space_before_twips.value
    )
    assert _resolved_before(final_education_to_skills) is None
    assert (
        _resolved_before(skill_to_experience)
        == skill_to_experience.destination_space_before_twips.value
    )
    assert section_to_education.empty_paragraph_count == 1
    assert section_to_education.drawing_separator_present is True


def test_semantic_spacing_is_compact_and_transition_specific(tmp_path: Path) -> None:
    profile = analyze_reference_docx(REFERENCE)
    output = tmp_path / "semantic-rhythm.docx"
    render_structured_resume(_with_second_experience(_resume()), profile, output)
    document = Document(output)
    paragraphs = document.paragraphs

    education_details = [
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith(("Awards:", "Relevant Courses:"))
    ]
    skills = [
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith(("Future Integration Toolchain:", "Prototype Hardware:"))
    ]
    title = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("Systems Integration Engineer")
    )
    employer = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("Example Research Cooperative")
    )
    first_bullet = next(
        paragraph for paragraph in paragraphs if paragraph.text.startswith("Integrated verified")
    )
    second_title = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("Verification Systems Engineer")
    )
    second_bullets = [
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith(("Validated a verified", "Documented a verified"))
    ]
    technical_skills = _paragraph(document, "Technical Skills")
    technical_experience = _paragraph(document, "Technical Experience")

    assert education_details[0].paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "education_program_location_row", "education_detail_bullet")
    )
    assert education_details[1].paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "education_detail_bullet", "final_paragraph_in_section")
    )
    assert (
        technical_skills.paragraph_format.space_before.twips
        == profile.semantic_roles["section_heading"].paragraph.space_before_twips.value
    )
    assert technical_experience.paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "skill_category_row", "section_heading", "experience_title_date_row")
    )
    assert skills[0].paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "section_heading", "skill_category_row")
    )
    assert title.paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "section_heading", "experience_title_date_row")
    )
    assert employer.paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "experience_title_date_row", "employer_location_row")
    )
    assert first_bullet.paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "employer_location_row", "experience_bullet")
    )
    assert second_title.paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "experience_bullet", "interior_entry_transition")
    )
    assert second_bullets[0].paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "employer_location_row", "experience_bullet")
    )
    assert second_bullets[1].paragraph_format.space_before.twips == _resolved_before(
        _transition(profile, "experience_bullet", "experience_bullet")
    )
    assert _resolved_before(
        _transition(profile, "experience_bullet", "interior_entry_transition")
    ) != _resolved_before(_transition(profile, "experience_bullet", "experience_bullet"))
    assert all(paragraph.text.strip() for paragraph in paragraphs)


def test_spacing_render_is_deterministic_without_content_or_reference_mutation(
    tmp_path: Path,
) -> None:
    before_reference = REFERENCE.read_bytes()
    profile = analyze_reference_docx(REFERENCE)
    resume = _with_second_experience(_resume())
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    render_structured_resume(resume, profile, first)
    render_structured_resume(resume, profile, second)

    def snapshot(path: Path) -> tuple[list[str], list[tuple[int | None, int | None]]]:
        document = Document(path)
        spacing = []
        for paragraph in document.paragraphs:
            before = paragraph.paragraph_format.space_before
            after = paragraph.paragraph_format.space_after
            spacing.append(
                (
                    before.twips if before is not None else None,
                    after.twips if after is not None else None,
                )
            )
        return [paragraph.text for paragraph in document.paragraphs], spacing

    assert snapshot(first) == snapshot(second)
    assert snapshot(first)[0] == snapshot(second)[0]
    assert REFERENCE.read_bytes() == before_reference
