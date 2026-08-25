from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from pathlib import Path

from docx import Document
from streamlit.testing.v1 import AppTest

import resume_tailor.infrastructure.dependencies as dependencies
from resume_tailor.application.generated_artifact import content_fingerprint
from resume_tailor.application.job_intake import build_job_posting
from resume_tailor.application.workflow_state import (
    COVER_LETTER_ARTIFACT_KEY,
    GENERATED_RESUME_REVIEW_STATE_KEY,
    GeneratedResumeReviewState,
)
from resume_tailor.domain.company_research import (
    CompanyFactConfidence,
    CompanyResearchStatus,
    CompanySourceType,
)
from resume_tailor.domain.cover_letter import (
    CoverLetterQualityGateStatus,
    CoverLetterRecipient,
    CoverLetterReviewState,
    CoverLetterValidationStatus,
)
from resume_tailor.domain.models import (
    JobPosting,
    MasterProfile,
    StructuredResume,
    TemplateConstraints,
)
from resume_tailor.domain.requirement_ranking import EvidenceRelationship
from resume_tailor.domain.resume_composition import ProjectRepresentationStatus
from resume_tailor.infrastructure.config import Settings
from resume_tailor.infrastructure.rendering import PageCountMeasurement
from tests.cover_letter_helpers import ControlledCoverLetterRenderer

ROOT = Path(__file__).resolve().parents[1]
PROFILE_FIXTURE = ROOT / "tests" / "fixtures" / "world_star_tech_production_profile.json"
POSTING_FIXTURE = (
    ROOT / "tests" / "fixtures" / "titan_haptics_mechatronics_integration_engineer.txt"
)


class _CountingPageProvider:
    def __init__(self, maximum_paragraphs: int = 40) -> None:
        self.maximum_paragraphs = maximum_paragraphs
        self.measure_calls = 0
        self.batch_calls = 0

    def measure(self, docx_path: Path) -> PageCountMeasurement:
        self.measure_calls += 1
        paragraph_count = sum(
            bool(paragraph.text.strip()) for paragraph in Document(docx_path).paragraphs
        )
        page_count = 1 if paragraph_count <= self.maximum_paragraphs else 2
        return PageCountMeasurement(
            page_count=page_count,
            provider=f"controlled paragraph limit {self.maximum_paragraphs}",
            confidence="exact",
            exact=True,
        )

    def measure_many(self, docx_paths: list[Path]) -> list[PageCountMeasurement]:
        self.batch_calls += 1
        return [self.measure(path) for path in docx_paths]


class _CountingResearchFetcher:
    def __init__(self) -> None:
        self.calls = 0

    def fetch(self, *_args: object, **_kwargs: object) -> object:
        self.calls += 1
        raise AssertionError("Blank optional company inputs must not make a network request")


@dataclass
class _CountingAdapters:
    page_provider: _CountingPageProvider
    research_fetcher: _CountingResearchFetcher
    cover_renderer: ControlledCoverLetterRenderer
    provider_constructions: list[str]


def _install_counting_adapters(monkeypatch) -> _CountingAdapters:
    page_provider = _CountingPageProvider(40)
    research_fetcher = _CountingResearchFetcher()
    cover_renderer = ControlledCoverLetterRenderer([0.94])
    provider_constructions: list[str] = []

    def forbidden_provider(*_args: object, **_kwargs: object) -> object:
        provider_constructions.append("gemini")
        raise AssertionError("Gemini must not be constructed in the offline production path")

    monkeypatch.setattr(
        dependencies,
        "ExactDocxPageCountProvider",
        lambda **_kwargs: page_provider,
    )
    monkeypatch.setattr(
        dependencies,
        "HttpxOfficialCompanySourceFetcher",
        lambda **_kwargs: research_fetcher,
    )
    monkeypatch.setattr(
        dependencies,
        "CoverLetterRenderer",
        lambda **_kwargs: cover_renderer,
    )
    monkeypatch.setattr(dependencies, "GeminiResumeLanguageModel", forbidden_provider)
    return _CountingAdapters(
        page_provider=page_provider,
        research_fetcher=research_fetcher,
        cover_renderer=cover_renderer,
        provider_constructions=provider_constructions,
    )


def _settings(data_directory: Path) -> Settings:
    return Settings(
        app_data_directory=data_directory,
        gemini_api_key=None,
        gemini_model=None,
        llm_enable_role_classification=False,
        llm_enable_opportunity_analysis=False,
        llm_enable_composition=False,
        llm_enable_bullet_rewrite=False,
        llm_enable_shortening=False,
        llm_enable_cover_letter=False,
        llm_deterministic_fallback=True,
    )


def _configure_app_environment(monkeypatch, data_directory: Path) -> None:
    monkeypatch.setenv("APPLICATION_VIEGO_DATA_DIR", str(data_directory))
    monkeypatch.setenv("GEMINI_API_KEY", "")
    monkeypatch.setenv("GEMINI_MODEL", "")
    for name in (
        "LLM_ENABLE_ROLE_CLASSIFICATION",
        "LLM_ENABLE_OPPORTUNITY_ANALYSIS",
        "LLM_ENABLE_COMPOSITION",
        "LLM_ENABLE_BULLET_REWRITE",
        "LLM_ENABLE_SHORTENING",
        "LLM_ENABLE_COVER_LETTER",
    ):
        monkeypatch.setenv(name, "false")


def _fixture_profile() -> MasterProfile:
    return MasterProfile.model_validate_json(PROFILE_FIXTURE.read_text(encoding="utf-8"))


def _load_profile_through_production_repository(
    settings: Settings,
    profile: MasterProfile | None = None,
) -> MasterProfile:
    repository = dependencies.create_profile_repository(settings)
    if profile is not None:
        repository.save(profile)
    loaded = repository.get((profile or _fixture_profile()).id)
    assert loaded is not None
    return loaded


def _posting(*, description: str | None = None) -> JobPosting:
    return build_job_posting(
        "local-posting",
        "Mechatronics Integration Engineer",
        description or POSTING_FIXTURE.read_text(encoding="utf-8"),
        company_name="",
    )


def _production_artifact(monkeypatch, tmp_path, *, profile: MasterProfile | None = None):
    adapters = _install_counting_adapters(monkeypatch)
    settings = _settings(tmp_path)
    loaded = _load_profile_through_production_repository(
        settings,
        profile or _fixture_profile(),
    )
    posting = _posting()
    service = dependencies.create_tailor_service(settings)
    service.start_generation()
    plan = service.create_plan(loaded, posting, TemplateConstraints())
    artifact = service.build_generated_artifact(plan, loaded, set())
    trace = service.resume_decision_trace(
        artifact,
        loaded,
        posting,
        plan,
        profile_source="application_profile_repository",
        posting_source="streamlit_job_intake",
    )
    return adapters, service, loaded, posting, plan, artifact, trace


def _selected_text(resume: StructuredResume) -> str:
    return "\n".join(
        bullet.text
        for section in (resume.experience_bullets, resume.project_bullets)
        for bullets in section.values()
        for bullet in bullets
    )


def test_titan_mechatronics_portfolio_prefers_direct_hardware_depth(
    monkeypatch,
    tmp_path,
) -> None:
    adapters, _service, _profile, _posting_value, _plan, artifact, trace = (
        _production_artifact(monkeypatch, tmp_path)
    )
    resume = artifact.final_resume
    diagnostic = resume.composition_diagnostic
    assert diagnostic is not None
    selected_text = _selected_text(resume)
    lowered = selected_text.casefold()

    assert "drive-by-wire" in lowered
    assert "sensor architecture" in lowered
    assert "wiring harness" in lowered
    substantive_mechanical_projects = [
        entry_id
        for entry_id, bullets in resume.project_bullets.items()
        if len(bullets) >= 2
        and any(
            term in " ".join(item.text for item in bullets).casefold()
            for term in ("solidworks", "mechanical", "actuator", "linkage", "cad assembly")
        )
    ]
    assert substantive_mechanical_projects
    substantive_circuit_projects = [
        bullets
        for bullets in resume.project_bullets.values()
        if len(bullets) >= 2
        and "ldr sensor" in " ".join(item.text for item in bullets).casefold()
        and "motor driver" in " ".join(item.text for item in bullets).casefold()
    ]
    assert substantive_circuit_projects
    assert diagnostic.project_representation is not None
    assert (
        diagnostic.project_representation.status
        is ProjectRepresentationStatus.SUBSTANTIVE_PROJECT
    )
    assert all(
        count >= 2
        for entry_id, count in diagnostic.bullet_counts.items()
        if entry_id in diagnostic.selected_experience_ids
    )
    assert all(count <= 6 for count in diagnostic.bullet_counts.values())
    selected_project_diagnostics = [
        item
        for item in diagnostic.selected_candidates
        if item.entry_id in diagnostic.selected_project_ids
        and item.evidence_relationship is EvidenceRelationship.DIRECT
    ]
    assert len(selected_project_diagnostics) >= 3

    unsupported = (
        "altium",
        "kicad",
        "oscilloscope",
        "formal bom ownership",
        "haptic integration",
        "customer workshop",
        "cnc machining",
        "formal durability testing",
    )
    assert not any(term in lowered for term in unsupported)
    skill_text = " ".join(resume.selected_skills).casefold()
    assert all(
        term in skill_text
        for term in ("solidworks", "fusion360", "stm32", "c++", "wiring", "3d print")
    )
    assert 0.80 <= diagnostic.final_utilization_ratio <= 0.95
    assert diagnostic.page_count == 1
    assert diagnostic.verification_status.value == "exact"
    assert all(
        "Employer identity was not scored" in item.final_reason
        for item in diagnostic.experience_package_selections
    )
    assert trace.selected_bullet_ids == diagnostic.selected_bullet_ids
    assert trace.page_fit_finalists
    assert [item for item in trace.page_fit_finalists if item.selected]
    assert trace.provider_call_count == 0
    assert not adapters.provider_constructions
    assert "exp-exl" not in diagnostic.selected_experience_ids
    assert "exp-stush" not in diagnostic.selected_experience_ids
    assert "proj-crest" not in diagnostic.selected_project_ids
    assert {"proj-robotic-arm", "proj-ventilation"} <= set(
        diagnostic.selected_project_ids
    )


def test_titan_portfolio_is_invariant_to_employer_identity(
    monkeypatch,
    tmp_path,
) -> None:
    _adapters, service, profile, posting, _plan, baseline_artifact, _trace = (
        _production_artifact(monkeypatch, tmp_path)
    )
    renamed = profile.model_copy(
        update={
            "experiences": [
                entry.model_copy(update={"organization": f"Organization {index}"})
                for index, entry in enumerate(profile.experiences, start=1)
            ]
        }
    )
    renamed = _load_profile_through_production_repository(_settings(tmp_path), renamed)
    service.start_generation()
    changed_plan = service.create_plan(renamed, posting, TemplateConstraints())
    changed_artifact = service.build_generated_artifact(changed_plan, renamed, set())

    baseline = baseline_artifact.final_resume.composition_diagnostic
    changed = changed_artifact.final_resume.composition_diagnostic
    assert changed is not None
    assert baseline is not None
    assert changed.selected_bullet_ids == baseline.selected_bullet_ids
    assert changed.bullet_counts == baseline.bullet_counts


def test_streamlit_titan_production_path_preserves_posting_authority_and_finalist(
    monkeypatch,
    tmp_path,
) -> None:
    _configure_app_environment(monkeypatch, tmp_path)
    adapters = _install_counting_adapters(monkeypatch)
    settings = Settings()
    profile = _fixture_profile()
    dependencies.create_profile_repository(settings).save(profile)

    app_path = ROOT / "src" / "resume_tailor" / "frontend" / "app.py"
    app = AppTest.from_file(str(app_path))
    app.session_state["profile"] = profile
    app.session_state["profile_id"] = profile.id
    app.run()
    for key in ("posting", "plan", "generated_resume_artifact", COVER_LETTER_ARTIFACT_KEY):
        assert key not in app.session_state
    loaded_profile = app.session_state["profile"]
    assert loaded_profile.id == profile.id
    assert content_fingerprint(loaded_profile) == content_fingerprint(profile)

    app.button(key="pw-route-sidebar-resume_studio").click().run()
    app.text_input(key="_resume_studio_job_title_widget").input(
        "Mechatronics Integration Engineer"
    )
    app.text_area(key="_resume_studio_job_description_widget").input(
        POSTING_FIXTURE.read_text(encoding="utf-8")
    )
    app.button(key="resume-create-strategy").click().run(timeout=60)

    posting = app.session_state["posting"]
    assert isinstance(posting, JobPosting)
    posting_fingerprint = content_fingerprint(posting)
    assert posting.id == "local-posting"
    assert posting.company_name == "TITAN Haptics"
    assert app.session_state["workflow_posting_fingerprint"] == posting_fingerprint
    assert app.session_state["plan"].posting == posting
    assert content_fingerprint(app.session_state["plan"].posting) == posting_fingerprint

    resume_artifact = app.session_state["generated_resume_artifact"]
    streamlit_trace = app.session_state["_tailor_service"].resume_decision_trace(
        resume_artifact,
        loaded_profile,
        posting,
        app.session_state["plan"],
        profile_source="application_profile_repository",
        posting_source="streamlit_job_intake",
    )
    assert resume_artifact.fingerprint_inputs.normalized_posting_fingerprint == (
        posting_fingerprint
    )
    assert resume_artifact.provider_diagnostic.call_count == 0
    assert app.session_state[GENERATED_RESUME_REVIEW_STATE_KEY] == (
        GeneratedResumeReviewState.GENERATED_AWAITING_REVIEW
    )
    assert streamlit_trace.selected_bullet_ids == (
        resume_artifact.composition_diagnostic.selected_bullet_ids
    )
    assert streamlit_trace.requirement_ids
    assert streamlit_trace.retrieval_candidates
    assert streamlit_trace.composition_candidates
    assert streamlit_trace.experience_package_candidates
    assert streamlit_trace.project_candidate_ids
    assert streamlit_trace.skill_candidate_ids
    assert sum(item.selected for item in streamlit_trace.page_fit_finalists) == 1
    assert streamlit_trace.profile_source == "application_profile_repository"
    assert streamlit_trace.profile_id == profile.id
    assert streamlit_trace.profile_version == profile.version
    assert streamlit_trace.profile_fingerprint == content_fingerprint(profile)
    assert streamlit_trace.posting_source == "streamlit_job_intake"
    assert streamlit_trace.posting_fingerprint == posting_fingerprint
    assert streamlit_trace.plan_posting_fingerprint == posting_fingerprint
    assert streamlit_trace.artifact_posting_fingerprint == posting_fingerprint
    serialized_trace = streamlit_trace.model_dump_json()
    assert profile.display_name not in serialized_trace
    assert "source_text" not in serialized_trace
    assert "contact" not in serialized_trace

    direct_profile = dependencies.create_profile_repository(settings).get(profile.id)
    assert direct_profile is not None
    direct_posting = _posting()
    direct_service = dependencies.create_tailor_service(settings)
    direct_service.start_generation()
    direct_plan = direct_service.create_plan(
        direct_profile,
        direct_posting,
        TemplateConstraints(),
    )
    direct_artifact = direct_service.build_generated_artifact(
        direct_plan,
        direct_profile,
        set(),
    )
    direct_trace = direct_service.resume_decision_trace(
        direct_artifact,
        direct_profile,
        direct_posting,
        direct_plan,
        profile_source="application_profile_repository",
        posting_source="streamlit_job_intake",
    )
    assert direct_trace == streamlit_trace
    assert direct_artifact.final_resume.composition_diagnostic.selected_bullet_ids == (
        resume_artifact.final_resume.composition_diagnostic.selected_bullet_ids
    )

    app.button(key="pw-route-sidebar-cover_letters").click().run()
    company_input = next(item for item in app.text_input if item.label == "Company")
    assert company_input.value == "TITAN Haptics"
    company_input.input("")
    assert all(
        not item.value
        for item in app.text_input
        if item.label
        in {
            "Company domain (optional)",
            "Recipient name (optional)",
            "Recipient title (optional)",
        }
    )
    assert all(
        not item.value
        for item in app.text_area
        if item.label.startswith(("Official company", "Verified company", "Your motivation"))
    )
    next(button for button in app.button if button.label == "Generate cover letter").click().run(
        timeout=60
    )

    assert not any("requires a posting" in item.value for item in app.error)
    artifact = app.session_state[COVER_LETTER_ARTIFACT_KEY]
    assert artifact.fingerprint_inputs.posting_fingerprint == posting_fingerprint
    assert artifact.company_research.status is CompanyResearchStatus.POSTING_ONLY
    posting_sources = {
        source.id
        for source in artifact.company_research.sources
        if source.source_type is CompanySourceType.JOB_POSTING
    }
    assert posting_sources
    assert all(
        source.content_fingerprint == sha256(posting.description.encode("utf-8")).hexdigest()
        for source in artifact.company_research.sources
        if source.id in posting_sources
    )
    posting_authority_ids = {
        fact.id
        for fact in artifact.company_research.facts
        if fact.source_id in posting_sources
        and fact.confidence is CompanyFactConfidence.POSTING_AUTHORITY
    }
    assert posting_authority_ids
    paragraph_authority_ids = {
        authority_id
        for paragraph in artifact.letter.paragraphs
        for authority_id in paragraph.company_research_ids
    }
    assert paragraph_authority_ids
    assert paragraph_authority_ids <= posting_authority_ids
    assert all(
        claim.status is CoverLetterValidationStatus.SUPPORTED
        and set(claim.company_research_ids) <= posting_authority_ids
        for paragraph in artifact.letter.paragraphs
        for claim in paragraph.claims
    )
    assert artifact.call_counts.provider_calls == 0
    assert artifact.call_counts.research_network_requests == 0
    assert adapters.research_fetcher.calls == 0
    assert not adapters.provider_constructions
    assert all(
        gate.status is not CoverLetterQualityGateStatus.FAILED
        for gate in artifact.quality_gates
    )
    assert artifact.review_state is CoverLetterReviewState.GENERATED_AWAITING_REVIEW
    assert artifact.ready_for_review
    assert app.session_state["posting"] == posting

    artifact_fingerprint = artifact.artifact_fingerprint
    artifact_bytes = artifact.docx_bytes
    cover_render_calls = adapters.cover_renderer.render_calls
    page_measure_calls = adapters.page_provider.measure_calls
    app.run()
    assert app.session_state["posting"] == posting
    assert app.session_state["plan"].posting == posting
    assert app.session_state[COVER_LETTER_ARTIFACT_KEY].artifact_fingerprint == (
        artifact_fingerprint
    )
    assert adapters.cover_renderer.render_calls == cover_render_calls
    assert adapters.page_provider.measure_calls == page_measure_calls
    assert adapters.research_fetcher.calls == 0

    next(
        checkbox
        for checkbox in app.checkbox
        if checkbox.label.startswith("I reviewed the complete letter")
    ).check().run()
    next(button for button in app.button if button.label == "Approve cover letter").click().run()
    approved = app.session_state[COVER_LETTER_ARTIFACT_KEY]
    assert approved.review_state is CoverLetterReviewState.APPROVED
    assert approved.artifact_fingerprint == artifact_fingerprint
    assert approved.docx_bytes == artifact_bytes
    assert adapters.cover_renderer.render_calls == cover_render_calls
    assert adapters.research_fetcher.calls == 0
    download_button = next(
        button for button in app.get("download_button") if button.label == "Download approved DOCX"
    )
    download_button.click().run()
    downloaded = app.session_state[COVER_LETTER_ARTIFACT_KEY]
    assert downloaded.review_state is CoverLetterReviewState.DOWNLOADED
    assert downloaded.artifact_fingerprint == artifact_fingerprint
    assert downloaded.docx_bytes == artifact_bytes
    assert adapters.cover_renderer.render_calls == cover_render_calls
    assert adapters.page_provider.measure_calls == page_measure_calls
    assert adapters.research_fetcher.calls == 0
    assert not adapters.provider_constructions

    app.button(key="pw-route-sidebar-resume_studio").click().run()
    app.session_state["resume_studio_pending_stage"] = "Job context"
    app.run()
    changed_description = POSTING_FIXTURE.read_text(encoding="utf-8") + (
        "\n- Validate an additional hardware interface."
    )
    app.text_area(key="_resume_studio_job_description_widget").input(changed_description)
    app.button(key="resume-create-strategy").click().run(timeout=30)

    changed_posting = app.session_state["posting"]
    assert isinstance(changed_posting, JobPosting)
    assert content_fingerprint(changed_posting) != posting_fingerprint
    assert COVER_LETTER_ARTIFACT_KEY not in app.session_state
    assert app.session_state["generated_resume_artifact"].artifact_fingerprint != (
        resume_artifact.artifact_fingerprint
    )
    assert "resume_decision_trace" not in app.session_state
    assert not app.session_state["_tailor_service"].cover_letter_artifact_is_current(
        artifact,
        profile,
        changed_posting,
        app.session_state["plan"],
        recipient=CoverLetterRecipient(company=changed_posting.company_name),
        final_resume=None,
        research_request=app.session_state[
            "_tailor_service"
        ].default_cover_letter_research_request(changed_posting),
        explicit_motivation=None,
    )
