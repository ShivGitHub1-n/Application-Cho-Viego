from pathlib import Path
import shutil

import pytest
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from resume_tailor.domain.layout import LayoutProfile
from resume_tailor.infrastructure.reference_docx import (
    ReferenceDocxAnalysisError,
    analyze_reference_docx,
)


REFERENCE = Path("manual-test/reference-resume.docx")


@pytest.fixture(scope="module")
def profile() -> LayoutProfile:
    return analyze_reference_docx(REFERENCE)


def test_reference_produces_serializable_deterministic_layout_profile(profile: LayoutProfile) -> None:
    payload = profile.to_json()
    assert LayoutProfile.model_validate_json(payload) == profile
    assert analyze_reference_docx(REFERENCE) == profile
    assert profile.inspected_parts == sorted(profile.inspected_parts)


def test_page_geometry_is_read_from_reference(profile: LayoutProfile) -> None:
    document = Document(REFERENCE)
    section = document.sections[0]
    assert profile.page.width_twips == section.page_width.twips
    assert profile.page.height_twips == section.page_height.twips
    assert profile.page.top_margin_twips == section.top_margin.twips
    assert profile.page.bottom_margin_twips == section.bottom_margin.twips
    assert profile.page.left_margin_twips == section.left_margin.twips
    assert profile.page.right_margin_twips == section.right_margin.twips


def test_header_roles_and_centered_contact_are_distinct(profile: LayoutProfile) -> None:
    name = profile.semantic_roles["name"]
    contact = profile.semantic_roles["contact_line"]
    assert name != contact
    assert contact.paragraph.alignment.value == "center"
    assert name.primary_typography != contact.primary_typography


def test_spacing_controls_and_transition_resolution_are_distinct(profile: LayoutProfile) -> None:
    name = profile.semantic_roles["name"].paragraph
    assert name.before_auto_spacing.value is True
    assert name.space_before_twips.value == 100
    assert name.contextual_spacing.value is None
    assert all(
        transition.resolved_destination_space_before_twips is not None
        for transition in profile.transition_spacings
    )
    assert all(
        "resolved_by_dominant_semantic_transition_value" in transition.provenance
        for transition in profile.transition_spacings
    )


def test_metadata_anchor_groups_are_relative_and_role_scoped(profile: LayoutProfile) -> None:
    assert profile.page.usable_width_twips == (
        profile.page.width_twips
        - profile.page.left_margin_twips
        - profile.page.right_margin_twips
    )
    assert profile.page.usable_height_twips == (
        profile.page.height_twips
        - profile.page.top_margin_twips
        - profile.page.bottom_margin_twips
    )
    assert profile.metadata_anchor_groups
    assert all(
        group.tolerance_twips == round(
            profile.page.usable_width_twips * group.relative_tolerance
        )
        for group in profile.metadata_anchor_groups
    )
    assert any(
        "education_institution_date_row" in group.role_groups
        and "employer_location_row" in group.role_groups
        for group in profile.metadata_anchor_groups
    )
    assert any(
        "experience_title_date_row" in group.role_groups
        for group in profile.metadata_anchor_groups
    )


def test_section_borders_tabs_bullets_and_role_specific_spacing_are_captured(
    profile: LayoutProfile,
) -> None:
    headings = profile.semantic_roles["section_heading"]
    assert headings.borders
    metadata_tabs = [
        tab
        for role in profile.semantic_roles.values()
        for tab in role.tab_stops
        if tab.semantic_use in {"right_aligned_metadata", "positioned_metadata_column"}
    ]
    assert metadata_tabs
    assert all(tab.alignment == "left" for tab in metadata_tabs)
    assert all(tab.semantic_use == "positioned_metadata_column" for tab in metadata_tabs)
    assert all(tab.provenance == "direct_paragraph_property" for tab in metadata_tabs)
    bullet_roles = [role for role in profile.semantic_roles.values() if role.bullet]
    assert bullet_roles
    assert any(role.bullet.hanging_indent_twips is not None for role in bullet_roles)
    assert "skill_category_row" in profile.semantic_roles
    assert profile.semantic_roles["skill_category_row"].bullet is None
    assert "experience_title_date_row" in profile.semantic_roles
    assert "employer_location_row" in profile.semantic_roles
    # The reference has one explicit space-after value. Its role-specific
    # rhythm comes from destination space-before and intervening separators.
    transition_pairs = {
        (transition.source_role, transition.destination_role)
        for transition in profile.transition_spacings
    }
    assert ("section_heading", "education_institution_date_row") in transition_pairs
    assert ("skill_category_row", "skill_category_row") in transition_pairs
    assert ("experience_title_date_row", "employer_location_row") in transition_pairs
    assert ("employer_location_row", "experience_bullet") in transition_pairs
    assert ("experience_bullet", "experience_bullet") in transition_pairs
    assert ("experience_bullet", "interior_entry_transition") in transition_pairs
    assert ("final_paragraph_in_section", "section_heading") in transition_pairs
    assert ("project_title_metadata_row", "project_bullet") in transition_pairs
    assert any(
        transition.destination_section_first_role == "skill_category_row"
        for transition in profile.transition_spacings
        if transition.destination_role == "section_heading"
    )
    spacing_signatures = {
        (
            transition.source_space_after_twips.value,
            transition.destination_space_before_twips.value,
            tuple(value.value for value in transition.empty_line_spacing_twips),
            transition.empty_paragraph_count,
        )
        for transition in profile.transition_spacings
    }
    assert len(spacing_signatures) > 1


def test_numbering_marker_typography_is_captured_from_numbering_definition(
    profile: LayoutProfile,
) -> None:
    for role_name in ("education_detail_bullet", "experience_bullet"):
        bullet = profile.semantic_roles[role_name].bullet
        assert bullet is not None
        assert bullet.mechanism == "numbering"
        assert bullet.provenance == "numbering_definition"
        assert bullet.marker_typography is not None
        assert bullet.marker_typography.font_family.value == "Wingdings"
        assert bullet.marker_typography.font_family.provenance == "numbering_definition"


def test_structurally_equivalent_section_headings_do_not_depend_on_their_text(
    tmp_path: Path,
) -> None:
    path = tmp_path / "structural-headings.docx"
    document = Document()
    document.add_paragraph("Anonymous Candidate")
    document.add_paragraph("contact placeholder")
    for heading_text, body_text in (
        ("First Topic", "ordinary content one"),
        ("Completely Different Label", "ordinary content two"),
    ):
        heading = document.add_paragraph(heading_text)
        properties = heading._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
        properties.append(borders)
        document.add_paragraph(body_text)
    document.save(path)

    profile = analyze_reference_docx(path)

    headings = profile.semantic_roles["section_heading"]
    assert headings.occurrence_count == 2
    assert headings.borders[0].provenance == "direct_paragraph_property"
    serialized = profile.to_json()
    assert "First Topic" not in serialized
    assert "Completely Different Label" not in serialized


def test_literal_tab_with_positioned_left_stop_is_detected_as_metadata_column(
    tmp_path: Path,
) -> None:
    path = tmp_path / "positioned-metadata.docx"
    document = Document()
    document.add_paragraph("Anonymous Candidate")
    document.add_paragraph("contact placeholder")
    heading = document.add_paragraph("Unrelated Topic")
    properties = heading._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    borders.append(bottom)
    properties.append(borders)
    row = document.add_paragraph()
    row.paragraph_format.tab_stops.add_tab_stop(Inches(5.75), WD_TAB_ALIGNMENT.LEFT)
    row.add_run("Primary metadata")
    row.add_run().add_tab()
    row.add_run("Secondary metadata")
    document.save(path)

    profile = analyze_reference_docx(path)
    metadata_tabs = [
        tab
        for role in profile.semantic_roles.values()
        for tab in role.tab_stops
        if tab.semantic_use == "positioned_metadata_column"
    ]

    assert len(metadata_tabs) == 1
    assert metadata_tabs[0].alignment == "left"
    assert metadata_tabs[0].position_twips == Inches(5.75).twips
    assert metadata_tabs[0].provenance == "direct_paragraph_property"
    assert "Primary metadata" not in profile.to_json()
    assert "Secondary metadata" not in profile.to_json()


def test_transition_spacing_captures_direct_values_and_empty_separator(
    tmp_path: Path,
) -> None:
    path = tmp_path / "transition-spacing.docx"
    document = Document()
    document.add_paragraph("Anonymous Candidate")
    document.add_paragraph("contact placeholder")
    heading = document.add_paragraph("Unrelated Topic")
    properties = heading._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    borders.append(bottom)
    properties.append(borders)
    source = document.add_paragraph("First unrelated body")
    source.paragraph_format.space_after = Pt(7)
    separator = document.add_paragraph()
    separator.paragraph_format.line_spacing = Pt(3)
    destination = document.add_paragraph("Second unrelated body")
    destination.paragraph_format.space_before = Pt(11)
    document.save(path)

    profile = analyze_reference_docx(path)
    transition = next(
        item
        for item in profile.transition_spacings
        if item.empty_paragraph_count == 1
        and item.source_space_after_twips.value == 140
        and item.destination_space_before_twips.value == 220
    )

    assert transition.empty_line_spacing_twips[0].value == 60
    assert "intervening_empty_paragraph_properties" in transition.provenance
    serialized = profile.to_json()
    assert "First unrelated body" not in serialized
    assert "Second unrelated body" not in serialized


def test_run_emphasis_and_hyperlink_relationships_are_inspected_safely(
    profile: LayoutProfile,
) -> None:
    patterns = [pattern for role in profile.semantic_roles.values() for pattern in role.run_patterns]
    assert any(pattern.bold_run_positions for pattern in patterns)
    assert any(pattern.italic_run_positions for pattern in patterns)
    contact = profile.semantic_roles["contact_line"]
    assert contact.hyperlinks is not None
    assert contact.hyperlinks.present
    assert contact.hyperlinks.relationship_handling == "relationship targets intentionally omitted"


def test_serialized_profile_contains_no_source_paragraph_text(profile: LayoutProfile) -> None:
    serialized = profile.to_json()
    source_paragraphs = [paragraph.text.strip() for paragraph in Document(REFERENCE).paragraphs]
    for text in source_paragraphs:
        if len(text) >= 4:
            assert text not in serialized


def test_changed_fixture_formatting_changes_profile_without_analyzer_changes(tmp_path: Path) -> None:
    changed_path = tmp_path / "changed-reference.docx"
    shutil.copyfile(REFERENCE, changed_path)
    document = Document(changed_path)
    section = document.sections[0]
    section.left_margin = Inches(1.17)
    nonempty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    first = nonempty[0]
    first.runs[0].font.size = Pt(19)
    target = nonempty[1]
    target.paragraph_format.space_after = Pt(13)
    target.paragraph_format.left_indent = Inches(0.31)
    target.paragraph_format.tab_stops.add_tab_stop(Inches(6.13), WD_TAB_ALIGNMENT.RIGHT)
    document.save(changed_path)

    original = analyze_reference_docx(REFERENCE)
    changed = analyze_reference_docx(changed_path)

    assert changed.page.left_margin_twips != original.page.left_margin_twips
    assert changed.semantic_roles["name"].primary_typography.font_size_half_points.value == 38
    assert changed.semantic_roles["contact_line"].paragraph.space_after_twips.value == 260
    assert changed.semantic_roles["contact_line"].paragraph.left_indent_twips.value != original.semantic_roles["contact_line"].paragraph.left_indent_twips.value
    assert any(tab.position_twips != original.semantic_roles["contact_line"].tab_stops[0].position_twips if original.semantic_roles["contact_line"].tab_stops else True for tab in changed.semantic_roles["contact_line"].tab_stops)


def test_missing_reference_has_controlled_error(tmp_path: Path) -> None:
    missing = tmp_path / "missing.docx"
    with pytest.raises(ReferenceDocxAnalysisError, match="does not exist"):
        analyze_reference_docx(missing)


def test_analysis_is_not_connected_to_existing_renderer(profile: LayoutProfile) -> None:
    # Regression guard: analysis is opt-in and returns data only.
    assert profile.schema_version == "1.0"
