from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

from resume_tailor.application.profile_editor import (
    editor_state_to_profile,
    profile_to_editor_state,
)
from resume_tailor.domain.layout import LayoutProfile, ObservedValue, PageUtilizationStatus
from resume_tailor.domain.models import (
    ClaimSupport,
    MasterProfile,
    ResumeStrategy,
    StructuredBullet,
    StructuredResume,
)
from resume_tailor.domain.resume_composition import (
    TEMPLATE_V1_UTILIZATION_TARGET_CEILING,
    TEMPLATE_V1_UTILIZATION_TARGET_FLOOR,
)
from resume_tailor.infrastructure.adaptive_docx import render_structured_resume
from resume_tailor.infrastructure.rendering import (
    PageCountMeasurement,
    diagnose_docx_page_utilization,
)
from resume_tailor.infrastructure.static_template_docx import render_template_v1_resume
from resume_tailor.infrastructure.template_v1 import (
    TEMPLATE_V1_ID,
    TEMPLATE_V1_REFERENCE_SHA256,
    load_template_v1_layout_profile,
)

FIXTURE = Path(__file__).parent / "fixtures" / "template_v1_complete_profile.json"
REFERENCE = Path("manual-test/reference-resume.docx")


def _profile() -> MasterProfile:
    return MasterProfile.model_validate(json.loads(FIXTURE.read_text(encoding="utf-8")))


def _resume(profile: MasterProfile | None = None) -> StructuredResume:
    profile = profile or _profile()
    evidence_by_entry: dict[str, list[StructuredBullet]] = {}
    for evidence in profile.evidence:
        evidence_by_entry.setdefault(evidence.entity_id, []).append(
            StructuredBullet(
                id=evidence.id,
                text=evidence.source_text,
                evidence_ids=[evidence.id],
                support=ClaimSupport.DIRECT,
            )
        )
    return StructuredResume(
        profile_id=profile.id,
        profile_version=profile.version,
        posting_id="template-v1-acceptance",
        template_id=TEMPLATE_V1_ID,
        display_name=profile.display_name,
        contact_line=" | ".join(
            value
            for value in (
                profile.contact.phone,
                profile.contact.email,
                *profile.contact.links,
                profile.contact.location,
            )
            if value
        ),
        strategy=ResumeStrategy(
            role_family="robotics_mechatronics",
            primary_focus="reviewed engineering evidence",
            rationale="Deterministic acceptance fixture.",
        ),
        entity_titles={item.id: item.title for item in [*profile.experiences, *profile.projects]},
        education=profile.education,
        technical_skills=profile.technical_skills,
        experiences=profile.experiences,
        projects=profile.projects,
        experience_bullets={item.id: evidence_by_entry[item.id] for item in profile.experiences},
        project_bullets={item.id: evidence_by_entry[item.id] for item in profile.projects},
        selected_coursework=profile.coursework,
    )


def test_complete_profile_editor_round_trip_preserves_template_fields() -> None:
    profile = _profile()

    round_tripped = editor_state_to_profile(profile_to_editor_state(profile))

    assert round_tripped == profile
    education = round_tripped.education[0]
    assert education.start_date == "Sept. 2024"
    assert education.expected_graduation_date == "Apr. 2029"
    assert education.location == "Toronto, Ontario, Canada"
    assert education.awards == [
        "University of Toronto Engineering International Scholar",
        "CSWA",
        "Dean's List",
    ]
    assert education.relevant_coursework == [
        "Mechanical Engineering Design",
        "Fundamentals of Computer Programming",
        "Electrical Fundamentals",
    ]
    assert [category.category for category in round_tripped.technical_skills] == [
        "Programming & Scripting",
        "Embedded Systems & Microcontrollers",
        "Robotics & Perception",
        "Data & AI",
    ]


def test_static_template_v1_contract_preserves_rows_anchors_and_all_selected_entries(
    tmp_path: Path,
) -> None:
    profile = load_template_v1_layout_profile()
    output = tmp_path / "template-v1-complete.docx"
    resume = _resume()
    resume_before = resume.model_dump(mode="json")

    render_structured_resume(resume, profile, output)

    document = Document(output)
    section = document.sections[0]
    paragraphs = document.paragraphs
    texts = [paragraph.text for paragraph in paragraphs]
    assert resume.model_dump(mode="json") == resume_before
    assert TEMPLATE_V1_REFERENCE_SHA256 == (
        "2b9dd1474b9e4a303a87b8a147f3511460988104efde7cfa053cad64294369cd"
    )
    assert (
        section.page_width.twips,
        section.page_height.twips,
        section.top_margin.twips,
        section.right_margin.twips,
        section.bottom_margin.twips,
        section.left_margin.twips,
    ) == (12_240, 15_840, 640, 360, 280, 720)
    section_order = [
        texts.index("Education"),
        texts.index("Technical Skills"),
        texts.index("Technical Experience"),
        texts.index("Projects"),
    ]
    assert section_order == sorted(section_order)
    institution = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("University of Toronto\t")
    )
    program = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("Bachelor of Applied Science")
    )
    awards = next(paragraph for paragraph in paragraphs if paragraph.text.startswith("Awards:"))
    coursework = next(
        paragraph for paragraph in paragraphs if paragraph.text.startswith("Relevant Courses:")
    )
    assert "Sept. 2024 – Expected Apr. 2029" in institution.text
    assert program.text.endswith("\tToronto, Ontario, Canada")
    assert "GPA: 3.70/4.00" in awards.text
    assert "Dean's List" in awards.text
    assert "Electrical Fundamentals" in coursework.text
    assert (
        texts.count(
            "Awards: University of Toronto Engineering International Scholar, "
            "CSWA, Dean's List, GPA: 3.70/4.00"
        )
        == 1
    )
    for label in (
        "Programming & Scripting:",
        "Embedded Systems & Microcontrollers:",
        "Robotics & Perception:",
        "Data & AI:",
    ):
        assert any(text.startswith(label) for text in texts)
    for expected in (
        "Software Engineering Intern | Python, Pandas, Power BI",
        "Stush Foods\tToronto, Ontario, Canada",
        "Robotics Systems Engineer | ROS2, OpenCV, YOLOv8",
        "Telebotics\tYork, Ontario, Canada",
        "Crest - AI-Powered Expense Intelligence Platform "
        "(3rd Place, MPC Hacks) | FastAPI, Gemini, MongoDB",
    ):
        assert any(text.startswith(expected) for text in texts)
    metadata_rows = [
        institution,
        program,
        *[
            paragraph
            for paragraph in paragraphs
            if paragraph.text.startswith(
                (
                    "Software Engineering Intern",
                    "Stush Foods",
                    "Robotics Systems Engineer",
                    "Telebotics",
                )
            )
        ],
    ]
    assert metadata_rows
    assert all(
        len(paragraph.paragraph_format.tab_stops) == 1
        and paragraph.paragraph_format.tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
        and paragraph.paragraph_format.tab_stops[0].position.twips == 11_160
        for paragraph in metadata_rows
    )
    assert profile.page.usable_width_twips == 11_160
    bullet_paragraphs = [
        paragraph
        for paragraph in paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    ]
    assert bullet_paragraphs
    education_bullets = [
        paragraph
        for paragraph in bullet_paragraphs
        if paragraph.text.startswith(("Awards:", "Relevant Courses:"))
    ]
    body_bullets = [
        paragraph for paragraph in bullet_paragraphs if paragraph not in education_bullets
    ]
    assert all(
        paragraph.paragraph_format.left_indent.twips == 525
        and paragraph.paragraph_format.first_line_indent.twips == -185
        for paragraph in education_bullets
    )
    assert all(
        paragraph.paragraph_format.left_indent.twips == 468
        and paragraph.paragraph_format.first_line_indent.twips == -185
        for paragraph in body_bullets
    )
    wrapped_bullet = max(body_bullets, key=lambda paragraph: len(paragraph.text))
    assert len(wrapped_bullet.text) > 150
    assert wrapped_bullet.paragraph_format.left_indent.twips == 468
    assert wrapped_bullet.paragraph_format.first_line_indent.twips == -185
    representative_typography = {
        "Shiv Arora": ("Georgia", 14.0),
        "Education": ("Bookman Old Style", 11.0),
        "University of Toronto": ("Times New Roman", 11.0),
        "Bachelor of Applied Science": ("Times New Roman", 10.0),
        "Programming & Scripting:": ("Times New Roman", 10.0),
        "Software Engineering Intern": ("Times New Roman", 11.0),
        "Stush Foods": ("Times New Roman", 10.0),
        "Crest - AI-Powered": ("Times New Roman", 10.0),
    }
    for prefix, expected in representative_typography.items():
        paragraph = next(item for item in paragraphs if item.text.startswith(prefix))
        assert paragraph.runs[0].font.name == expected[0]
        assert paragraph.runs[0].font.size.pt == expected[1]
    for item in [*resume.experiences, *resume.projects]:
        assert any(text.startswith(item.title) for text in texts)
    assert all(paragraph.text.strip() for paragraph in paragraphs)


def test_template_v1_spacing_is_explicit_for_every_rendered_semantic_paragraph(
    tmp_path: Path,
) -> None:
    profile = load_template_v1_layout_profile()
    output = tmp_path / "template-v1-explicit-spacing.docx"

    render_structured_resume(_resume(), profile, output)

    document = Document(output)
    paragraphs = document.paragraphs
    assert paragraphs
    expected_role_before = {
        "name": 100,
        "contact_line": 11,
        "section_heading": 246,
        "education_institution_date_row": 48,
        "education_program_location_row": 34,
        "education_detail_bullet": 53,
        "final_paragraph_in_section": 52,
        "skill_category_row": 82,
        "experience_title_date_row": 115,
        "employer_location_row": 35,
        "experience_bullet": 52,
        "interior_entry_transition": 111,
        "project_title_metadata_row": 61,
        "project_bullet": 52,
        "section_transition": 246,
        "education_awards_row": 53,
        "education_coursework_row": 53,
    }
    assert {
        role_name: role.paragraph.space_before_twips.value
        for role_name, role in profile.semantic_roles.items()
    } == expected_role_before
    assert all(
        role.paragraph.space_after_twips.value
        == (32 if role_name in {"section_heading", "section_transition"} else 0)
        for role_name, role in profile.semantic_roles.items()
    )
    assert all(
        role.paragraph.line_spacing_twips.value == 240
        and role.paragraph.line_spacing_rule.value == "exact"
        for role in profile.semantic_roles.values()
    )
    assert all(
        transition.resolved_source_space_after_twips is not None
        and isinstance(transition.resolved_source_space_after_twips.value, int)
        and transition.resolved_destination_space_before_twips is not None
        and isinstance(transition.resolved_destination_space_before_twips.value, int)
        for transition in profile.transition_spacings
    )

    for paragraph in paragraphs:
        properties = paragraph._p.pPr
        assert properties is not None
        spacing = properties.find(qn("w:spacing"))
        assert spacing is not None
        expected_after = 32 if paragraph.text == "Education" else 0
        assert spacing.get(qn("w:after")) == str(expected_after)
        assert spacing.get(qn("w:line")) == "240"
        assert spacing.get(qn("w:lineRule")) == "exact"
        assert paragraph.paragraph_format.space_after is not None
        assert paragraph.paragraph_format.space_after.twips == expected_after
        assert paragraph.paragraph_format.space_after.twips != 200
        assert paragraph.paragraph_format.line_spacing.twips == 240
        assert paragraph.paragraph_format.line_spacing.twips != 276
        assert paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY

    rendered_before = {
        "Shiv Arora": 100,
        "Education": 246,
        "University of Toronto": 48,
        "Bachelor of Applied Science": 34,
        "Awards:": 53,
        "Relevant Courses:": 52,
        "Technical Skills": 0,
        "Programming & Scripting:": 82,
        "Technical Experience": 1,
        "Software Engineering Intern": 115,
        "Stush Foods": 35,
        "Robotics Systems Engineer": 111,
        "Telebotics": 34,
        "Projects": 61,
        "Crest - AI-Powered": 61,
    }
    for prefix, expected_before in rendered_before.items():
        paragraph = next(item for item in paragraphs if item.text.startswith(prefix))
        assert paragraph.paragraph_format.space_before is not None
        assert paragraph.paragraph_format.space_before.twips == expected_before
    assert all(paragraph.text.strip() for paragraph in paragraphs)


def test_template_v1_transition_spacing_replaces_without_adding_or_erasing(
    tmp_path: Path,
) -> None:
    profile = load_template_v1_layout_profile()
    missing = ObservedValue(value=None, provenance="not_present")

    def changed_profile(
        source_after: ObservedValue,
        destination_before: ObservedValue,
    ) -> LayoutProfile:
        transitions = [
            transition.model_copy(
                update={
                    "source_space_after_twips": source_after,
                    "resolved_source_space_after_twips": source_after,
                    "destination_space_before_twips": destination_before,
                    "resolved_destination_space_before_twips": destination_before,
                }
            )
            if transition.source_role == "section_heading"
            and transition.destination_role == "education_institution_date_row"
            else transition
            for transition in profile.transition_spacings
        ]
        return profile.model_copy(update={"transition_spacings": transitions})

    missing_output = tmp_path / "template-v1-missing-transition-spacing.docx"
    render_structured_resume(
        _resume(),
        changed_profile(missing, missing),
        missing_output,
    )
    missing_document = Document(missing_output)
    education_heading = next(
        paragraph for paragraph in missing_document.paragraphs if paragraph.text == "Education"
    )
    institution = next(
        paragraph
        for paragraph in missing_document.paragraphs
        if paragraph.text.startswith("University of Toronto")
    )
    assert education_heading.paragraph_format.space_after.twips == 32
    assert institution.paragraph_format.space_before.twips == 48

    replacement_output = tmp_path / "template-v1-replacement-transition-spacing.docx"
    render_structured_resume(
        _resume(),
        changed_profile(
            ObservedValue(value=17, provenance="inferred_recurring_pattern"),
            ObservedValue(value=29, provenance="inferred_recurring_pattern"),
        ),
        replacement_output,
    )
    replacement_document = Document(replacement_output)
    education_heading = next(
        paragraph for paragraph in replacement_document.paragraphs if paragraph.text == "Education"
    )
    institution = next(
        paragraph
        for paragraph in replacement_document.paragraphs
        if paragraph.text.startswith("University of Toronto")
    )
    assert education_heading.paragraph_format.space_after.twips == 17
    assert institution.paragraph_format.space_before.twips == 29


def test_page_utilization_uses_static_template_v1_calibration_floor(
    tmp_path: Path,
) -> None:
    profile = load_template_v1_layout_profile()
    measurement = PageCountMeasurement(
        page_count=1,
        provider="acceptance-fixture",
        confidence="exact",
        exact=True,
    )
    sparse = _resume().model_copy(
        update={
            "education": [],
            "technical_skills": [],
            "experiences": [],
            "projects": [],
            "experience_bullets": {},
            "project_bullets": {},
        }
    )
    sparse_path = tmp_path / "sparse.docx"
    complete_path = tmp_path / "complete.docx"
    render_template_v1_resume(sparse, sparse_path)
    render_template_v1_resume(_resume(), complete_path)

    sparse_diagnostic = diagnose_docx_page_utilization(
        sparse_path,
        profile,
        measurement,
    )
    complete_diagnostic = diagnose_docx_page_utilization(
        complete_path,
        profile,
        measurement,
    )
    canonical_diagnostic = diagnose_docx_page_utilization(
        REFERENCE,
        profile,
        measurement,
    )
    overflow_diagnostic = diagnose_docx_page_utilization(
        complete_path,
        profile,
        PageCountMeasurement(
            page_count=2,
            provider="acceptance-fixture",
            confidence="exact",
            exact=True,
        ),
    )

    assert sparse_diagnostic.status is PageUtilizationStatus.SEVERE_UNDERFILL
    assert complete_diagnostic.status is PageUtilizationStatus.SEVERE_UNDERFILL
    assert canonical_diagnostic.status is PageUtilizationStatus.ACCEPTABLE_ONE_PAGE
    assert overflow_diagnostic.status is PageUtilizationStatus.OVERFLOW
    assert complete_diagnostic.estimated_utilization_ratio == pytest.approx(0.6669571045576408)
    assert canonical_diagnostic.estimated_utilization_ratio == pytest.approx(0.964343163538874)
    assert complete_diagnostic.severe_underfill_threshold == TEMPLATE_V1_UTILIZATION_TARGET_FLOOR
    assert (
        canonical_diagnostic.estimated_utilization_ratio <= TEMPLATE_V1_UTILIZATION_TARGET_CEILING
    )
    assert complete_diagnostic.uncontrolled_blank_paragraph_count == 0
    assert (
        complete_diagnostic.estimated_utilization_ratio
        > sparse_diagnostic.estimated_utilization_ratio
    )


def test_selected_entries_render_even_when_no_bullet_is_selected(
    tmp_path: Path,
) -> None:
    profile = load_template_v1_layout_profile()
    output = tmp_path / "selected-entry-headers.docx"
    resume = _resume().model_copy(
        update={
            "experience_bullets": {},
            "project_bullets": {},
        }
    )

    render_structured_resume(resume, profile, output)

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "Software Engineering Intern" in text
    assert "Robotics Systems Engineer" in text
    assert "Crest - AI-Powered Expense Intelligence Platform" in text
