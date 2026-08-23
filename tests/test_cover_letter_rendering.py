from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from resume_tailor.application.cover_letter import CoverLetterService
from resume_tailor.domain.cover_letter import (
    CoverLetterPageFitStatus,
    CoverLetterQualityGateStatus,
    CoverLetterReviewState,
)
from resume_tailor.infrastructure.cover_letter_rendering import (
    CoverLetterRenderer,
    audit_cover_letter_docx_structure,
)
from resume_tailor.infrastructure.rendering import (
    PageCountMeasurement,
    PageCountVerificationError,
)
from tests.cover_letter_helpers import ControlledCoverLetterRenderer, cover_letter_case


class _ExactBatchProvider:
    def __init__(self, page_count: int = 1) -> None:
        self.page_count = page_count
        self.batch_calls = 0
        self.single_calls = 0

    def measure(self, path: Path) -> PageCountMeasurement:
        del path
        self.single_calls += 1
        raise AssertionError("Cover-letter finalists must use bounded batch pagination")

    def measure_many(self, paths: list[Path]) -> list[PageCountMeasurement]:
        self.batch_calls += 1
        return [
            PageCountMeasurement(
                page_count=self.page_count,
                provider="controlled exact Word pagination",
                confidence="exact",
                exact=True,
            )
            for _path in paths
        ]


class _UnavailableProvider:
    def measure(self, path: Path) -> PageCountMeasurement:
        del path
        raise PageCountVerificationError("Microsoft Word is unavailable")

    def measure_many(self, paths: list[Path]) -> list[PageCountMeasurement]:
        del paths
        raise PageCountVerificationError("Microsoft Word is unavailable")


def _letter():
    profile, posting, plan = cover_letter_case()
    return (
        CoverLetterService(renderer=ControlledCoverLetterRenderer([0.94]))
        .generate_artifact(profile, posting, plan)
        .letter
    )


def test_renderer_creates_one_page_docx_with_fixed_professional_tokens(tmp_path: Path) -> None:
    provider = _ExactBatchProvider()
    renderer = CoverLetterRenderer(page_count_provider=provider)
    letter = _letter()

    result = renderer.render_candidate(letter, tmp_path)

    assert result.page_count == 1
    assert result.measurement.exact
    assert result.docx_bytes == result.docx_path.read_bytes()
    assert provider.batch_calls == 1
    assert provider.single_calls == 0
    document = Document(BytesIO(result.docx_bytes))
    section = document.sections[0]
    assert round(section.top_margin.inches, 2) == 1.0
    assert round(section.left_margin.inches, 2) == 1.0
    assert document.styles["Cover Letter Body"].font.size.pt == 11
    assert document.styles["Cover Letter Name"].font.size.pt == 16
    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == letter.paragraphs[0].text
    )
    assert body.runs[0].font.size.pt == 11
    assert body.paragraph_format.space_after.pt == 8
    assert body.paragraph_format.line_spacing == 1.1
    assert "explicit_page_break" not in audit_cover_letter_docx_structure(result.docx_path)


def test_renderer_detects_blank_or_structurally_suspect_second_page(tmp_path: Path) -> None:
    renderer = CoverLetterRenderer(page_count_provider=_ExactBatchProvider(page_count=2))

    result = renderer.render_candidate(_letter(), tmp_path)

    assert result.page_count == 2
    assert result.blank_trailing_page


def test_pagination_failure_falls_back_to_honest_estimate(tmp_path: Path) -> None:
    renderer = CoverLetterRenderer(page_count_provider=_UnavailableProvider())

    result = renderer.render_candidate(_letter(), tmp_path)

    assert not result.measurement.exact
    assert result.pagination_failure == "Microsoft Word is unavailable"
    assert result.measurement.provider == "deterministic cover-letter occupancy estimate"
    assert result.estimated_utilization >= 0
    assert result.estimated_remaining_lines >= 0


def test_unavailable_pagination_retains_review_copy_but_fails_final_authority() -> None:
    profile, posting, plan = cover_letter_case()
    renderer = CoverLetterRenderer(page_count_provider=_UnavailableProvider())

    artifact = CoverLetterService(renderer=renderer).generate_artifact(
        profile,
        posting,
        plan,
    )

    page_gate = next(gate for gate in artifact.quality_gates if gate.gate == "page_fit")
    assert artifact.docx_bytes.startswith(b"PK\x03\x04")
    assert artifact.page_fit.status is CoverLetterPageFitStatus.PAGINATION_UNVERIFIED
    assert not artifact.page_fit.exact_pagination
    assert artifact.page_fit.manual_word_inspection_required
    assert page_gate.status is CoverLetterQualityGateStatus.FAILED
    assert not artifact.ready_for_review
    assert artifact.review_state is CoverLetterReviewState.GENERATION_FAILED


def test_contact_header_has_no_duplicate_raw_values(tmp_path: Path) -> None:
    result = CoverLetterRenderer(page_count_provider=_ExactBatchProvider()).render_candidate(
        _letter(), tmp_path
    )
    document = Document(BytesIO(result.docx_bytes))
    header_text = "\n".join(paragraph.text for paragraph in document.paragraphs[:3])

    assert header_text.count("candidate@example.com") == 1
    assert header_text.count("Toronto, ON") == 1
    assert "GitHub" in header_text
    assert "Website" not in header_text
    assert any(
        relationship.reltype.endswith("/hyperlink")
        and relationship.target_ref == "https://github.com/candidate"
        for relationship in document.part.rels.values()
    )


def test_signoff_stays_attached_without_hiding_or_negative_spacing(tmp_path: Path) -> None:
    result = CoverLetterRenderer(page_count_provider=_ExactBatchProvider()).render_candidate(
        _letter(), tmp_path
    )
    document = Document(BytesIO(result.docx_bytes))
    signoff_index, signoff = next(
        (index, paragraph)
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "Sincerely,"
    )
    signoff_name = document.paragraphs[signoff_index + 1]

    assert signoff.paragraph_format.keep_with_next
    assert signoff.paragraph_format.space_after.pt >= 0
    assert signoff_name.paragraph_format.space_before.pt == 0
    assert signoff_name.text == "Candidate Name"
    assert signoff_name._p.pPr.find(qn("w:vanish")) is None
