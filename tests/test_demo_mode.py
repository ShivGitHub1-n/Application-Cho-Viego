from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from resume_tailor.application import demo_mode
from resume_tailor.application.cover_letter import CoverLetterService
from resume_tailor.application.generated_artifact import content_fingerprint
from resume_tailor.config import Settings
from resume_tailor.domain.company_research import CompanyResearchRequest
from resume_tailor.domain.cover_letter import (
    CoverLetterProviderStatus,
    CoverLetterReviewState,
)
from resume_tailor.domain.models import (
    ContactInfo,
    EducationRecord,
    EntityKind,
    EvidenceItem,
    JobPosting,
    MasterProfile,
    ResumeItem,
    TechnicalSkillCategory,
    TemplateConstraints,
)
from resume_tailor.frontend import cover_letter_view, resume_studio_page
from resume_tailor.infrastructure.dependencies import create_tailor_service
from resume_tailor.infrastructure.optimization import DeterministicResumeOptimizer
from tests.cover_letter_helpers import cover_letter_case


def _posting(
    company: str = "Anduril Industries",
    title: str = "2027 Electrical Engineer Intern",
) -> JobPosting:
    return JobPosting(
        id="demo-posting",
        title=title,
        company_name=company,
        description="Build and test electrical hardware and embedded systems.",
    )


def _profile() -> MasterProfile:
    entries = [
        ResumeItem(
            id="telebotics-mechatronics-engineer",
            title="Mechatronics Engineer",
            kind=EntityKind.EXPERIENCE,
        ),
        ResumeItem(
            id="lassonde-rd-hardware-engineer",
            title="R&D Hardware Engineer",
            kind=EntityKind.EXPERIENCE,
        ),
        ResumeItem(
            id="robotic-hand",
            title="Vision Controlled Robotic Hand",
            kind=EntityKind.PROJECT,
        ),
    ]
    return MasterProfile(
        id="demo-profile",
        user_id="demo-user",
        display_name="Demo Candidate",
        experiences=entries[:2],
        projects=entries[2:],
    )


def _evidence(profile: MasterProfile) -> list[EvidenceItem]:
    return [
        EvidenceItem(
            id=f"demo-{index}",
            entity_id=entry.id,
            source_text=f"Reviewed fact {index}.",
        )
        for index, entry in enumerate(profile.experiences + profile.projects)
        for _ in range(1)
    ]


def _complete_demo_profile() -> MasterProfile:
    entries = {
        "telebotics-mechatronics-engineer": ResumeItem(
            id="telebotics-mechatronics-engineer",
            title="Mechatronics Engineer",
            organization="Telebotics",
            kind=EntityKind.EXPERIENCE,
        ),
        "lassonde-rd-hardware-engineer": ResumeItem(
            id="lassonde-rd-hardware-engineer",
            title="R&D Hardware Engineer",
            organization="Lassonde School of Engineering – York University",
            kind=EntityKind.EXPERIENCE,
        ),
        "robotic-hand": ResumeItem(
            id="robotic-hand",
            title="Vision Controlled Robotic Hand",
            kind=EntityKind.PROJECT,
        ),
        "sodium-silicate": ResumeItem(
            id="sodium-silicate",
            title="Preventing Sodium Silicate Crystal Build-up in Holding Tanks",
            kind=EntityKind.PROJECT,
        ),
        "robotic-arm": ResumeItem(
            id="robotic-arm",
            title="Long Reach Robotic Arm Manipulator",
            kind=EntityKind.PROJECT,
        ),
    }
    evidence = [
        EvidenceItem(
            id=f"demo-evidence-{index:02d}",
            entity_id=entry_id,
            source_text=f"{snippet}. {text}",
            source_reference="sanitized reviewed profile",
        )
        for index, (entry_id, snippet, text) in enumerate(
            demo_mode._DEMO_RESUME_SPECS,
            start=1,
        )
    ]
    reviewed_skills = list(
        dict.fromkeys(
            skill
            for skills in demo_mode._DEMO_SKILL_GROUPS.values()
            for skill in skills
        )
    )
    return MasterProfile(
        id="demo-profile-complete",
        user_id="demo-user",
        display_name="Shiv Arora",
        contact=ContactInfo(
            email="demo@example.com",
            phone="555-0100",
            location="Toronto, ON",
        ),
        education=[
            EducationRecord(
                school="University of Toronto",
                program="Mechanical Engineering",
                minor_or_specialization="Robotics & Mechatronics",
            )
        ],
        experiences=[
            entries["telebotics-mechatronics-engineer"],
            entries["lassonde-rd-hardware-engineer"],
        ],
        projects=[
            entries["robotic-hand"],
            entries["sodium-silicate"],
            entries["robotic-arm"],
        ],
        technical_skills=[
            TechnicalSkillCategory(category="Reviewed skills", values=reviewed_skills)
        ],
        evidence=evidence,
    )


def _disabled_settings() -> Settings:
    return Settings(
        llm_enable_composition=False,
        llm_enable_bullet_rewrite=False,
        llm_enable_cover_letter=False,
        llm_enable_role_classification=False,
        llm_enable_opportunity_analysis=False,
    )


def test_demo_activation_requires_explicit_flag_not_company_or_title(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.delenv("VIEGO_DEMO_MODE", raising=False)
    assert not demo_mode.is_demo_application(_posting())
    assert not demo_mode.is_demo_details("Cohere", "Software Engineering Intern")
    monkeypatch.setenv("VIEGO_DEMO_MODE", "1")
    arbitrary = _posting(
        company="Cohere",
        title="Software Engineering Intern",
    )
    assert demo_mode.is_demo_application(arbitrary)
    assert demo_mode.is_demo_details("Cohere", "Software Engineering Intern")
    monkeypatch.setenv("VIEGO_DEMO_MODE", "0")
    assert not demo_mode.is_demo_application(arbitrary)
    assert not demo_mode.is_demo_details("Cohere", "Software Engineering Intern")


def test_demo_plan_and_cover_letter_boundary_use_only_resolved_reviewed_atoms(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("VIEGO_DEMO_MODE", "1")
    profile = _profile()
    evidence = _evidence(profile)
    profile = profile.model_copy(update={"evidence": evidence})
    monkeypatch.setattr(demo_mode, "_resolve_demo_evidence", lambda _profile: evidence)
    arbitrary_posting = _posting(
        company="Cohere",
        title="Software Engineering Intern",
    )
    base = DeterministicResumeOptimizer().create_plan(
        profile,
        arbitrary_posting,
        TemplateConstraints(),
    )
    plan = demo_mode.build_demo_resume_plan(
        profile,
        arbitrary_posting,
        TemplateConstraints(),
        base,
    )
    assert plan.selected_claim_ids == [item.id for item in evidence]
    assert plan.selected_entity_ids == [item.entity_id for item in evidence]
    demo_mode.validate_demo_plan(plan, profile)
    records, diagnostic = demo_mode.build_demo_cover_letter_evidence(profile, arbitrary_posting)
    assert diagnostic.selected_evidence_ids == [item.id for item in evidence]
    assert {record.entity_id for record in records} == set(plan.selected_entity_ids)


def test_demo_cover_letter_service_disables_provider_for_arbitrary_posting(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("VIEGO_DEMO_MODE", "1")
    profile, _, _ = cover_letter_case(
        company="Cohere",
        title="Software Engineering Intern",
    )
    evidence = list(profile.evidence)
    monkeypatch.setattr(demo_mode, "_resolve_demo_evidence", lambda _profile: evidence)
    posting = JobPosting(
        id="demo-posting",
        title="Software Engineering Intern",
        company_name="Cohere",
        description="Build and test electrical hardware and embedded systems.",
    )
    plan = DeterministicResumeOptimizer().create_plan(profile, posting, TemplateConstraints())
    request = CoverLetterService().create_request(profile, posting, plan)
    _, diagnostic = CoverLetterService()._provider_output(request)
    assert diagnostic.request_count == 0
    assert diagnostic.failure_code == "temporary_demo_override"


def test_demo_plan_rejects_changed_canonical_text(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("VIEGO_DEMO_MODE", "1")
    profile = _profile()
    evidence = _evidence(profile)
    profile = profile.model_copy(update={"evidence": evidence})
    monkeypatch.setattr(demo_mode, "_resolve_demo_evidence", lambda _profile: evidence)
    base = DeterministicResumeOptimizer().create_plan(
        profile,
        _posting(),
        TemplateConstraints(),
    )
    plan = demo_mode.build_demo_resume_plan(profile, _posting(), TemplateConstraints(), base)
    changed = plan.claim_candidates[0].model_copy(update={"text": "unsupported change"})
    with pytest.raises(ValueError, match="changed canonical evidence"):
        demo_mode.validate_demo_plan(
            plan.model_copy(
                update={"claim_candidates": [changed, *plan.claim_candidates[1:]]}
            ),
            profile,
        )


def test_final_demo_resume_and_cover_letter_artifacts_are_exact_and_provider_free(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("VIEGO_DEMO_MODE", "1")
    profile = _complete_demo_profile()
    posting = _posting()
    service = create_tailor_service(_disabled_settings())
    plan = service.create_plan(profile, posting, TemplateConstraints())

    resume_artifact = service.build_generated_artifact(plan, profile, set())
    resume = resume_artifact.final_resume
    assert [item.title for item in resume.experiences] == [
        "Mechatronics Engineer",
        "R&D Hardware Engineer",
    ]
    assert [item.title for item in resume.projects] == [
        "Vision Controlled Robotic Hand",
        "Preventing Sodium Silicate Crystal Build-up in Holding Tanks",
        "Long Reach Robotic Arm Manipulator",
    ]
    assert [len(resume.experience_bullets[item.id]) for item in resume.experiences] == [5, 4]
    assert [len(resume.project_bullets[item.id]) for item in resume.projects] == [4, 3, 2]
    actual_bullets = [
        bullet.text
        for item in [*resume.experiences, *resume.projects]
        for bullet in (
            resume.experience_bullets[item.id]
            if item.kind is EntityKind.EXPERIENCE
            else resume.project_bullets[item.id]
        )
    ]
    assert actual_bullets == [text for _, _, text in demo_mode._DEMO_RESUME_SPECS]
    rendered_text = "\n".join(
        paragraph.text for paragraph in Document(BytesIO(resume_artifact.docx_bytes)).paragraphs
    )
    assert all(text in rendered_text for text in actual_bullets)
    assert not any(
        name in rendered_text for name in ("EXL", "Stush", "Crest", "Resume Tailor")
    )
    assert resume_artifact.pagination_diagnostic.status == "exact"
    assert resume_artifact.call_counts.provider_calls == 0
    assert resume_studio_page._is_demo_resume_artifact(resume_artifact)

    research_request = CompanyResearchRequest(
        company_name=posting.company_name,
        role_title=posting.title,
        posting_fingerprint=content_fingerprint(posting),
        posting_description=posting.description,
        enabled=False,
    )
    letter_artifact = service.generate_cover_letter_artifact(
        profile,
        posting,
        plan,
        final_resume=resume,
        research_request=research_request,
        date_text="August 26, 2026",
    )
    assert tuple(
        paragraph.text for paragraph in letter_artifact.letter.paragraphs
    ) == demo_mode._DEMO_COVER_LETTER_PARAGRAPHS
    assert letter_artifact.letter.salutation == "Dear Anduril Hiring Team,"
    assert letter_artifact.letter.signoff == "Sincerely,"
    assert letter_artifact.letter.signoff_name == "Shiv Arora"
    assert letter_artifact.ready_for_review
    assert letter_artifact.review_state is CoverLetterReviewState.GENERATED_AWAITING_REVIEW
    assert letter_artifact.provider_diagnostic.status is CoverLetterProviderStatus.DETERMINISTIC
    assert letter_artifact.provider_diagnostic.fallback_reason is None
    assert letter_artifact.call_counts.provider_calls == 0
    assert letter_artifact.page_fit.exact_pagination
    assert letter_artifact.page_fit.page_count == 1
    approved_letter = service.approve_cover_letter_artifact(
        letter_artifact,
        expected_fingerprint=letter_artifact.artifact_fingerprint,
    )
    downloaded_letter = service.prepare_cover_letter_download(
        approved_letter,
        expected_fingerprint=approved_letter.artifact_fingerprint,
    )
    assert downloaded_letter.docx_bytes == letter_artifact.docx_bytes
    rendered_letter_paragraphs = [
        paragraph.text
        for paragraph in Document(BytesIO(letter_artifact.docx_bytes)).paragraphs
        if paragraph.text
    ]
    expected_letter_sequence = [
        "Dear Anduril Hiring Team,",
        *demo_mode._DEMO_COVER_LETTER_PARAGRAPHS,
        "Sincerely,",
        "Shiv Arora",
    ]
    body_start = rendered_letter_paragraphs.index("Dear Anduril Hiring Team,")
    assert rendered_letter_paragraphs[
        body_start : body_start + len(expected_letter_sequence)
    ] == expected_letter_sequence
    status_items: dict[str, str] = {}
    captions: list[str] = []
    monkeypatch.setattr(
        cover_letter_view,
        "render_status_strip",
        lambda _streamlit, items: status_items.update(items),
    )
    monkeypatch.setattr(
        cover_letter_view,
        "st",
        SimpleNamespace(caption=captions.append),
    )
    cover_letter_view._render_status(letter_artifact, True)
    assert status_items == {
        "Cover Letter": "Ready for review",
        "Pagination": "Fits one page",
    }
    assert captions == ["Current job"]
    assert cover_letter_view._writer_status(letter_artifact) == "Prepared"
    complete_letter = " ".join(
        paragraph.text for paragraph in letter_artifact.letter.paragraphs
    ).casefold()
    assert not any(
        phrase in complete_letter
        for phrase in (
            "u.s. person",
            "work authorization",
            "altium",
            "oscilloscope",
            "device tree",
            "bootloader",
            "pcie",
        )
    )
