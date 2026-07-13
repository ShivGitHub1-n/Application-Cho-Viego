from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from resume_tailor.application.services import TailorResumeService
from resume_tailor.domain.models import (
    JobPosting,
    MasterProfile,
    StructuredResume,
    TailoringPlan,
    TemplateConstraints,
)
from resume_tailor.domain.profile_completeness import (
    validate_master_profile_completeness,
)
from resume_tailor.infrastructure.optimization import (
    DeterministicResumeOptimizer,
    EvidenceBoundResumeWriter,
)
from resume_tailor.infrastructure.rendering import (
    ManagedResumeRenderer,
    PageCountVerificationError,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "manual-test" / "generated-reference-layout-resume.docx"


def build_manual_resume() -> tuple[MasterProfile, TailoringPlan, StructuredResume]:
    profile_payload = json.loads(
        (ROOT / "manual-test" / "profile.json").read_text(encoding="utf-8")
    )
    profile = MasterProfile.model_validate(profile_payload)
    description = (ROOT / "manual-test" / "job-description.txt.txt").read_text(
        encoding="utf-8"
    )
    posting = JobPosting(
        id="manual-reference-render",
        title="AI Researcher Co-op",
        description=description,
    )
    service = TailorResumeService(
        DeterministicResumeOptimizer(),
        EvidenceBoundResumeWriter(),
    )
    plan = service.create_plan(profile, posting, TemplateConstraints())
    resume = service.build_document(plan, profile, set())
    return profile, plan, resume


def field_presence_diagnostic(
    profile: MasterProfile,
    plan: TailoringPlan,
    resume: StructuredResume,
) -> dict[str, object]:
    """Report structure only; never emit contact values or evidence text."""

    def education(records: list[object]) -> dict[str, object]:
        return {
            "records": len(records),
            "start_date_present": [bool(item.start_date) for item in records],
            "graduation_date_present": [
                bool(item.expected_graduation_date or item.graduation_date)
                for item in records
            ],
            "location_present": [bool(item.location) for item in records],
            "awards_count": [len(item.awards) for item in records],
            "coursework_count": [len(item.relevant_coursework) for item in records],
        }

    def entries(records: list[object], bullet_counts: dict[str, int]) -> dict[str, object]:
        return {
            "count": len(records),
            "dates_present": [bool(item.start_date or item.end_date) for item in records],
            "locations_present": [bool(item.location) for item in records],
            "selected_bullet_counts": [bullet_counts.get(item.id, 0) for item in records],
        }

    return {
        "master_profile": {
            "education": education(profile.education),
            "technical_skill_category_count": len(profile.technical_skills),
            "skills_per_category": [len(item.skills) for item in profile.technical_skills],
            "experience_count": len(profile.experiences),
            "project_count": len(profile.projects),
            "top_level_coursework_count": len(profile.coursework),
        },
        "tailoring_plan": {
            "education": education(plan.education),
            "technical_skill_category_count": len(plan.technical_skills),
            "skills_per_category": [len(item.skills) for item in plan.technical_skills],
            "legacy_selected_skill_count": len(plan.selected_skills),
            "selected_coursework_count": len(plan.selected_coursework),
            "experience_count": len(plan.selected_experiences),
            "project_count": len(plan.selected_projects),
        },
        "structured_resume": {
            "education": education(resume.education),
            "technical_skill_category_count": len(resume.technical_skills),
            "skills_per_category": [len(item.skills) for item in resume.technical_skills],
            "legacy_selected_skill_count": len(resume.selected_skills),
            "selected_coursework_count": len(resume.selected_coursework),
            "experiences": entries(
                resume.experiences,
                {key: len(value) for key, value in resume.experience_bullets.items()},
            ),
            "projects": entries(
                resume.projects,
                {key: len(value) for key, value in resume.project_bullets.items()},
            ),
        },
    }


def geometry_report(renderer: ManagedResumeRenderer, output: Path | None) -> dict[str, object]:
    page = renderer.layout_profile.page

    def paragraph_metrics(path: Path) -> dict[str, object]:
        document = Document(path)
        paragraphs = document.paragraphs

        def spacing(paragraph: object, name: str) -> int:
            value = getattr(paragraph.paragraph_format, name)
            return value.twips if value is not None else 0

        before = sum(spacing(paragraph, "space_before") for paragraph in paragraphs)
        after = sum(spacing(paragraph, "space_after") for paragraph in paragraphs)
        return {
            "paragraph_count": len(paragraphs),
            "empty_paragraph_count": sum(not paragraph.text.strip() for paragraph in paragraphs),
            "explicit_before_twips": before,
            "explicit_after_twips": after,
            "estimated_spacing_height_points": (before + after) / 20,
        }

    report: dict[str, object] = {
        "reference": {
            "width_twips": page.width_twips,
            "height_twips": page.height_twips,
            "margins_twips": {
                "top": page.top_margin_twips,
                "right": page.right_margin_twips,
                "bottom": page.bottom_margin_twips,
                "left": page.left_margin_twips,
            },
            "usable_width_twips": page.usable_width_twips,
            "usable_height_twips": page.usable_height_twips,
        },
        "generated": None,
        "reference_spacing_diagnostics": paragraph_metrics(
            ROOT / "manual-test" / "reference-resume.docx"
        ),
        "metadata_anchor_groups": [
            {
                "group_id": group.group_id,
                "observed_positions_twips": group.observed_positions_twips,
                "representative_position_twips": group.representative_position_twips,
                "role_groups": group.role_groups,
                "tolerance_twips": group.tolerance_twips,
                "relative_tolerance": group.relative_tolerance,
                "provenance": group.provenance,
            }
            for group in renderer.layout_profile.metadata_anchor_groups
        ],
        "role_geometry": {
            role_name: {
                "left_indent_twips": role.paragraph.left_indent_twips.value,
                "first_line_indent_twips": role.paragraph.first_line_indent_twips.value,
                "hanging_indent_twips": role.paragraph.hanging_indent_twips.value,
                "font_size_half_points": role.primary_typography.font_size_half_points.value,
                "line_spacing_twips": role.paragraph.line_spacing_twips.value,
                "line_spacing_rule": role.paragraph.line_spacing_rule.value,
                "metadata_anchor_group_ids": role.metadata_anchor_group_ids,
                "bullet_left_indent_twips": role.bullet.left_indent_twips if role.bullet else None,
                "bullet_hanging_indent_twips": role.bullet.hanging_indent_twips if role.bullet else None,
            }
            for role_name, role in renderer.layout_profile.semantic_roles.items()
            if role_name in {
                "name",
                "contact_line",
                "section_heading",
                "education_institution_date_row",
                "education_program_location_row",
                "experience_title_date_row",
                "employer_location_row",
                "experience_bullet",
            }
        },
        "initial_page_count": (
            renderer.initial_measurement.page_count
            if renderer.initial_measurement is not None
            else None
        ),
        "final_page_count": (
            renderer.last_measurement.page_count
            if renderer.last_measurement is not None
            else None
        ),
        "measurement_provider": (
            renderer.last_measurement.provider
            if renderer.last_measurement is not None
            else None
        ),
        "measurement_confidence": (
            renderer.last_measurement.confidence
            if renderer.last_measurement is not None
            else None
        ),
        "exact_one_page_verified": (
            renderer.last_measurement.exact and renderer.last_measurement.page_count == 1
            if renderer.last_measurement is not None
            else False
        ),
        "underfill_expansion_disabled": not renderer.underfill_expansion_enabled,
        "overflow_reduction_count": renderer.last_overflow_reduction_count,
    }
    if output is not None and output.is_file():
        section = Document(output).sections[0]
        report["generated"] = {
            "width_twips": section.page_width.twips,
            "height_twips": section.page_height.twips,
            "margins_twips": {
                "top": section.top_margin.twips,
                "right": section.right_margin.twips,
                "bottom": section.bottom_margin.twips,
                "left": section.left_margin.twips,
            },
            "usable_width_twips": (
                section.page_width.twips
                - section.left_margin.twips
                - section.right_margin.twips
            ),
            "usable_height_twips": (
                section.page_height.twips
                - section.top_margin.twips
                - section.bottom_margin.twips
            ),
            "spacing_diagnostics": paragraph_metrics(output),
        }
    return report


def main() -> None:
    profile, plan, resume = build_manual_resume()
    completeness = validate_master_profile_completeness(profile)
    print(json.dumps({"profile_completeness": completeness.model_dump(mode="json")}, indent=2))
    print(json.dumps({"field_presence": field_presence_diagnostic(profile, plan, resume)}, indent=2))
    renderer = ManagedResumeRenderer()
    try:
        renderer.render_docx(resume, OUTPUT)
    except PageCountVerificationError as error:
        report = geometry_report(renderer, OUTPUT)
        report["verification_error"] = str(error)
        print(json.dumps({"geometry_report": report}, indent=2))
        raise
    print(json.dumps({"geometry_report": geometry_report(renderer, OUTPUT)}, indent=2))
    print(json.dumps({"output": str(OUTPUT), "selected_claim_count": len(plan.selected_claim_ids)}))


if __name__ == "__main__":
    main()
