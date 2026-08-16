from __future__ import annotations

import re
import zlib
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from resume_tailor.domain.models import ContactHyperlink


class ResumeExtractionError(ValueError):
    pass


class UnsupportedResumeFileError(ResumeExtractionError):
    pass


class EmptyResumeFileError(ResumeExtractionError):
    pass


class ImageOnlyResumeError(ResumeExtractionError):
    pass


@dataclass(frozen=True)
class ExtractedResumeText:
    filename: str
    source_format: str
    text: str
    contact_links: tuple[ContactHyperlink, ...] = ()


def extract_resume_text(filename: str, content: bytes) -> ExtractedResumeText:
    suffix = Path(filename).suffix.casefold()
    if suffix not in {".docx", ".pdf"}:
        raise UnsupportedResumeFileError("Only .docx and text-based .pdf files are supported.")
    if not content:
        raise EmptyResumeFileError("The uploaded resume file is empty.")
    if suffix == ".docx":
        text, contact_links = _extract_docx(content)
        source_format = "docx"
    else:
        text = _extract_pdf(content)
        contact_links = ()
        source_format = "pdf"
    normalized = "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").split("\n")).strip()
    if not normalized:
        raise ImageOnlyResumeError(
            "No selectable text was found. Image-only resumes require OCR, which is not enabled."
        )
    return ExtractedResumeText(
        filename=filename,
        source_format=source_format,
        text=normalized,
        contact_links=contact_links,
    )


def _extract_docx(content: bytes) -> tuple[str, tuple[ContactHyperlink, ...]]:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        body_lines = _docx_story_lines(document)
        contact_links = _docx_story_hyperlinks(document)
        other_lines: list[str] = []
        visited_parts: set[str] = set()
        for section in document.sections:
            stories = (
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            )
            for story in stories:
                part_name = str(story.part.partname)
                if part_name in visited_parts:
                    continue
                visited_parts.add(part_name)
                other_lines.extend(_docx_story_lines(story))
                contact_links.extend(_docx_story_hyperlinks(story))
        seen = {line.casefold() for line in body_lines if line}
        for line in other_lines:
            normalized = line.casefold()
            if line and normalized not in seen:
                body_lines.append(line)
                seen.add(normalized)
        unique_links: list[ContactHyperlink] = []
        seen_links: set[tuple[str, str]] = set()
        for link in contact_links:
            key = (link.display_text.casefold(), link.destination.casefold().rstrip("/"))
            if key not in seen_links:
                unique_links.append(link)
                seen_links.add(key)
        return "\n".join(body_lines), tuple(unique_links)
    except (BadZipFile, ValueError, OSError) as error:
        raise ResumeExtractionError("The DOCX file is corrupt or unreadable.") from error
    except Exception as error:
        raise ResumeExtractionError("The DOCX file could not be read.") from error


def _docx_story_lines(story: object) -> list[str]:
    lines = [
        text
        for paragraph in getattr(story, "paragraphs", ())
        if (text := _docx_paragraph_text(paragraph))
    ]
    for table in getattr(story, "tables", ()):
        for row in table.rows:
            cell_texts: list[str] = []
            visited_cells: set[int] = set()
            for cell in row.cells:
                cell_identity = id(cell._tc)
                if cell_identity in visited_cells:
                    continue
                visited_cells.add(cell_identity)
                cell_text = " ".join(
                    text
                    for paragraph in cell.paragraphs
                    if (text := _docx_paragraph_text(paragraph))
                ).strip()
                if cell_text:
                    cell_texts.append(cell_text)
            if cell_texts:
                lines.append(" | ".join(cell_texts))
    return lines


def _docx_story_hyperlinks(story: object) -> list[ContactHyperlink]:
    links = [
        link
        for paragraph in getattr(story, "paragraphs", ())
        for link in _docx_paragraph_hyperlinks(paragraph)
    ]
    for table in getattr(story, "tables", ()):
        visited_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_identity = id(cell._tc)
                if cell_identity in visited_cells:
                    continue
                visited_cells.add(cell_identity)
                links.extend(
                    link
                    for paragraph in cell.paragraphs
                    for link in _docx_paragraph_hyperlinks(paragraph)
                )
    return links


def _docx_paragraph_hyperlinks(paragraph: object) -> list[ContactHyperlink]:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn

    links: list[ContactHyperlink] = []
    for child in paragraph._p.iterchildren():
        if child.tag != qn("w:hyperlink"):
            continue
        relationship_id = child.get(qn("r:id"))
        relationship = paragraph.part.rels.get(relationship_id) if relationship_id else None
        if relationship is None or relationship.reltype != RT.HYPERLINK:
            continue
        target = str(relationship.target_ref).strip()
        display = "".join(
            node.text or "" for node in child.iter() if node.tag == qn("w:t")
        ).strip()
        if display and _is_visible_contact_link(target):
            links.append(ContactHyperlink(display_text=display, destination=target))
    return links


def _docx_paragraph_text(paragraph: object) -> str:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn

    fragments: list[str] = []
    for child in paragraph._p.iterchildren():
        child_fragments: list[str] = []
        for node in child.iter():
            if node.tag == qn("w:t") and node.text:
                child_fragments.append(node.text)
            elif node.tag == qn("w:tab"):
                child_fragments.append("\t")
            elif node.tag in {qn("w:br"), qn("w:cr")}:
                child_fragments.append("\n")
        visible = "".join(child_fragments)
        fragments.append(visible)
        if child.tag != qn("w:hyperlink"):
            continue
        relationship_id = child.get(qn("r:id"))
        relationship = paragraph.part.rels.get(relationship_id) if relationship_id else None
        if relationship is None or relationship.reltype != RT.HYPERLINK:
            continue
        target = str(relationship.target_ref).strip()
        if not _is_visible_contact_link(target):
            continue
        if target.casefold() not in visible.casefold():
            fragments.append(f" <{target}>")
    return re.sub(r"[ \t]+", " ", "".join(fragments)).strip()


def _is_visible_contact_link(target: str) -> bool:
    normalized = target.casefold()
    return normalized.startswith(
        ("https://", "http://", "https\\://", "http\\://", "mailto:", "www.")
    )


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return _extract_pdf_text_fallback(content)
    try:
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as error:
        raise ResumeExtractionError("The PDF file is corrupt or unreadable.") from error


def _extract_pdf_text_fallback(content: bytes) -> str:
    """Extract common literal-string PDF text operators when no PDF library is installed."""

    try:
        source = content.decode("latin-1")
    except UnicodeDecodeError as error:
        raise ResumeExtractionError("The PDF encoding could not be read.") from error
    if not source.startswith("%PDF-"):
        raise ResumeExtractionError("The PDF file is corrupt or unreadable.")
    compressed_text = []
    for match in re.finditer(
        rb"<<(?P<dictionary>.*?)>>\s*stream\r?\n(?P<data>.*?)\r?\nendstream",
        content,
        re.DOTALL,
    ):
        if b"/FlateDecode" not in match.group("dictionary"):
            continue
        try:
            compressed_text.append(zlib.decompress(match.group("data")).decode("latin-1"))
        except (zlib.error, UnicodeDecodeError):
            continue
    source = "\n".join([source, *compressed_text])
    strings = re.findall(r"\(((?:[^\\()]|\\.)*)\)\s*T[Jj]", source)
    decoded = []
    for value in strings:
        decoded.append(
            value.replace(r"\(", "(").replace(r"\)", ")").replace(r"\\", "\\")
        )
    return "\n".join(decoded)
