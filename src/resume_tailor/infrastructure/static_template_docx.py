from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from lxml import etree  # type: ignore[import-untyped]

from resume_tailor.domain.models import (
    EducationRecord,
    EntityKind,
    ResumeItem,
    StructuredBullet,
    StructuredResume,
)
from resume_tailor.domain.resume_metadata import (
    compose_date_range,
    education_end_date,
    validate_structured_resume_metadata,
)
from resume_tailor.infrastructure.template_v1 import template_v1_docx_path


class StaticTemplateRenderError(ValueError):
    pass


_PLACEHOLDER = re.compile(r"\{\{[A-Z0-9_]+\}\}")
_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def render_template_v1_resume(
    resume: StructuredResume,
    output_path: Path,
    *,
    template_path: Path | None = None,
) -> Path:
    """Populate the packaged static Template V1 DOCX without reconstructing formatting."""

    return StaticTemplateV1Renderer(template_path=template_path).render(
        resume,
        output_path,
    )


class StaticTemplateV1Renderer:
    def __init__(self, *, template_path: Path | None = None) -> None:
        self._template_path = (template_path or template_v1_docx_path()).resolve()

    @property
    def template_path(self) -> Path:
        return self._template_path

    def render(self, resume: StructuredResume, output_path: Path) -> Path:
        if not self._template_path.is_file():
            raise StaticTemplateRenderError(
                f"Packaged Template V1 DOCX is unavailable: {self._template_path}"
            )
        validate_structured_resume_metadata(resume)
        output_path = Path(output_path).expanduser().resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Passing the packaged path is deliberate: Template V1 must never start
        # from python-docx's blank-document default.
        document = Document(str(self._template_path))
        catalog = _TemplateCatalog(document)
        rendered = self._rendered_paragraphs(resume, catalog)
        _replace_document_body(document, rendered)
        _validate_rendered_body(document)
        document.save(str(output_path))
        return output_path

    def _rendered_paragraphs(
        self,
        resume: StructuredResume,
        catalog: _TemplateCatalog,
    ) -> list[etree._Element]:
        paragraphs = [
            catalog.populate("{{NAME}}", {"{{NAME}}": resume.display_name}),
        ]
        if resume.contact_line:
            paragraphs.append(
                catalog.populate(
                    "{{CONTACT_LINE}}",
                    {"{{CONTACT_LINE}}": resume.contact_line},
                )
            )
        if resume.education:
            paragraphs.append(catalog.heading("Education"))
            paragraphs.extend(self._education_paragraphs(resume, catalog))
        if resume.technical_skills or resume.selected_skills:
            skill_paragraphs = self._skill_paragraphs(resume, catalog)
            if skill_paragraphs:
                paragraphs.append(catalog.heading("Technical Skills"))
                paragraphs.extend(skill_paragraphs)
        if resume.experiences or resume.experience_bullets:
            paragraphs.append(catalog.heading("Technical Experience"))
            paragraphs.extend(self._experience_paragraphs(resume, catalog))
        if resume.projects or resume.project_bullets:
            paragraphs.append(catalog.heading("Projects"))
            paragraphs.extend(self._project_paragraphs(resume, catalog))
        return paragraphs

    def _education_paragraphs(
        self,
        resume: StructuredResume,
        catalog: _TemplateCatalog,
    ) -> list[etree._Element]:
        paragraphs: list[etree._Element] = []
        for record in resume.education:
            dates = compose_date_range(record.start_date, education_end_date(record))
            paragraphs.append(
                catalog.populate(
                    "{{EDUCATION_INSTITUTION}}",
                    {
                        "{{EDUCATION_INSTITUTION}}": record.school,
                        "{{EDUCATION_DATES}}": dates,
                    },
                    retain_tab=bool(dates),
                )
            )
            program = _education_program(record)
            if program or record.location:
                paragraphs.append(
                    catalog.populate(
                        "{{EDUCATION_PROGRAM}}",
                        {
                            "{{EDUCATION_PROGRAM}}": program,
                            "{{EDUCATION_LOCATION}}": record.location,
                        },
                        retain_tab=bool(record.location),
                    )
                )
            awards = f"Awards: {', '.join(record.awards)}" if record.awards else ""
            gpa = f"{', ' if awards else ''}GPA: {record.gpa}" if record.gpa else ""
            if awards or gpa:
                paragraphs.append(
                    catalog.populate(
                        "{{EDUCATION_AWARDS}}",
                        {
                            "{{EDUCATION_AWARDS}}": awards,
                            "{{EDUCATION_GPA}}": gpa,
                        },
                    )
                )
            coursework = record.relevant_coursework or resume.selected_coursework
            if coursework:
                paragraphs.append(
                    catalog.populate(
                        "{{EDUCATION_COURSEWORK}}",
                        {
                            "{{EDUCATION_COURSEWORK}}": (
                                f"Relevant Courses: {', '.join(coursework)}"
                            )
                        },
                    )
                )
        return paragraphs

    def _skill_paragraphs(
        self,
        resume: StructuredResume,
        catalog: _TemplateCatalog,
    ) -> list[etree._Element]:
        rows: list[tuple[str, list[str]]] = []
        if resume.technical_skills:
            for category in resume.technical_skills:
                values = category.values or [skill.value for skill in category.skills]
                if values:
                    rows.append((category.category, values))
        elif resume.selected_skills:
            rows.append(("Skills", resume.selected_skills))
        return [
            catalog.populate(
                "{{SKILL_CATEGORY}}",
                {
                    "{{SKILL_CATEGORY}}": category,
                    "{{SKILL_SEPARATOR}}": ": ",
                    "{{SKILL_VALUES}}": ", ".join(values),
                },
            )
            for category, values in rows
        ]

    def _experience_paragraphs(
        self,
        resume: StructuredResume,
        catalog: _TemplateCatalog,
    ) -> list[etree._Element]:
        paragraphs: list[etree._Element] = []
        records = _ordered_records(
            resume.experiences,
            resume.experience_bullets,
            resume.entity_titles,
            EntityKind.EXPERIENCE,
        )
        for index, item in enumerate(records):
            repeated = index > 0
            prefix = "EXPERIENCE_REPEAT" if repeated else "EXPERIENCE"
            marker = f"{{{{{prefix}_TITLE}}}}"
            subtitle = item.subtitle or item.technology_label
            if subtitle and subtitle.casefold() in item.title.casefold():
                subtitle = None
            dates = compose_date_range(item.start_date, item.end_date)
            paragraphs.append(
                catalog.populate(
                    marker,
                    {
                        marker: item.title,
                        f"{{{{{prefix}_SUBTITLE_SEPARATOR}}}}": (" | " if subtitle else ""),
                        f"{{{{{prefix}_SUBTITLE}}}}": subtitle,
                        f"{{{{{prefix}_DATES}}}}": dates,
                    },
                    retain_tab=bool(dates),
                )
            )
            if item.organization or item.location:
                paragraphs.append(
                    catalog.populate(
                        f"{{{{{prefix}_ORGANIZATION}}}}",
                        {
                            f"{{{{{prefix}_ORGANIZATION}}}}": item.organization,
                            f"{{{{{prefix}_LOCATION}}}}": item.location,
                        },
                        retain_tab=bool(item.location),
                    )
                )
            bullet_marker = f"{{{{{prefix}_BULLET}}}}"
            for bullet in resume.experience_bullets.get(item.id, []):
                paragraphs.append(
                    catalog.populate(
                        bullet_marker,
                        {bullet_marker: bullet.text},
                    )
                )
        return paragraphs

    def _project_paragraphs(
        self,
        resume: StructuredResume,
        catalog: _TemplateCatalog,
    ) -> list[etree._Element]:
        paragraphs: list[etree._Element] = []
        records = _ordered_records(
            resume.projects,
            resume.project_bullets,
            resume.entity_titles,
            EntityKind.PROJECT,
        )
        for item in records:
            title = item.title
            if (
                item.award_or_placement
                and item.award_or_placement.casefold() not in title.casefold()
            ):
                title = f"{title} ({item.award_or_placement})"
            technologies = item.technology_label or ", ".join(item.technologies)
            if technologies and technologies.casefold() in title.casefold():
                technologies = ""
            dates = compose_date_range(item.start_date, item.end_date)
            paragraphs.append(
                catalog.populate(
                    "{{PROJECT_TITLE}}",
                    {
                        "{{PROJECT_TITLE}}": title,
                        "{{PROJECT_TECHNOLOGY_SEPARATOR}}": (" | " if technologies else ""),
                        "{{PROJECT_TECHNOLOGIES}}": technologies,
                        "{{PROJECT_DATES}}": dates,
                    },
                    retain_tab=bool(dates),
                )
            )
            if item.organization or item.location:
                paragraphs.append(
                    catalog.populate(
                        "{{PROJECT_ORGANIZATION}}",
                        {
                            "{{PROJECT_ORGANIZATION}}": item.organization,
                            "{{PROJECT_LOCATION}}": item.location,
                        },
                        retain_tab=bool(item.location),
                    )
                )
            for bullet in resume.project_bullets.get(item.id, []):
                paragraphs.append(
                    catalog.populate(
                        "{{PROJECT_BULLET}}",
                        {"{{PROJECT_BULLET}}": bullet.text},
                    )
                )
        return paragraphs


class _TemplateCatalog:
    def __init__(self, document: DocumentType) -> None:
        self._paragraphs = [deepcopy(paragraph._p) for paragraph in document.paragraphs]
        self._by_heading = {
            paragraph.text: deepcopy(paragraph._p)
            for paragraph in document.paragraphs
            if paragraph.text
            in {"Education", "Technical Skills", "Technical Experience", "Projects"}
        }
        if set(self._by_heading) != {
            "Education",
            "Technical Skills",
            "Technical Experience",
            "Projects",
        }:
            raise StaticTemplateRenderError(
                "Packaged Template V1 is missing one or more section-heading prototypes."
            )

    def heading(self, label: str) -> etree._Element:
        paragraph = deepcopy(self._by_heading[label])
        _strip_template_markup(paragraph)
        return paragraph

    def populate(
        self,
        locator: str,
        replacements: dict[str, str | None],
        *,
        retain_tab: bool = True,
    ) -> etree._Element:
        matches = [
            paragraph for paragraph in self._paragraphs if locator in _paragraph_text(paragraph)
        ]
        if len(matches) != 1:
            raise StaticTemplateRenderError(
                f"Template placeholder locator {locator!r} matched {len(matches)} paragraphs."
            )
        paragraph = deepcopy(matches[0])
        _strip_template_markup(paragraph)
        for text_node in list(paragraph.iter(qn("w:t"))):
            marker = text_node.text or ""
            if marker not in replacements:
                continue
            value = replacements[marker] or ""
            if value:
                text_node.text = value
                if value.startswith(" ") or value.endswith(" "):
                    text_node.set(_XML_SPACE, "preserve")
                elif _XML_SPACE in text_node.attrib:
                    del text_node.attrib[_XML_SPACE]
            else:
                _remove_containing_run(text_node)
        if not retain_tab:
            _remove_tabs(paragraph)
        remaining = _PLACEHOLDER.findall(_paragraph_text(paragraph))
        if remaining:
            raise StaticTemplateRenderError(
                f"Unpopulated Template V1 placeholders remain: {sorted(set(remaining))}"
            )
        if not _paragraph_text(paragraph).strip():
            raise StaticTemplateRenderError(
                f"Template population produced a blank paragraph for {locator!r}."
            )
        return paragraph


def _replace_document_body(
    document: DocumentType,
    paragraphs: list[etree._Element],
) -> None:
    body = document._element.body
    section_properties = body.sectPr
    if section_properties is None:
        raise StaticTemplateRenderError("Packaged Template V1 has no section properties.")
    for child in list(body):
        if child is not section_properties:
            body.remove(child)
    for index, paragraph in enumerate(paragraphs, start=1):
        _assign_paragraph_identity(paragraph, index)
        body.insert(len(body) - 1, paragraph)


def _validate_rendered_body(document: DocumentType) -> None:
    paragraphs = document.paragraphs
    if not paragraphs:
        raise StaticTemplateRenderError("Template V1 population produced no body paragraphs.")
    blanks = [index for index, paragraph in enumerate(paragraphs) if not paragraph.text.strip()]
    if blanks:
        raise StaticTemplateRenderError(
            f"Template V1 population produced blank paragraphs at indexes {blanks}."
        )
    body_text = "\n".join(paragraph.text for paragraph in paragraphs)
    placeholders = _PLACEHOLDER.findall(body_text)
    if placeholders:
        raise StaticTemplateRenderError(
            f"Template V1 population left placeholders: {sorted(set(placeholders))}"
        )
    headings = [
        paragraph.text
        for paragraph in paragraphs
        if paragraph.text in {"Education", "Technical Skills", "Technical Experience", "Projects"}
    ]
    if len(headings) != len(set(headings)):
        raise StaticTemplateRenderError("Template V1 population duplicated a section heading.")


def _strip_template_markup(paragraph: etree._Element) -> None:
    elements = [
        *list(paragraph.iter(qn("w:bookmarkStart"))),
        *list(paragraph.iter(qn("w:bookmarkEnd"))),
    ]
    for element in elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _assign_paragraph_identity(paragraph: etree._Element, index: int) -> None:
    paragraph.set(f"{{{_W14}}}paraId", f"{index:08X}")
    paragraph.set(f"{{{_W14}}}textId", f"{(index + 0x10000000):08X}")
    for attribute in list(paragraph.attrib):
        namespace, _, local_name = attribute.rpartition("}")
        if namespace == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main" and (
            local_name.startswith("rsid")
        ):
            del paragraph.attrib[attribute]


def _remove_containing_run(text_node: etree._Element) -> None:
    parent = text_node.getparent()
    while parent is not None and parent.tag != qn("w:r"):
        parent = parent.getparent()
    if parent is not None and parent.getparent() is not None:
        parent.getparent().remove(parent)


def _remove_tabs(paragraph: etree._Element) -> None:
    for tab in list(paragraph.iter(qn("w:tab"))):
        parent = tab.getparent()
        if parent is not None and parent.tag == qn("w:r") and parent.getparent() is not None:
            parent.getparent().remove(parent)
        elif parent is not None:
            parent.remove(tab)


def _paragraph_text(paragraph: etree._Element) -> str:
    pieces: list[str] = []
    for element in paragraph.iter():
        if element.tag == qn("w:t") and element.text:
            pieces.append(element.text)
        elif element.tag == qn("w:tab"):
            pieces.append("\t")
    return "".join(pieces)


def _ordered_records(
    records: list[ResumeItem],
    bullets: dict[str, list[StructuredBullet]],
    titles: dict[str, str],
    kind: EntityKind,
) -> list[ResumeItem]:
    by_id = {item.id: item for item in records}
    ordered = list(records)
    for entity_id in bullets:
        if entity_id not in by_id:
            ordered.append(
                ResumeItem(
                    id=entity_id,
                    title=titles.get(entity_id, entity_id),
                    kind=kind,
                )
            )
    return ordered


def _education_program(record: EducationRecord) -> str:
    values = [
        record.program,
        record.minor_or_specialization,
        record.co_op_designation,
    ]
    unique: list[str] = []
    for value in values:
        if value and value.casefold() not in " ".join(unique).casefold():
            unique.append(value)
    return ", ".join(unique)


__all__ = [
    "StaticTemplateRenderError",
    "StaticTemplateV1Renderer",
    "render_template_v1_resume",
]
