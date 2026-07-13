from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from resume_tailor.domain.layout import (
    Border,
    BulletLayout,
    LayoutProfile,
    ParagraphLayout,
    SemanticRoleLayout,
    TabStop,
    Typography,
    TransitionSpacing,
)
from resume_tailor.domain.models import EntityKind, ResumeItem, StructuredBullet, StructuredResume


class AdaptiveDocxRenderError(ValueError):
    pass


@dataclass(frozen=True)
class _Fallbacks:
    """Generic values used only when a property is absent from LayoutProfile."""

    font_family: str = "Times New Roman"
    font_size_points: float = 10.0
    font_color: str = "000000"
    bullet_marker: str = "•"
    contact_separator: str = " | "


_FALLBACKS = _Fallbacks()


def render_structured_resume(
    resume: StructuredResume,
    layout_profile: LayoutProfile,
    output_path: Path,
) -> Path:
    renderer = AdaptiveStructuredResumeRenderer(layout_profile)
    return renderer.render(resume, output_path)


class AdaptiveStructuredResumeRenderer:
    def __init__(self, layout_profile: LayoutProfile) -> None:
        self._profile = layout_profile
        self._last_paragraph: Paragraph | None = None
        self._last_role: str | None = None
        self._numbering_ids: dict[str, int] = {}

    def render(self, resume: StructuredResume, output_path: Path) -> Path:
        self._last_paragraph = None
        self._last_role = None
        self._numbering_ids = {}
        required_roles = {"name", "contact_line", "section_heading"}
        missing = required_roles - set(self._profile.semantic_roles)
        if missing:
            raise AdaptiveDocxRenderError(
                f"LayoutProfile is missing required semantic roles: {sorted(missing)}"
            )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._apply_page_layout(document)
        self._add_text_paragraph(document, "name", [(resume.display_name, None)])
        if resume.contact_line:
            self._add_contact_line(document, resume.contact_line)
        if resume.education:
            self._add_section_heading(
                document,
                "Education",
                first_content_role="education_institution_date_row",
            )
            self._add_education(document, resume)
        if resume.technical_skills or resume.selected_skills:
            self._add_section_heading(
                document,
                "Technical Skills",
                first_content_role="skill_category_row",
            )
            self._add_skills(document, resume)
        if resume.experience_bullets:
            self._add_section_heading(
                document,
                "Technical Experience",
                first_content_role="experience_title_date_row",
            )
            self._add_experiences(document, resume)
        if resume.project_bullets:
            self._add_section_heading(
                document,
                "Projects",
                first_content_role="project_title_metadata_row",
            )
            self._add_projects(document, resume)
        document.save(output_path)
        return output_path

    def _apply_page_layout(self, document: DocumentType) -> None:
        page = self._profile.page
        section = document.sections[0]
        section.page_width = Twips(page.width_twips)
        section.page_height = Twips(page.height_twips)
        section.top_margin = Twips(page.top_margin_twips)
        section.bottom_margin = Twips(page.bottom_margin_twips)
        section.left_margin = Twips(page.left_margin_twips)
        section.right_margin = Twips(page.right_margin_twips)
        section.header_distance = Twips(page.header_distance_twips)
        section.footer_distance = Twips(page.footer_distance_twips)
        section.orientation = (
            WD_ORIENT.LANDSCAPE if page.orientation == "landscape" else WD_ORIENT.PORTRAIT
        )
        # python-docx swaps dimensions when orientation changes. Reapply the
        # profile dimensions after setting orientation so the profile remains
        # authoritative for both portrait and landscape references.
        section.page_width = Twips(page.width_twips)
        section.page_height = Twips(page.height_twips)
        section_properties = section._sectPr
        columns = section_properties.find(qn("w:cols"))
        if columns is None:
            columns = OxmlElement("w:cols")
            section_properties.append(columns)
        columns.set(qn("w:num"), str(page.column_count))
        if page.column_spacing_twips is not None:
            columns.set(qn("w:space"), str(page.column_spacing_twips))
        if page.column_equal_width is not None:
            columns.set(qn("w:equalWidth"), "1" if page.column_equal_width else "0")

    def _add_contact_line(self, document: DocumentType, contact_line: str) -> None:
        parts = [part.strip() for part in contact_line.split("|") if part.strip()]
        paragraph = self._new_paragraph(document, "contact_line")
        for index, part in enumerate(parts):
            if index:
                run = paragraph.add_run(_FALLBACKS.contact_separator)
                self._apply_run(run, self._role("contact_line"), index)
            target = _hyperlink_target(part)
            if target:
                self._add_hyperlink(paragraph, part, target, self._role("contact_line"))
            else:
                run = paragraph.add_run(part)
                self._apply_run(run, self._role("contact_line"), index)

    def _add_section_heading(
        self,
        document: DocumentType,
        label: str,
        *,
        first_content_role: str,
    ) -> None:
        self._add_text_paragraph(
            document,
            "section_heading",
            [(label, None)],
            transition_context=first_content_role,
        )

    def _add_education(self, document: DocumentType, resume: StructuredResume) -> None:
        for record in resume.education:
            date = _date_range(
                record.start_date,
                record.expected_graduation_date or record.graduation_date,
            )
            self._add_metadata_row(
                document,
                "education_institution_date_row",
                record.school,
                date,
            )
            self._add_metadata_row(
                document,
                "education_program_location_row",
                record.program,
                record.location,
            )
            details: list[str] = []
            if record.gpa:
                details.append(f"GPA: {record.gpa}")
            if record.awards:
                details.append(f"Awards: {', '.join(record.awards)}")
            coursework = record.relevant_coursework or resume.selected_coursework
            if coursework:
                details.append(f"Relevant Coursework: {', '.join(coursework)}")
            for detail_index, detail in enumerate(details):
                role = (
                    "final_paragraph_in_section"
                    if detail_index == len(details) - 1
                    else "education_detail_bullet"
                )
                self._add_bullet(
                    document,
                    role,
                    detail,
                    fallback_role="education_detail_bullet",
                )

    def _add_skills(self, document: DocumentType, resume: StructuredResume) -> None:
        if resume.technical_skills:
            for category in resume.technical_skills:
                values = category.values or [skill.value for skill in category.skills]
                if not values:
                    continue
                self._add_text_paragraph(
                    document,
                    "skill_category_row",
                    [(f"{category.category}: ", "label"), (", ".join(values), "values")],
                )
            return
        if resume.selected_skills:
            self._add_text_paragraph(
                document,
                "skill_category_row",
                [("Skills: ", "label"), (", ".join(resume.selected_skills), "values")],
            )

    def _add_experiences(self, document: DocumentType, resume: StructuredResume) -> None:
        records = _ordered_records(
            resume.experiences,
            resume.experience_bullets,
            resume.entity_titles,
            EntityKind.EXPERIENCE,
        )
        for entry_index, item in enumerate(records):
            role = "experience_title_date_row" if entry_index == 0 else "interior_entry_transition"
            left = item.title
            subtitle = item.subtitle or item.technology_label
            if subtitle and subtitle.casefold() not in left.casefold():
                left = f"{left} | {subtitle}"
            self._add_metadata_row(
                document,
                role,
                left,
                _date_range(item.start_date, item.end_date),
                layout_role="experience_title_date_row" if entry_index else None,
            )
            if item.organization or item.location:
                self._add_metadata_row(
                    document,
                    "employer_location_row",
                    item.organization,
                    item.location,
                )
            bullets = resume.experience_bullets.get(item.id, [])
            for bullet_index, bullet in enumerate(bullets):
                role = (
                    "final_paragraph_in_section"
                    if entry_index == len(records) - 1 and bullet_index == len(bullets) - 1
                    else "experience_bullet"
                )
                self._add_bullet(document, role, bullet.text, fallback_role="experience_bullet")

    def _add_projects(self, document: DocumentType, resume: StructuredResume) -> None:
        records = _ordered_records(
            resume.projects,
            resume.project_bullets,
            resume.entity_titles,
            EntityKind.PROJECT,
        )
        for entry_index, item in enumerate(records):
            left = item.title
            technology_label = item.technology_label or ", ".join(item.technologies)
            if technology_label and technology_label.casefold() not in left.casefold():
                left = f"{left} | {technology_label}"
            role = "project_title_metadata_row" if entry_index == 0 else "interior_entry_transition"
            self._add_metadata_row(
                document,
                role,
                left,
                _date_range(item.start_date, item.end_date),
                fallback_role="experience_title_date_row",
                layout_role="project_title_metadata_row" if entry_index else None,
            )
            bullets = resume.project_bullets.get(item.id, [])
            for bullet_index, bullet in enumerate(bullets):
                role = (
                    "final_paragraph_in_section"
                    if entry_index == len(records) - 1 and bullet_index == len(bullets) - 1
                    else "project_bullet"
                )
                self._add_bullet(document, role, bullet.text, fallback_role="experience_bullet")

    def _add_metadata_row(
        self,
        document: DocumentType,
        role_name: str,
        left: str | None,
        right: str | None,
        fallback_role: str | None = None,
        layout_role: str | None = None,
        transition_context: str | None = None,
    ) -> Paragraph | None:
        if not left and not right:
            return None
        paragraph = self._new_paragraph(
            document,
            role_name,
            fallback_role,
            layout_role=layout_role,
            transition_context=transition_context,
        )
        role = self._role(layout_role or role_name, fallback_role)
        if left:
            left_run = paragraph.add_run(left)
            self._apply_run(left_run, role, 0)
        if right:
            collision = left and self._metadata_needs_following_line(
                left,
                right,
                role,
            )
            if left and not collision:
                paragraph.add_run().add_tab()
            if collision:
                paragraph.add_run().add_break()
                if not self._metadata_right_fits(right, role):
                    right_run = paragraph.add_run(right)
                    self._apply_run(right_run, role, 1)
                    return paragraph
                paragraph.add_run().add_tab()
            right_run = paragraph.add_run(right)
            self._apply_run(right_run, role, 1)
        return paragraph

    def _add_bullet(
        self,
        document: DocumentType,
        role_name: str,
        text: str,
        fallback_role: str | None = None,
        transition_context: str | None = None,
    ) -> Paragraph:
        layout_role = fallback_role if role_name == "final_paragraph_in_section" else None
        paragraph = self._new_paragraph(
            document,
            role_name,
            fallback_role,
            layout_role=layout_role,
            transition_context=transition_context,
        )
        role = self._role(layout_role or role_name, fallback_role)
        if role.bullet and role.bullet.mechanism == "numbering":
            self._apply_numbering(paragraph, role.bullet)
        else:
            marker = _bullet_marker(role)
            marker_run = paragraph.add_run(f"{marker}\t")
            self._apply_run(
                marker_run,
                role,
                0,
                typography=role.bullet.marker_typography if role.bullet else None,
                apply_pattern=False,
            )
        text_run = paragraph.add_run(text)
        self._apply_run(
            text_run,
            role,
            0,
            typography=role.primary_typography,
            apply_pattern=False,
        )
        return paragraph

    def _add_text_paragraph(
        self,
        document: DocumentType,
        role_name: str,
        pieces: list[tuple[str, str | None]],
        fallback_role: str | None = None,
        transition_context: str | None = None,
    ) -> Paragraph:
        paragraph = self._new_paragraph(
            document,
            role_name,
            fallback_role,
            transition_context=transition_context,
        )
        role = self._role(role_name, fallback_role)
        for index, (text, semantic) in enumerate(pieces):
            run = paragraph.add_run(text)
            if semantic == "values":
                self._apply_run(
                    run,
                    role,
                    index,
                    typography=_regular_typography(role),
                    apply_pattern=False,
                )
            elif semantic == "label":
                self._apply_run(run, role, index, apply_pattern=False)
            else:
                self._apply_run(run, role, index)
            if semantic == "label":
                run.bold = True
        return paragraph

    def _new_paragraph(
        self,
        document: DocumentType,
        role_name: str,
        fallback_role: str | None = None,
        layout_role: str | None = None,
        transition_context: str | None = None,
    ) -> Paragraph:
        paragraph = document.add_paragraph()
        role = self._role(layout_role or role_name, fallback_role)
        geometry_role_name, geometry_role = self._geometry_role(
            role_name,
            layout_role,
            fallback_role,
            role,
        )
        self._apply_paragraph(paragraph, geometry_role, geometry_role_name)
        self._apply_transition(paragraph, role_name, transition_context)
        self._last_paragraph = paragraph
        self._last_role = role_name
        return paragraph

    def _geometry_role(
        self,
        role_name: str,
        layout_role: str | None,
        fallback_role: str | None,
        role: SemanticRoleLayout,
    ) -> tuple[str, SemanticRoleLayout]:
        candidate_name = layout_role or role_name
        candidate = role
        if not candidate.metadata_anchor_group_ids and fallback_role:
            fallback = self._profile.semantic_roles.get(fallback_role)
            if fallback is not None and fallback.metadata_anchor_group_ids:
                return fallback_role, fallback
        return candidate_name, candidate

    def _role(
        self,
        role_name: str,
        fallback_role: str | None = None,
    ) -> SemanticRoleLayout:
        role = self._profile.semantic_roles.get(role_name)
        if role is None and fallback_role:
            role = self._profile.semantic_roles.get(fallback_role)
        if role is None:
            raise AdaptiveDocxRenderError(
                f"LayoutProfile is missing semantic role {role_name!r}"
            )
        return role

    def _apply_transition(
        self,
        destination: Paragraph,
        destination_role: str,
        destination_section_first_role: str | None,
    ) -> None:
        if self._last_paragraph is None or self._last_role is None:
            return
        matches = self._transition_matches(
            self._last_role,
            destination_role,
            destination_section_first_role,
        )
        if not matches:
            return
        transition = max(matches, key=lambda item: item.occurrence_count)
        source_spacing = (
            transition.resolved_source_space_after_twips
            or transition.source_space_after_twips
        )
        destination_spacing = (
            transition.resolved_destination_space_before_twips
            or transition.destination_space_before_twips
        )
        _apply_transition_spacing(
            self._last_paragraph,
            "space_after",
            source_spacing.value,
        )
        _apply_transition_spacing(
            destination,
            "space_before",
            destination_spacing.value,
        )

    def _transition_matches(
        self,
        source_role: str,
        destination_role: str,
        destination_section_first_role: str | None,
    ) -> list[TransitionSpacing]:
        matches = [
            transition
            for transition in self._profile.transition_spacings
            if transition.source_role == source_role
            and transition.destination_role == destination_role
        ]
        if destination_role == "section_heading" and matches:
            contextual = [
                transition
                for transition in matches
                if transition.destination_section_first_role == destination_section_first_role
            ]
            if contextual:
                return contextual
        if matches:
            return matches

        # Some compact reference blocks do not repeat an exact semantic role
        # pair (for example, an education detail followed by the final detail
        # role). Reuse the nearest observed bullet transition only when both
        # roles have the same structural bullet contract.
        if self._has_bullet_role(source_role) and self._has_bullet_role(destination_role):
            return [
                transition
                for transition in self._profile.transition_spacings
                if transition.source_role == source_role
                and self._has_bullet_role(transition.destination_role)
                and transition.destination_role != "section_heading"
            ]
        return []

    def _has_bullet_role(self, role_name: str) -> bool:
        role = self._profile.semantic_roles.get(role_name)
        return role is not None and role.bullet is not None

    def _apply_paragraph(
        self,
        paragraph: Paragraph,
        role: SemanticRoleLayout,
        role_name: str,
    ) -> None:
        layout = role.paragraph
        paragraph.alignment = _paragraph_alignment(layout.alignment.value)
        formatting = paragraph.paragraph_format
        _set_twips(formatting, "space_before", layout.space_before_twips.value)
        _set_twips(formatting, "space_after", layout.space_after_twips.value)
        _set_twips(formatting, "left_indent", layout.left_indent_twips.value)
        _set_twips(formatting, "right_indent", layout.right_indent_twips.value)
        if _is_twips(layout.hanging_indent_twips.value):
            formatting.first_line_indent = Twips(-layout.hanging_indent_twips.value)
        elif _is_twips(layout.first_line_indent_twips.value):
            formatting.first_line_indent = Twips(layout.first_line_indent_twips.value)
        if _is_twips(layout.line_spacing_twips.value):
            formatting.line_spacing = Twips(layout.line_spacing_twips.value)
        formatting.line_spacing_rule = _line_spacing_rule(layout.line_spacing_rule.value)
        formatting.keep_with_next = _optional_bool(layout.keep_with_next.value)
        formatting.keep_together = _optional_bool(layout.keep_together.value)
        formatting.widow_control = _optional_bool(layout.widow_control.value)
        formatting.page_break_before = _optional_bool(layout.page_break_before.value)
        formatting.tab_stops.clear_all()
        for tab in self._effective_tab_stops(role, role_name):
            formatting.tab_stops.add_tab_stop(
                Twips(tab.position_twips),
                _tab_alignment(tab),
                _tab_leader(tab.leader),
            )
        if role.bullet is not None:
            if role.bullet.left_indent_twips is not None:
                formatting.left_indent = Twips(role.bullet.left_indent_twips)
            if role.bullet.hanging_indent_twips is not None:
                formatting.first_line_indent = Twips(-role.bullet.hanging_indent_twips)
            if role.bullet.space_before_twips is not None:
                formatting.space_before = Twips(role.bullet.space_before_twips)
            if role.bullet.space_after_twips is not None:
                formatting.space_after = Twips(role.bullet.space_after_twips)
        self._apply_borders(paragraph, role.borders)
        _apply_spacing_flags(paragraph, layout)

    def _apply_run(
        self,
        run: Run,
        role: SemanticRoleLayout,
        index: int,
        *,
        typography: Typography | None = None,
        apply_pattern: bool = True,
    ) -> None:
        typography = typography or _typography_for_run(role, index)
        family = typography.font_family.value or _FALLBACKS.font_family
        size = typography.font_size_half_points.value
        run.font.name = str(family)
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), str(family))
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), str(family))
        run.font.size = Pt(float(size) / 2 if isinstance(size, (int, float)) else _FALLBACKS.font_size_points)
        run.bold = _optional_bool(typography.bold.value)
        run.italic = _optional_bool(typography.italic.value)
        underline = typography.underline.value
        run.underline = False if underline in {None, False, "none"} else True
        color = typography.color.value
        color_text = _valid_rgb(color) or _FALLBACKS.font_color
        run.font.color.rgb = RGBColor.from_string(color_text)
        spacing = typography.character_spacing_twips.value
        if _is_twips(spacing):
            spacing_element = OxmlElement("w:spacing")
            spacing_element.set(qn("w:val"), str(spacing))
            run._element.get_or_add_rPr().append(spacing_element)
        pattern = role.run_patterns[0] if role.run_patterns else None
        if apply_pattern and pattern and index in pattern.bold_run_positions:
            run.bold = True
        if apply_pattern and pattern and index in pattern.italic_run_positions:
            run.italic = True

    def _effective_tab_stops(
        self,
        role: SemanticRoleLayout,
        role_name: str,
    ) -> list[TabStop]:
        anchor_groups = [
            group
            for group in self._profile.metadata_anchor_groups
            if role_name in group.role_groups
        ]
        if anchor_groups:
            group = max(
                anchor_groups,
                key=lambda item: item.representative_position_twips,
            )
            candidates = [
                tab
                for tab in role.tab_stops
                if tab.position_twips in group.observed_positions_twips
            ]
            representative = min(
                candidates or role.tab_stops,
                key=lambda tab: (
                    abs(tab.position_twips - group.representative_position_twips),
                    tab.position_twips,
                ),
            ) if role.tab_stops else None
            if representative is not None:
                return [
                    representative.model_copy(
                        update={
                            "position_twips": group.representative_position_twips,
                            "semantic_use": "normalized_metadata_column",
                            "provenance": "inferred_recurring_pattern",
                        }
                    )
                ]

        metadata_tabs = [
            tab
            for candidate in self._profile.semantic_roles.values()
            for tab in candidate.tab_stops
            if tab.semantic_use in {
                "right_aligned_metadata",
                "positioned_metadata_column",
            }
        ]
        role_has_metadata = any(
            tab.semantic_use in {
                "right_aligned_metadata",
                "positioned_metadata_column",
            }
            for tab in role.tab_stops
        )
        if not metadata_tabs:
            return role.tab_stops
        if not role_has_metadata and len(role.tab_stops) <= 1:
            return role.tab_stops
        positions = sorted(tab.position_twips for tab in metadata_tabs)
        canonical_position = positions[len(positions) // 2]
        representative = min(
            metadata_tabs,
            key=lambda tab: (abs(tab.position_twips - canonical_position), tab.position_twips),
        )
        return [
            representative.model_copy(
                update={
                    "position_twips": canonical_position,
                    "semantic_use": "normalized_metadata_column",
                    "provenance": "inferred_recurring_pattern",
                }
            )
        ]

    def _metadata_needs_following_line(
        self,
        left: str,
        right: str,
        role: SemanticRoleLayout,
    ) -> bool:
        anchor = self._metadata_anchor(role)
        if anchor is None:
            return False
        left_indent = role.paragraph.left_indent_twips.value
        if not isinstance(left_indent, int):
            left_indent = 0
        left_end = left_indent + _estimated_text_width_twips(left, role)
        right_fits = anchor + _estimated_text_width_twips(right, role) <= self._profile.page.usable_width_twips
        return left_end + _metadata_collision_tolerance(self._profile) > anchor or not right_fits

    def _metadata_right_fits(self, right: str, role: SemanticRoleLayout) -> bool:
        anchor = self._metadata_anchor(role)
        return anchor is not None and (
            anchor + _estimated_text_width_twips(right, role)
            <= self._profile.page.usable_width_twips
        )

    def _metadata_anchor(self, role: SemanticRoleLayout) -> int | None:
        positions = [
            group.representative_position_twips
            for group in self._profile.metadata_anchor_groups
            if any(
                tab.position_twips in group.observed_positions_twips
                for tab in role.tab_stops
            )
        ]
        if positions:
            return max(positions)
        metadata_tabs = [
            tab.position_twips
            for tab in role.tab_stops
            if tab.semantic_use in {
                "right_aligned_metadata",
                "positioned_metadata_column",
            }
        ]
        return max(metadata_tabs) if metadata_tabs else None

    def _apply_numbering(self, paragraph: Paragraph, bullet: BulletLayout) -> None:
        signature = bullet.model_dump_json()
        numbering_id = self._numbering_ids.get(signature)
        if numbering_id is None:
            numbering = paragraph.part.numbering_part.element
            abstract_ids = [
                int(element.get(qn("w:abstractNumId")))
                for element in numbering.findall(qn("w:abstractNum"))
            ]
            num_ids = [
                int(element.get(qn("w:numId")))
                for element in numbering.findall(qn("w:num"))
            ]
            abstract_id = max(abstract_ids, default=-1) + 1
            numbering_id = max(num_ids, default=0) + 1
            abstract = OxmlElement("w:abstractNum")
            abstract.set(qn("w:abstractNumId"), str(abstract_id))
            level = OxmlElement("w:lvl")
            level.set(qn("w:ilvl"), str(bullet.numbering_level or 0))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            number_format = OxmlElement("w:numFmt")
            number_format.set(qn("w:val"), bullet.numbering_format or "bullet")
            level_text = OxmlElement("w:lvlText")
            level_text.set(qn("w:val"), bullet.representation or _FALLBACKS.bullet_marker)
            justification = OxmlElement("w:lvlJc")
            justification.set(qn("w:val"), "left")
            level.extend([start, number_format, level_text, justification])
            if bullet.left_indent_twips is not None or bullet.hanging_indent_twips is not None:
                level_p_pr = OxmlElement("w:pPr")
                indent = OxmlElement("w:ind")
                if bullet.left_indent_twips is not None:
                    indent.set(qn("w:left"), str(bullet.left_indent_twips))
                if bullet.hanging_indent_twips is not None:
                    indent.set(qn("w:hanging"), str(bullet.hanging_indent_twips))
                level_p_pr.append(indent)
                level.append(level_p_pr)
            if bullet.marker_typography is not None:
                level.append(_numbering_run_properties(bullet.marker_typography))
            abstract.append(level)
            numbering.append(abstract)
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(numbering_id))
            abstract_reference = OxmlElement("w:abstractNumId")
            abstract_reference.set(qn("w:val"), str(abstract_id))
            num.append(abstract_reference)
            numbering.append(num)
            self._numbering_ids[signature] = numbering_id
        properties = paragraph._p.get_or_add_pPr()
        num_pr = properties.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            properties.append(num_pr)
        level_reference = OxmlElement("w:ilvl")
        level_reference.set(qn("w:val"), str(bullet.numbering_level or 0))
        num_reference = OxmlElement("w:numId")
        num_reference.set(qn("w:val"), str(numbering_id))
        num_pr.extend([level_reference, num_reference])

    def _apply_borders(self, paragraph: Paragraph, borders: list[Border]) -> None:
        if not borders:
            return
        properties = paragraph._p.get_or_add_pPr()
        existing = properties.find(qn("w:pBdr"))
        if existing is not None:
            properties.remove(existing)
        container = OxmlElement("w:pBdr")
        for border in borders:
            element = OxmlElement(f"w:{border.position}")
            style = "single" if border.style == "solid" else border.style
            element.set(qn("w:val"), style)
            if border.thickness_eighth_points is not None:
                element.set(qn("w:sz"), str(border.thickness_eighth_points))
            if border.spacing_points is not None:
                element.set(qn("w:space"), str(border.spacing_points))
            element.set(qn("w:color"), _valid_rgb(border.color) or _FALLBACKS.font_color)
            container.append(element)
        properties.append(container)

    def _add_hyperlink(
        self,
        paragraph: Paragraph,
        display: str,
        target: str,
        role: SemanticRoleLayout,
    ) -> None:
        relationship_id = paragraph.part.relate_to(
            target,
            RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run_element = OxmlElement("w:r")
        text_element = OxmlElement("w:t")
        text_element.text = display
        run_element.append(text_element)
        hyperlink.append(run_element)
        paragraph._p.append(hyperlink)
        run = Run(run_element, paragraph)
        hyperlink_typography = (
            role.hyperlinks.display_typography
            if role.hyperlinks and role.hyperlinks.display_typography
            else role.primary_typography
        )
        hyperlink_role = role.model_copy(
            update={"primary_typography": hyperlink_typography, "run_patterns": []}
        )
        self._apply_run(run, hyperlink_role, 0)
        if role.hyperlinks:
            run.underline = role.hyperlinks.underline_behavior not in {None, "none", "False"}
            color = _valid_rgb(role.hyperlinks.color_behavior)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def _ordered_records(
    records: list[ResumeItem],
    bullets: dict[str, list[StructuredBullet]],
    titles: dict[str, str],
    kind: EntityKind,
) -> list[ResumeItem]:
    by_id = {item.id: item for item in records}
    ordered = [by_id[entity_id] for entity_id in bullets if entity_id in by_id]
    for entity_id in bullets:
        if entity_id not in by_id:
            ordered.append(
                ResumeItem(
                    id=entity_id,
                    title=titles.get(entity_id, entity_id),
                    kind=kind,
                )
            )
    return ordered


def _date_range(start: str | None, end: str | None) -> str | None:
    return " - ".join(part for part in (start, end) if part) or None


def _hyperlink_target(value: str) -> str | None:
    stripped = value.strip()
    if "@" in stripped and not re.search(r"\s", stripped):
        return stripped if stripped.casefold().startswith("mailto:") else f"mailto:{stripped}"
    if stripped.casefold().startswith(("http://", "https://")):
        return stripped
    if re.match(r"^(www\.|[a-z0-9-]+\.[a-z]{2,}/)", stripped, re.IGNORECASE):
        return f"https://{stripped}"
    return None


def _paragraph_alignment(value: object) -> WD_ALIGN_PARAGRAPH | None:
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(str(value)) if value is not None else None


def _line_spacing_rule(value: object) -> WD_LINE_SPACING | None:
    return {
        "exact": WD_LINE_SPACING.EXACTLY,
        "atLeast": WD_LINE_SPACING.AT_LEAST,
        "auto": WD_LINE_SPACING.SINGLE,
    }.get(str(value)) if value is not None else None


def _tab_alignment(tab: TabStop) -> WD_TAB_ALIGNMENT:
    return {
        "right": WD_TAB_ALIGNMENT.RIGHT,
        "center": WD_TAB_ALIGNMENT.CENTER,
        "decimal": WD_TAB_ALIGNMENT.DECIMAL,
        "bar": WD_TAB_ALIGNMENT.BAR,
    }.get(tab.alignment, WD_TAB_ALIGNMENT.LEFT)


def _tab_leader(value: str | None) -> WD_TAB_LEADER:
    return {
        "dot": WD_TAB_LEADER.DOTS,
        "hyphen": WD_TAB_LEADER.DASHES,
        "underscore": WD_TAB_LEADER.LINES,
        "heavy": WD_TAB_LEADER.HEAVY,
        "middleDot": WD_TAB_LEADER.MIDDLE_DOT,
    }.get(value or "", WD_TAB_LEADER.SPACES)


def _typography_for_run(role: SemanticRoleLayout, index: int) -> Typography:
    if role.run_patterns:
        variants = role.run_patterns[0].typography_variants
        if index < len(variants):
            return variants[index]
    return role.primary_typography


def _regular_typography(role: SemanticRoleLayout) -> Typography:
    for pattern in role.run_patterns:
        for typography in pattern.typography_variants:
            if typography.bold.value is False and typography.italic.value is False:
                return typography
    return role.primary_typography


def _bullet_marker(role: SemanticRoleLayout) -> str:
    marker = role.bullet.representation if role.bullet else None
    if marker and "%" not in marker and len(marker) <= 3:
        return marker
    return _FALLBACKS.bullet_marker


def _set_twips(target: object, attribute: str, value: object) -> None:
    if _is_twips(value):
        setattr(target, attribute, Twips(value))


def _apply_transition_spacing(
    paragraph: Paragraph,
    attribute: str,
    value: object,
) -> None:
    formatting = paragraph.paragraph_format
    if _is_twips(value):
        setattr(formatting, attribute, Twips(value))
    else:
        # An absent reference contribution means no direct paragraph spacing,
        # not permission to retain the role representative's spacing.
        setattr(formatting, attribute, None)


def _apply_spacing_flags(paragraph: Paragraph, layout: ParagraphLayout) -> None:
    properties = paragraph._p.get_or_add_pPr()
    spacing_values = {
        "beforeAutospacing": layout.before_auto_spacing.value,
        "afterAutospacing": layout.after_auto_spacing.value,
    }
    if any(isinstance(value, bool) for value in spacing_values.values()):
        spacing = properties.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            properties.append(spacing)
        for attribute, value in spacing_values.items():
            if isinstance(value, bool):
                spacing.set(qn(f"w:{attribute}"), "1" if value else "0")

    contextual = layout.contextual_spacing.value
    existing = properties.find(qn("w:contextualSpacing"))
    if isinstance(contextual, bool):
        if existing is None:
            existing = OxmlElement("w:contextualSpacing")
            properties.append(existing)
        existing.set(qn("w:val"), "1" if contextual else "0")


def _is_twips(value: object) -> bool:
    return isinstance(value, int) and not isinstance(value, bool)


def _numbering_run_properties(typography: Typography) -> OxmlElement:
    properties = OxmlElement("w:rPr")
    family = typography.font_family.value
    if isinstance(family, str) and family:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), family)
        fonts.set(qn("w:hAnsi"), family)
        fonts.set(qn("w:cs"), family)
        properties.append(fonts)
    size = typography.font_size_half_points.value
    if isinstance(size, (int, float)) and not isinstance(size, bool):
        size_element = OxmlElement("w:sz")
        size_element.set(qn("w:val"), str(size))
        properties.append(size_element)
    if typography.bold.value is True:
        properties.append(OxmlElement("w:b"))
    if typography.italic.value is True:
        properties.append(OxmlElement("w:i"))
    spacing = typography.character_spacing_twips.value
    if _is_twips(spacing):
        spacing_element = OxmlElement("w:spacing")
        spacing_element.set(qn("w:val"), str(spacing))
        properties.append(spacing_element)
    return properties


def _estimated_text_width_twips(text: str, role: SemanticRoleLayout) -> int:
    size = role.primary_typography.font_size_half_points.value
    half_points = size if isinstance(size, (int, float)) and not isinstance(size, bool) else 20
    return round(len(text) * half_points * 5)


def _metadata_collision_tolerance(profile: LayoutProfile) -> int:
    return max(1, round(profile.page.usable_width_twips * 0.01))


def _optional_bool(value: object) -> bool | None:
    return value if isinstance(value, bool) else None


def _valid_rgb(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    normalized = value.removeprefix("#").upper()
    return normalized if re.fullmatch(r"[0-9A-F]{6}", normalized) else None
