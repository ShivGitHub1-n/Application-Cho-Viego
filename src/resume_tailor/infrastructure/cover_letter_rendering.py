from __future__ import annotations

import math
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from resume_tailor.domain.cover_letter import CoverLetter, CoverLetterLayoutProfile
from resume_tailor.domain.layout import LayoutProfile, PageLayout
from resume_tailor.infrastructure.rendering import (
    DocxPageCountProvider,
    ExactDocxPageCountProvider,
    PageCountMeasurement,
    PageCountVerificationError,
    diagnose_docx_page_utilization,
)


@dataclass(frozen=True)
class CoverLetterRenderResult:
    docx_path: Path
    docx_bytes: bytes
    measurement: PageCountMeasurement
    estimated_utilization: float
    estimated_remaining_lines: int
    pagination_failure: str | None = None
    blank_trailing_page: bool = False
    structural_issues: tuple[str, ...] = ()

    @property
    def page_count(self) -> int:
        return self.measurement.page_count


class CoverLetterRenderer:
    """Render fixed-geometry correspondence and measure a bounded candidate batch."""

    def __init__(
        self,
        layout_profile: CoverLetterLayoutProfile | None = None,
        page_count_provider: DocxPageCountProvider | None = None,
    ) -> None:
        self._profile = layout_profile or CoverLetterLayoutProfile()
        self._page_count_provider = page_count_provider or ExactDocxPageCountProvider()
        self.pagination_attempt_count = 0

    @property
    def layout_profile(self) -> CoverLetterLayoutProfile:
        return self._profile

    def render_candidates(
        self,
        letters: list[CoverLetter],
        output_directory: Path,
    ) -> list[CoverLetterRenderResult]:
        if not letters:
            return []
        output_directory = Path(output_directory).expanduser().resolve()
        output_directory.mkdir(parents=True, exist_ok=True)
        paths: list[Path] = []
        for index, letter in enumerate(letters):
            path = output_directory / f"cover-letter-{index:02d}-{uuid4().hex}.docx"
            self._render_docx(letter, path)
            if not path.is_file() or path.stat().st_size <= 0:
                raise PageCountVerificationError(
                    "Cover-letter renderer did not produce a usable DOCX."
                )
            paths.append(path)
        failure: str | None = None
        self.pagination_attempt_count += 1
        try:
            measure_many = getattr(self._page_count_provider, "measure_many", None)
            measurements = (
                list(measure_many(paths))
                if callable(measure_many)
                else [self._page_count_provider.measure(path) for path in paths]
            )
            if len(measurements) != len(paths):
                raise PageCountVerificationError(
                    "Exact page-count provider returned an incomplete cover-letter batch."
                )
        except PageCountVerificationError as error:
            failure = str(error)
            measurements = [self._estimated_measurement(path) for path in paths]
        return [
            self._result(path, measurement, pagination_failure=failure)
            for path, measurement in zip(paths, measurements, strict=True)
        ]

    def render_candidate(
        self,
        letter: CoverLetter,
        output_directory: Path,
    ) -> CoverLetterRenderResult:
        return self.render_candidates([letter], output_directory)[0]

    def render(self, letter: CoverLetter, output_directory: Path) -> CoverLetterRenderResult:
        result = self.render_candidate(letter, output_directory)
        if not result.measurement.exact or result.measurement.page_count != 1:
            raise PageCountVerificationError(
                "Cover letter did not receive exact one-page verification."
            )
        if result.blank_trailing_page:
            raise PageCountVerificationError("Cover letter contains a blank trailing page.")
        return result

    def _result(
        self,
        path: Path,
        measurement: PageCountMeasurement,
        *,
        pagination_failure: str | None,
    ) -> CoverLetterRenderResult:
        diagnostic = diagnose_docx_page_utilization(
            path,
            self._diagnostic_layout(),
            measurement,
            severe_underfill_threshold=self._profile.acceptable_utilization_floor,
        )
        remaining = self._remaining_lines(diagnostic.estimated_utilization_ratio)
        structural = tuple(audit_cover_letter_docx_structure(path))
        blank_trailing = bool(
            measurement.page_count > 1
            and (
                diagnostic.estimated_utilization_ratio <= 1.0
                or "trailing_blank_paragraph" in structural
                or "explicit_page_break" in structural
                or "final_keep_with_next" in structural
            )
        )
        return CoverLetterRenderResult(
            docx_path=path,
            docx_bytes=path.read_bytes(),
            measurement=measurement,
            estimated_utilization=diagnostic.estimated_utilization_ratio,
            estimated_remaining_lines=remaining,
            pagination_failure=pagination_failure,
            blank_trailing_page=blank_trailing,
            structural_issues=structural,
        )

    def _estimated_measurement(self, path: Path) -> PageCountMeasurement:
        provisional = PageCountMeasurement(
            page_count=1,
            provider="deterministic cover-letter occupancy estimate",
            confidence="estimated",
            exact=False,
        )
        diagnostic = diagnose_docx_page_utilization(
            path,
            self._diagnostic_layout(),
            provisional,
            severe_underfill_threshold=self._profile.acceptable_utilization_floor,
        )
        return PageCountMeasurement(
            page_count=max(1, math.ceil(diagnostic.estimated_utilization_ratio)),
            provider=provisional.provider,
            confidence=provisional.confidence,
            exact=False,
        )

    def _diagnostic_layout(self) -> LayoutProfile:
        p = self._profile
        return LayoutProfile(
            page=PageLayout(
                width_twips=round(p.page_width_inches * 1440),
                height_twips=round(p.page_height_inches * 1440),
                top_margin_twips=round(p.top_margin_inches * 1440),
                bottom_margin_twips=round(p.bottom_margin_inches * 1440),
                left_margin_twips=round(p.left_margin_inches * 1440),
                right_margin_twips=round(p.right_margin_inches * 1440),
                header_distance_twips=round(p.header_distance_inches * 1440),
                footer_distance_twips=round(p.footer_distance_inches * 1440),
                orientation="portrait",
            ),
            semantic_roles={},
        )

    def _remaining_lines(self, utilization: float) -> int:
        p = self._profile
        line_height_twips = p.body_size_pt * 20 * p.line_spacing
        remaining_twips = max(0.0, (1.0 - utilization) * p.usable_height_inches * 1440)
        return max(0, round(remaining_twips / line_height_twips))

    def _render_docx(self, letter: CoverLetter, output_path: Path) -> None:
        p = self._profile
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(p.page_width_inches)
        section.page_height = Inches(p.page_height_inches)
        section.top_margin = Inches(p.top_margin_inches)
        section.bottom_margin = Inches(p.bottom_margin_inches)
        section.left_margin = Inches(p.left_margin_inches)
        section.right_margin = Inches(p.right_margin_inches)
        section.header_distance = Inches(p.header_distance_inches)
        section.footer_distance = Inches(p.footer_distance_inches)
        self._configure_styles(document, p)

        name = document.add_paragraph(style="Cover Letter Name")
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.paragraph_format.space_before = Pt(0)
        name.paragraph_format.space_after = Pt(p.contact_spacing_pt)
        name.paragraph_format.line_spacing = 1.0
        self._add_run(
            name,
            letter.candidate_name,
            bold=True,
            size=p.candidate_name_size_pt,
            font_family=p.body_font,
        )

        for line in self._contact_lines(letter):
            contact = document.add_paragraph(style="Cover Letter Contact")
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_before = Pt(0)
            contact.paragraph_format.space_after = Pt(0)
            contact.paragraph_format.line_spacing = 1.0
            for index, part in enumerate(line):
                if index:
                    self._add_run(
                        contact,
                        p.contact_separator,
                        size=p.contact_size_pt,
                        font_family=p.body_font,
                    )
                if isinstance(part, tuple):
                    self._add_hyperlink(
                        contact,
                        part[0],
                        part[1],
                        p.contact_size_pt,
                        p.body_font,
                    )
                else:
                    self._add_run(
                        contact,
                        part,
                        size=p.contact_size_pt,
                        font_family=p.body_font,
                    )

        self._add_body(document, letter.date_text)
        recipient_values = [
            letter.recipient.name,
            letter.recipient.title,
            letter.recipient.company,
            *letter.recipient.address_lines,
        ]
        for value in recipient_values:
            if value:
                self._add_body(document, value, after=0)
        self._add_body(document, letter.salutation)
        for paragraph in letter.paragraphs:
            self._add_body(document, paragraph.text)
        self._add_body(
            document,
            letter.signoff,
            after=p.signoff_spacing_pt,
            keep_with_next=True,
        )
        signoff_name = document.add_paragraph(style="Cover Letter Body")
        signoff_name.paragraph_format.space_before = Pt(0)
        signoff_name.paragraph_format.space_after = Pt(p.paragraph_spacing_pt)
        signoff_name.paragraph_format.line_spacing = p.line_spacing
        signoff_name.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        self._add_run(
            signoff_name,
            letter.signoff_name,
            bold=True,
            size=p.body_size_pt,
            font_family=p.body_font,
        )
        document.core_properties.title = f"Cover letter - {letter.job_title}"
        document.core_properties.subject = "Evidence-grounded job application cover letter"
        document.core_properties.author = letter.candidate_name
        document.save(str(output_path))

    @staticmethod
    def _configure_styles(
        document: DocxDocument,
        profile: CoverLetterLayoutProfile,
    ) -> None:
        normal = document.styles["Normal"]
        normal.font.name = profile.body_font
        normal.font.size = Pt(profile.body_size_pt)
        normal.font.color.rgb = None
        normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
        normal_fonts.set(qn("w:ascii"), profile.body_font)
        normal_fonts.set(qn("w:hAnsi"), profile.body_font)
        normal_fonts.set(qn("w:eastAsia"), profile.body_font)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(profile.paragraph_spacing_pt)
        normal.paragraph_format.line_spacing = profile.line_spacing
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        styles = document.styles
        name = styles.add_style("Cover Letter Name", WD_STYLE_TYPE.PARAGRAPH)
        name.base_style = normal
        name.font.name = profile.body_font
        name.font.size = Pt(profile.candidate_name_size_pt)
        name.font.bold = True
        name.paragraph_format.space_before = Pt(0)
        name.paragraph_format.space_after = Pt(profile.contact_spacing_pt)
        name.paragraph_format.line_spacing = 1.0

        contact = styles.add_style("Cover Letter Contact", WD_STYLE_TYPE.PARAGRAPH)
        contact.base_style = normal
        contact.font.name = profile.body_font
        contact.font.size = Pt(profile.contact_size_pt)
        contact.paragraph_format.space_before = Pt(0)
        contact.paragraph_format.space_after = Pt(0)
        contact.paragraph_format.line_spacing = 1.0

        body = styles.add_style("Cover Letter Body", WD_STYLE_TYPE.PARAGRAPH)
        body.base_style = normal
        body.font.name = profile.body_font
        body.font.size = Pt(profile.body_size_pt)
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(profile.paragraph_spacing_pt)
        body.paragraph_format.line_spacing = profile.line_spacing
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    def _add_body(
        self,
        document: DocxDocument,
        text: str,
        *,
        after: float | None = None,
        keep_with_next: bool = False,
    ) -> Paragraph:
        profile = self._profile
        paragraph = document.add_paragraph(style="Cover Letter Body")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(
            profile.paragraph_spacing_pt if after is None else after
        )
        paragraph.paragraph_format.line_spacing = profile.line_spacing
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.keep_with_next = keep_with_next
        self._add_run(
            paragraph,
            text,
            size=profile.body_size_pt,
            font_family=profile.body_font,
        )
        return paragraph

    def _contact_lines(self, letter: CoverLetter) -> list[list[str | tuple[str, str]]]:
        values: list[str | tuple[str, str]] = []
        seen_values: set[str] = set()
        for value in (letter.contact.location, letter.contact.phone, letter.contact.email):
            cleaned = (value or "").strip()
            key = cleaned.casefold()
            if cleaned and key not in seen_values:
                values.append(cleaned)
                seen_values.add(key)
        used_labels: set[str] = set()
        for link in letter.contact.links:
            cleaned = link.strip()
            key = cleaned.casefold().rstrip("/")
            if not cleaned or key in seen_values:
                continue
            label = self._unique_link_label(cleaned, used_labels)
            values.append((label, cleaned))
            seen_values.add(key)
            used_labels.add(label.casefold())
        if not values:
            return []
        rough_length = sum(
            len(value[0] if isinstance(value, tuple) else value) for value in values
        ) + len(self._profile.contact_separator) * (len(values) - 1)
        if rough_length <= 100:
            return [values]
        midpoint = max(1, (len(values) + 1) // 2)
        return [values[:midpoint], values[midpoint:]]

    @staticmethod
    def _unique_link_label(link: str, used: set[str]) -> str:
        host = (urlparse(link).hostname or "").casefold().removeprefix("www.")
        if "linkedin" in host:
            base = "LinkedIn"
        elif "github" in host:
            base = "GitHub"
        elif host:
            base = host
        else:
            base = link.rstrip("/")
        if base.casefold() not in used:
            return base
        parsed = urlparse(link)
        path_label = parsed.path.strip("/")
        alternate = f"{host}/{path_label}" if host and path_label else link.rstrip("/")
        if alternate.casefold() not in used:
            return alternate
        index = 2
        while f"{alternate} {index}".casefold() in used:
            index += 1
        return f"{alternate} {index}"

    @staticmethod
    def _add_run(
        paragraph: Paragraph,
        text: str,
        *,
        bold: bool = False,
        size: float,
        font_family: str,
    ) -> Run:
        run = paragraph.add_run(text)
        run.font.name = font_family
        run_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        run_fonts.set(qn("w:ascii"), font_family)
        run_fonts.set(qn("w:hAnsi"), font_family)
        run_fonts.set(qn("w:eastAsia"), font_family)
        run.font.size = Pt(size)
        run.bold = bold
        return run

    @staticmethod
    def _add_hyperlink(
        paragraph: Paragraph,
        label: str,
        url: str,
        size: float,
        font_family: str,
    ) -> None:
        relationship = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship)
        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font_family)
        fonts.set(qn("w:hAnsi"), font_family)
        properties.append(fonts)
        size_element = OxmlElement("w:sz")
        size_element.set(qn("w:val"), str(round(size * 2)))
        properties.append(size_element)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        properties.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(underline)
        run.append(properties)
        text = OxmlElement("w:t")
        text.text = label
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)


def audit_cover_letter_docx_structure(path: Path) -> list[str]:
    document = Document(str(path))
    issues: list[str] = []
    xml = document._element.xml
    if 'w:type="page"' in xml:
        issues.append("explicit_page_break")
    nonblank = [
        index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip()
    ]
    if nonblank and nonblank[-1] < len(document.paragraphs) - 1:
        issues.append("trailing_blank_paragraph")
    if document.paragraphs:
        final_properties = document.paragraphs[-1]._p.pPr
        if final_properties is not None and final_properties.find(qn("w:keepNext")) is not None:
            issues.append("final_keep_with_next")
    if xml.count("<w:sectPr") > 1:
        issues.append("multiple_section_breaks")
    return issues


__all__ = [
    "CoverLetterRenderResult",
    "CoverLetterRenderer",
    "audit_cover_letter_docx_structure",
]
